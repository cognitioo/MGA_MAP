"""
Document Data Extractor
Extracts project data from PDF, DOCX, and XLSX files to auto-fill forms
"""

import os
import json
import re
from io import BytesIO
from typing import Dict, Optional, Any

# PDF parsing
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# DOCX parsing
try:
    from docx import Document
except ImportError:
    Document = None

# XLSX parsing
try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import openpyxl
except ImportError:
    openpyxl = None


class DocumentDataExtractor:
    """Extract data from uploaded documents to fill forms"""
    
    # Field mappings for each document type - Spanish keywords to field names
    FIELD_MAPPINGS = {
        "estudios_previos": {
            "entidad": ["entidad", "entidad contratante", "institution"],
            "objeto": ["objeto", "objeto contractual", "objeto del contrato"],
            "presupuesto": ["presupuesto", "valor", "monto", "valor estimado"],
            "plazo": ["plazo", "duración", "tiempo de ejecución"],
            "modalidad": ["modalidad", "modalidad de selección"],
            "sector": ["sector", "área"],
            "descripcion": ["descripción", "justificación", "necesidad"],
        },
        "analisis_sector": {
            "sector": ["sector", "sector económico"],
            "entidad": ["entidad", "entidad contratante"],
            "objeto": ["objeto", "objeto a contratar"],
            "valor_estimado": ["valor", "presupuesto", "valor estimado"],
        },
        "dts": {
            "municipio": ["municipio", "ciudad"],
            "departamento": ["departamento"],
            "entidad": ["entidad"],
            "proyecto": ["proyecto", "nombre del proyecto"],
            "valor": ["valor", "presupuesto", "costo total"],
        },
        "certificaciones": {
            "nombre_proyecto": ["proyecto", "nombre del proyecto"],
            "municipio": ["municipio"],
            "departamento": ["departamento"],
            "valor": ["valor", "presupuesto"],
            "responsable": ["responsable", "formulador", "secretario"],
        },
        "mga_subsidios": {
            "municipio": ["municipio", "ciudad"],
            "departamento": ["departamento"],
            "entidad": ["entidad", "alcaldía"],
            "bpin": ["bpin", "código bpin"],
            "nombre_proyecto": ["proyecto", "nombre del proyecto", "título"],
            "valor_total": ["valor", "valor total", "presupuesto", "monto"],
            "duracion": ["duración", "plazo", "días"],
            "responsable": ["responsable", "formulador", "secretario"],
            "cargo": ["cargo", "puesto"],
            "plan_nacional": ["plan nacional", "pnd"],
            "plan_departamental": ["plan departamental"],
            "plan_municipal": ["plan municipal", "pdm"],
        }
    }
    
    def __init__(self, llm=None):
        """Initialize extractor with optional LLM for AI extraction"""
        self.llm = llm
    
    def extract_from_file(self, file, file_type: str, doc_type: str, user_context: str = "") -> Dict[str, Any]:
        """
        Extract data from uploaded file
        
        Args:
            file: Uploaded file object (Streamlit UploadedFile)
            file_type: Extension (.pdf, .docx, .xlsx)
            doc_type: Document type (estudios_previos, mga_subsidios, etc.)
            user_context: Optional user-provided context for extraction
            
        Returns:
            Dictionary with extracted field values
        """
        # Read file content
        if hasattr(file, 'read'):
            content = file.read()
            file.seek(0)  # Reset for potential re-read
        else:
            with open(file, 'rb') as f:
                content = f.read()
        
        # Extract text based on file type
        if file_type.lower() in ['.pdf', 'pdf']:
            text = self._extract_pdf_text(content)
        elif file_type.lower() in ['.docx', 'docx']:
            text = self._extract_docx_text(content)
        elif file_type.lower() in ['.xlsx', 'xlsx', '.xls', 'xls']:
            text = self._extract_xlsx_text(content)
        else:
            return {"error": f"Unsupported file type: {file_type}"}
        
        
        # Extract structured data
        if self.llm:
            result = self._extract_with_ai(text, doc_type, user_context)
        else:
            result = self._extract_with_patterns(text, doc_type)
            
        # Add raw text dump for context fallback
        result["context_dump"] = text
        
        # Store user context if provided
        if user_context:
            result["user_context"] = user_context
        
        return result
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF"""
        text_parts = []
        
        # Try PyMuPDF first (faster)
        if fitz:
            try:
                doc = fitz.open(stream=content, filetype="pdf")
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()
                return "\n".join(text_parts)
            except Exception as e:
                print(f"PyMuPDF error: {e}")
        
        # Fall back to pdfplumber
        if pdfplumber:
            try:
                with pdfplumber.open(BytesIO(content)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                return "\n".join(text_parts)
            except Exception as e:
                print(f"pdfplumber error: {e}")
        
        return ""
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX"""
        if not Document:
            return ""
        
        try:
            doc = Document(BytesIO(content))
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            return ""
    
    def _extract_xlsx_text(self, content: bytes) -> str:
        """Extract text from XLSX"""
        text_parts = []
        
        if pd:
            try:
                # Read all sheets
                xlsx = pd.ExcelFile(BytesIO(content))
                for sheet_name in xlsx.sheet_names:
                    df = pd.read_excel(xlsx, sheet_name=sheet_name)
                    # Convert to text format
                    for col in df.columns:
                        text_parts.append(f"{col}: {df[col].tolist()}")
                return "\n".join(text_parts)
            except Exception as e:
                print(f"Pandas Excel error: {e}")
        
        if openpyxl:
            try:
                wb = openpyxl.load_workbook(BytesIO(content))
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows(values_only=True):
                        row_text = [str(cell) for cell in row if cell]
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                return "\n".join(text_parts)
            except Exception as e:
                print(f"openpyxl error: {e}")
        
        return ""
    
    def _extract_with_patterns(self, text: str, doc_type: str) -> Dict[str, str]:
        """Extract data using regex patterns and keyword matching"""
        result = {}
        mappings = self.FIELD_MAPPINGS.get(doc_type, {})
        
        lines = text.split('\n')
        text_lower = text.lower()
        
        for field_name, keywords in mappings.items():
            for keyword in keywords:
                # Try to find keyword followed by colon and value
                pattern = rf'{keyword}\s*[:\-]?\s*(.+?)(?:\n|$)'
                match = re.search(pattern, text_lower, re.IGNORECASE)
                if match:
                    value = match.group(1).strip()
                    # Clean up the value
                    value = re.sub(r'^[:\-\s]+', '', value)
                    if value and len(value) > 1:
                        result[field_name] = value[:500]  # Limit length
                        break
        
        return result
    
    def _extract_with_ai(self, text: str, doc_type: str, user_context: str = "") -> Dict[str, str]:
        """Extract data using AI/LLM with improved prompt for better quality"""
        from langchain_core.prompts import ChatPromptTemplate
        from langchain_core.output_parsers import StrOutputParser
        
        # Comprehensive field list for all document types (unified)
        all_fields = [
            "municipio", "departamento", "entidad", "bpin", "nombre_proyecto",
            "valor_total", "duracion", "responsable", "cargo", "alcalde",
            "objeto", "necesidad", "alcance", "modalidad", "fuente_financiacion",
            "sector", "codigo_ciiu", "codigos_unspsc", "programa", "subprograma",
            "plan_nacional", "plan_departamental", "plan_municipal",
            "poblacion_beneficiada", "indicador_producto", "meta_producto",
            "es_actualizacion"
        ]
        
        fields_str = ", ".join(all_fields)
        
        # Use MORE text for better extraction (increased from 6000 to 15000)
        # Also sample from different parts of the document to get data from all pages
        text_length = len(text)
        if text_length > 15000:
            # Smart sampling: take from beginning, middle, and end
            beginning = text[:6000]
            middle_start = max(0, (text_length // 2) - 3000)
            middle = text[middle_start:middle_start + 6000]
            end = text[-3000:]
            text_to_analyze = f"{beginning}\n\n[...SECCIÓN INTERMEDIA...]\n\n{middle}\n\n[...SECCIÓN FINAL...]\n\n{end}"
        else:
            text_to_analyze = text
        
        # Build context section if provided
        context_section = ""
        if user_context:
            context_section = """

===== CONTEXTO DEL USUARIO =====
""" + user_context + """
===== FIN DEL CONTEXTO =====

NOTA: El usuario ha proporcionado contexto adicional. Si indica que es una ACTUALIZACIÓN, 
agrega "es_actualizacion": "Si" al JSON. Prioriza la información según las instrucciones del usuario.
"""
        
        # Build the human message without f-string to avoid brace escaping issues
        human_message = """Extrae los siguientes campos del documento gubernamental colombiano:
""" + fields_str + context_section + """

===== DOCUMENTO COMPLETO =====
""" + text_to_analyze + """
===== FIN DEL DOCUMENTO =====

⚠️ REGLAS CRÍTICAS - LEE CUIDADOSAMENTE:

1. EXTRAE VALORES REALES, NO ETIQUETAS:
   ❌ INCORRECTO: "bpin": "Código BPIN" o "01 - datos básicos"
   ✅ CORRECTO: "bpin": "202500000011507"
   
2. BPIN es un NÚMERO de 10+ dígitos (ej: 202500000011507)
   NUNCA extraer textos como "datos básicos", "Identificador:", "Código BPIN"
   
3. MUNICIPIO es un NOMBRE DE CIUDAD (ej: "San Pablo", "Cartagena")
   NUNCA extraer "municipio de" o nombres genéricos

4. NOMBRE_PROYECTO es el TÍTULO COMPLETO del proyecto
   NUNCA extraer "Tipología", "Nombre", u otras etiquetas

5. RESPONSABLE es un NOMBRE DE PERSONA (ej: "Roxana Cáceres Quiñonez")
   NUNCA extraer "Formulador Ciudadano:", "ciudadano:", u otras etiquetas

6. VALOR_TOTAL es un NÚMERO (ej: 309909217)
   NUNCA extraer "valor total", "presupuesto", u otras etiquetas

Responde con JSON válido. Ejemplo de respuesta CORRECTA:
{{
  "municipio": "San Pablo",
  "departamento": "Bolívar", 
  "bpin": "202500000011507",
  "nombre_proyecto": "Apoyo a pequeños productores del municipio",
  "valor_total": "309909217",
  "responsable": "Roxana Cáceres Quiñonez",
  "sector": "Agricultura y desarrollo rural"
}}

Si no encuentras un valor REAL (no etiquetas), omite ese campo."""

        prompt = ChatPromptTemplate.from_messages([
            ("system", """Eres un experto en extracción de datos de documentos MGA colombianos.

REGLA #1 MÁS IMPORTANTE: 
NUNCA extraigas etiquetas, títulos de sección, o nombres de campos como valores.
Busca los DATOS REALES que aparecen DESPUÉS de cada etiqueta.

EJEMPLOS DE LO QUE NO DEBES EXTRAER:
- "01 - Datos básicos del proyecto" → esto es un TÍTULO de sección
- "Tipología" → esto es una ETIQUETA
- "Código BPIN" → esto es una ETIQUETA
- "Formulador Ciudadano:" → esto es una ETIQUETA
- "valor extraído" → esto es placeholder

EJEMPLOS DE LO QUE SÍ DEBES EXTRAER:
- El número "202500000011507" que aparece después de "Código BPIN"
- El nombre "San Pablo" que aparece después de "Municipio"
- El número "309909217" que aparece después de "Valor Total"
- El nombre "Roxana Cáceres Quiñonez" que aparece después de "Formulador"

CAMPOS A EXTRAER:
- bpin: número largo (10+ dígitos) - SOLO NÚMEROS
- municipio: nombre de ciudad colombiana
- departamento: nombre de departamento colombiano
- nombre_proyecto: título descriptivo del proyecto
- valor_total: cantidad numérica (sin símbolo $)
- responsable: nombre completo de persona
- sector: nombre del sector económico

Responde SOLO con JSON válido, sin explicaciones."""),
            ("human", human_message)
        ])
        
        try:
            chain = prompt | self.llm | StrOutputParser()
            response = chain.invoke({})
            
            # Parse JSON response - try multiple patterns
            import json
            
            # Try to find JSON in response
            json_patterns = [
                r'```json\s*([\s\S]*?)\s*```',  # Code block (try first)
                r'```\s*([\s\S]*?)\s*```',  # Generic code block
                r'\{[\s\S]*\}',  # Any JSON object (fallback)
            ]
            
            for pattern in json_patterns:
                match = re.search(pattern, response, re.DOTALL)
                if match:
                    try:
                        json_str = match.group(1) if '```' in pattern else match.group(0)
                        result = json.loads(json_str)
                        
                        # Clean and filter the extracted data
                        cleaned_result = {}
                        for k, v in result.items():
                            if not v or v == "null" or not str(v).strip():
                                continue
                            
                            str_v = str(v).strip()
                            
                            # Skip placeholder values
                            placeholder_phrases = [
                                "valor extraído", "no encontrado", "no disponible",
                                "n/a", "por definir", "pendiente", "null"
                            ]
                            if str_v.lower() in placeholder_phrases:
                                continue
                            
                            # Clean numeric fields (remove currency symbols and formatting)
                            if k in ["valor_total", "duracion"]:
                                # Remove $, dots, commas, spaces
                                clean_num = re.sub(r'[\$\.,\s]', '', str_v)
                                # Try to extract just numbers
                                num_match = re.search(r'\d+', clean_num)
                                if num_match:
                                    str_v = num_match.group(0)
                            
                            cleaned_result[k] = str_v
                        
                        return cleaned_result
                    except json.JSONDecodeError:
                        continue
                        
        except Exception as e:
            print(f"AI extraction error: {e}")
        
        # Fall back to pattern extraction
        return self._extract_with_patterns(text, doc_type)


def extract_data_from_upload(uploaded_file, doc_type: str, llm=None, user_context: str = "") -> Dict[str, Any]:
    """
    Convenience function to extract data from an uploaded file
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        doc_type: Type of document to generate
        llm: Optional LLM for AI-powered extraction
        user_context: Optional user-provided context (e.g., "this is for updating existing document")
        
    Returns:
        Dictionary with extracted field values
    """
    if not uploaded_file:
        return {}
    
    # Get file extension
    filename = uploaded_file.name if hasattr(uploaded_file, 'name') else str(uploaded_file)
    ext = os.path.splitext(filename)[1].lower()
    
    extractor = DocumentDataExtractor(llm=llm)
    return extractor.extract_from_file(uploaded_file, ext, doc_type, user_context=user_context)


def summarize_development_plan(uploaded_file, llm_cheap=None) -> Dict[str, Any]:
    """
    Summarize a Development Plan PDF using a cheap model for token efficiency.
    
    This function extracts the full text from a PDF, then uses a cheap LLM
    (Gemini Flash with temp=0.1) to create a structured JSON summary.
    
    Args:
        uploaded_file: Streamlit UploadedFile object (PDF)
        llm_cheap: Optional LLM for summarization (default: gemini_flash_summarizer)
        
    Returns:
        Dictionary with structured summary:
        {
            "resumen_global": "Short overview...",
            "paginas_relevantes": [...],
            "datos_programa": {...},
            "raw_text_length": int,
            "summary_length": int
        }
    """
    import hashlib
    from langchain_core.prompts import ChatPromptTemplate
    from langchain_core.output_parsers import StrOutputParser
    
    if not uploaded_file:
        return {"error": "No file provided"}
    
    # Get file extension
    filename = uploaded_file.name if hasattr(uploaded_file, 'name') else str(uploaded_file)
    ext = os.path.splitext(filename)[1].lower()
    
    if ext not in ['.pdf', 'pdf']:
        return {"error": f"Expected PDF, got {ext}"}
    
    # Extract full text
    extractor = DocumentDataExtractor(llm=None)  # No AI for extraction
    
    if hasattr(uploaded_file, 'read'):
        content = uploaded_file.read()
        uploaded_file.seek(0)
    else:
        with open(uploaded_file, 'rb') as f:
            content = f.read()
    
    full_text = extractor._extract_pdf_text(content)
    
    if not full_text or len(full_text) < 100:
        return {"error": "Could not extract text from PDF", "raw_text_length": len(full_text) if full_text else 0}
    
    # Generate document hash for caching
    doc_hash = hashlib.md5(content).hexdigest()[:12]
    
    # Get cheap LLM if not provided
    if not llm_cheap:
        try:
            from config import get_llm
            llm_cheap = get_llm("gemini_flash_summarizer")
        except Exception as e:
            return {"error": f"Could not initialize summarizer: {e}", "raw_text_length": len(full_text)}
    
    # Summarization prompt
    summarization_prompt = """Eres un experto en extracción de datos de Planes de Desarrollo colombianos.

TAREA: Analiza este documento y extrae información estructurada para un proyecto MGA.

DOCUMENTO COMPLETO:
{full_text}

EXTRAE en formato JSON estricto:

{{
    "resumen_global": "Resumen de 2-3 oraciones del plan completo (máximo 100 palabras)",
    "paginas_relevantes": [
        {{
            "pagina": "Número o rango de página",
            "seccion": "Nombre de la sección",
            "contenido_clave": ["Dato 1", "Dato 2", "Dato 3"]
        }}
    ],
    "datos_programa": {{
        "codigos_programa": ["2402 - Nombre", "2403 - Nombre"],
        "presupuestos": ["5,000,000,000 - Programa X"],
        "metas": ["100km vías", "500 beneficiarios"],
        "indicadores": ["Km construidos", "Familias atendidas"]
    }},
    "poblacion": {{
        "total": "Número si se encuentra",
        "urbana": "Número",
        "rural": "Número"
    }}
}}

REGLAS CRÍTICAS:
1. Extrae SOLO datos factuales (números, nombres, códigos)
2. NO inventes ni interpretes datos
3. Prioriza: presupuestos, metas, indicadores, códigos de programa
4. Máximo 400 palabras en total
5. Si no encuentras un dato, omítelo (no pongas "No encontrado")

Responde SOLO con JSON válido, sin explicaciones."""

    prompt = ChatPromptTemplate.from_messages([
        ("system", "Eres un asistente de extracción de datos. Responde SOLO en JSON válido."),
        ("human", summarization_prompt)
    ])
    
    try:
        chain = prompt | llm_cheap | StrOutputParser()
        
        # Use full text (Gemini Flash handles large context)
        response = chain.invoke({"full_text": full_text[:50000]})  # Cap at 50k chars just in case
        
        # Parse JSON response
        json_patterns = [
            r'```json\s*([\s\S]*?)\s*```',
            r'```\s*([\s\S]*?)\s*```',
            r'\{[\s\S]*\}',
        ]
        
        for pattern in json_patterns:
            match = re.search(pattern, response, re.DOTALL)
            if match:
                try:
                    json_str = match.group(1) if '```' in pattern else match.group(0)
                    result = json.loads(json_str)
                    
                    # Add metadata
                    result["raw_text_length"] = len(full_text)
                    result["summary_length"] = len(response)
                    result["doc_hash"] = doc_hash
                    result["success"] = True
                    
                    return result
                except json.JSONDecodeError:
                    continue
        
        # JSON parsing failed, return raw response
        return {
            "error": "Could not parse JSON from summarizer",
            "raw_response": response[:500],
            "raw_text_length": len(full_text),
            "doc_hash": doc_hash
        }
        
    except Exception as e:
        return {
            "error": f"Summarization failed: {str(e)}",
            "raw_text_length": len(full_text),
            "doc_hash": doc_hash
        }

