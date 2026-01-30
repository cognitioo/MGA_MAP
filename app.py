"""
MGA AI Agent - Streamlit Web Interface
Herramienta de apoyo para formulación de proyectos públicos en la plataforma MGA del DNP

Features:
- Multi-model support (Groq, Gemini, OpenAI, Anthropic)
- Document generation for Estudios Previos and Análisis del Sector
- Word document download
- Spanish language interface
"""
 
import streamlit as st
import os
import sys
from datetime import datetime
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from config import get_llm, get_available_providers, LLM_PROVIDERS, APP_TITLE, APP_DESCRIPTION, get_secret, GROQ_API_KEY, OFFICIAL_SECTORS
from generators.estudios_previos_generator import EstudiosPreviosGenerator
from generators.analisis_sector_generator import AnalisisSectorGenerator
from generators.dts_generator import DTSGenerator
from generators.certificaciones_generator import CertificacionesGenerator
from generators.mga_subsidios_generator import MGASubsidiosGenerator
from generators.docx_builder import DocumentBuilder
from extractors.document_data_extractor import extract_data_from_upload
from generators.unified_generator import UnifiedGenerator
from editors.mga_editor import edit_mga_document, MGAEditor


# ═══════════════════════════════════════════════════════════════
# AUTHENTICATION SYSTEM
# ═══════════════════════════════════════════════════════════════

# Credentials (can be overridden via st.secrets)
ADMIN_PASSWORD = get_secret("ADMIN_PASSWORD", "MGA_Admin_2026!")
USER_PASSWORD = get_secret("USER_PASSWORD", "MGA_User_2026")

# Rate limiting settings
USER_DAILY_LIMIT = int(get_secret("USER_DAILY_LIMIT", "10"))  # Max generations per day for normal users
ADMIN_DAILY_LIMIT = 999999  # Unlimited for admins

def check_authentication():
    """Check if user is authenticated and return their role"""
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
        st.session_state.user_role = None
        st.session_state.generation_count_today = 0
        st.session_state.last_generation_date = None
    
    return st.session_state.authenticated, st.session_state.user_role

def login_page():
    """Render the login page"""
    st.set_page_config(page_title="MGA AI Agent - Login", page_icon="🔐", layout="centered")
    
    st.markdown("""
    <style>
        .login-container { max-width: 400px; margin: 0 auto; padding: 2rem; }
        .login-title { text-align: center; font-size: 2rem; margin-bottom: 1rem; }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown("# 🔐 MGA AI Agent")
    st.markdown("### Iniciar Sesión")
    st.markdown("---")
    
    password = st.text_input("Contraseña", type="password", key="login_password")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔑 Entrar como Admin", use_container_width=True):
            if password == ADMIN_PASSWORD:
                st.session_state.authenticated = True
                st.session_state.user_role = "admin"
                st.session_state.generation_count_today = 0
                st.success("✅ Bienvenido, Administrador!")
                st.rerun()
            else:
                st.error("❌ Contraseña incorrecta")
    
    with col2:
        if st.button("👤 Entrar como Usuario", use_container_width=True):
            if password == USER_PASSWORD:
                st.session_state.authenticated = True
                st.session_state.user_role = "user"
                st.session_state.generation_count_today = 0
                st.success("✅ Bienvenido, Usuario!")
                st.rerun()
            else:
                st.error("❌ Contraseña incorrecta")
    
    st.markdown("---")
    st.caption("Contacte al administrador si no tiene acceso.")

def check_rate_limit():
    """Check if user has exceeded their daily generation limit"""
    from datetime import date
    
    today = date.today().isoformat()
    
    # Reset count if it's a new day
    if st.session_state.get('last_generation_date') != today:
        st.session_state.generation_count_today = 0
        st.session_state.last_generation_date = today
    
    # Get limit based on role
    if st.session_state.user_role == "admin":
        limit = ADMIN_DAILY_LIMIT
    else:
        limit = USER_DAILY_LIMIT
    
    return st.session_state.generation_count_today < limit, limit

def increment_generation_count():
    """Increment the generation counter"""
    st.session_state.generation_count_today = st.session_state.get('generation_count_today', 0) + 1

# ═══════════════════════════════════════════════════════════════
# PRE-SUBMIT VALIDATION AI
# ═══════════════════════════════════════════════════════════════

def validate_form_data(data: dict, doc_type: str) -> list:
    """
    Validate form data before submission - checks presence AND quality.
    Uses full MGA Subsidios 24-page structure knowledge.
    Returns list of tuples: (field_name, severity, message)
    """
    issues = []
    
    # ═══════════════════════════════════════════════════════════════
    # MGA SUBSIDIOS COMPLETE STRUCTURE (24 PAGES)
    # Based on mga_subsidios_builder.py structure
    # ═══════════════════════════════════════════════════════════════
    
    MGA_SUBSIDIOS_STRUCTURE = {
        # PAGE 1: Datos Básicos - Only minimal required (AI extracts rest from files)
        "pagina_1_datos_basicos": {
            "fields": {
                "municipio": {"required": True, "min_len": 3, "desc": "Nombre del municipio"},
                "departamento": {"required": True, "min_len": 3, "desc": "Nombre del departamento"},
                "nombre_proyecto": {"required": False, "min_len": 10, "desc": "Nombre del proyecto (extraído del POAI)"},
                "bpin": {"required": False, "min_len": 8, "desc": "Código BPIN (extraído del POAI)"},
                "valor_total": {"required": False, "numeric": True, "desc": "Valor total (extraído del POAI)"},
                "duracion": {"required": False, "numeric": True, "desc": "Duración en días"},
                "entidad": {"required": False, "min_len": 5, "desc": "Entidad ejecutora"},
                "responsable": {"required": True, "min_len": 3, "desc": "Nombre del responsable"},
                "cargo": {"required": False, "min_len": 3, "desc": "Cargo del responsable"}
            }
        },
        # PAGE 2: Plan de Desarrollo - AI extracts from uploaded files
        "pagina_2_plan_desarrollo": {
            "fields": {
                "plan_nacional": {"required": False, "min_len": 10, "desc": "Plan Nacional de Desarrollo"},
                "plan_departamental": {"required": False, "min_len": 10, "desc": "Plan Departamental (extraído del PDF)"},
                "plan_municipal": {"required": False, "min_len": 10, "desc": "Plan Municipal (extraído del PDF)"}
            }
        },
        # PAGE 3: Problemática
        "pagina_3_problematica": {
            "fields": {
                "problema_central": {"required": True, "min_len": 30, "desc": "Descripción del problema central"},
                "descripcion_situacion": {"required": True, "min_len": 50, "desc": "Descripción de la situación existente"},
                "magnitud_problema": {"required": False, "min_len": 20, "desc": "Magnitud del problema - indicadores"}
            }
        },
        # PAGE 4: Causas y Efectos
        "pagina_4_causas_efectos": {
            "fields": {
                "causas_directas": {"required": True, "is_list": True, "desc": "Lista de causas directas"},
                "causas_indirectas": {"required": False, "is_list": True, "desc": "Lista de causas indirectas"},
                "efectos_directos": {"required": True, "is_list": True, "desc": "Lista de efectos directos"},
                "efectos_indirectos": {"required": False, "is_list": True, "desc": "Lista de efectos indirectos"}
            }
        },
        # PAGE 5: Participantes
        "pagina_5_participantes": {
            "fields": {
                "participantes": {"required": True, "is_list": True, "desc": "Lista de actores/participantes"}
            }
        },
        # PAGE 6: Población
        "pagina_6_poblacion": {
            "fields": {
                "poblacion_afectada": {"required": True, "min_len": 5, "desc": "Población afectada"},
                "poblacion_objetivo": {"required": True, "min_len": 5, "desc": "Población objetivo"}
            }
        },
        # PAGE 7: Objetivos
        "pagina_7_objetivos": {
            "fields": {
                "objetivo_general": {"required": True, "min_len": 30, "desc": "Objetivo general del proyecto"},
                "objetivos_especificos": {"required": True, "is_list": True, "desc": "Lista de objetivos específicos"}
            }
        },
        # PAGE 12: Análisis Técnico
        "pagina_12_analisis": {
            "fields": {
                "descripcion_alternativa": {"required": True, "min_len": 30, "desc": "Descripción de la alternativa"}
            }
        },
        # PAGE 13: Localización
        "pagina_13_localizacion": {
            "fields": {
                "region": {"required": False, "min_len": 3, "desc": "Región"},
                "latitud": {"required": False, "desc": "Coordenada latitud"},
                "longitud": {"required": False, "desc": "Coordenada longitud"}
            }
        },
        # PAGE 14: Cadena de Valor
        "pagina_14_cadena_valor": {
            "fields": {
                "indicador_producto": {"required": True, "min_len": 10, "desc": "Indicador de producto"},
                "meta_producto": {"required": True, "desc": "Meta del producto"}
            }
        }
    }
    
    # Flatten for form field validation
    critical_fields = {
        "municipio": {"min_len": 3, "desc": "Municipio"},
        "departamento": {"min_len": 3, "desc": "Departamento"},
        "nombre_proyecto": {"min_len": 20, "desc": "Nombre del proyecto (min. 20 caracteres)"},
        "valor_total": {"min_len": 1, "desc": "Valor total", "numeric": True},
        "bpin": {"min_len": 8, "desc": "Código BPIN (min. 8 dígitos)"}
    }
    
    warning_fields = {
        "responsable": {"min_len": 5, "desc": "Nombre del responsable"},
        "duracion": {"min_len": 1, "desc": "Duración en días", "numeric": True},
        "entidad": {"min_len": 5, "desc": "Entidad ejecutora"},
        "objeto": {"min_len": 30, "desc": "Objeto del contrato (min. 30 caracteres)"},
        "necesidad": {"min_len": 20, "desc": "Descripción de la necesidad"},
        "plan_nacional": {"min_len": 10, "desc": "Plan Nacional de Desarrollo"},
        "plan_departamental": {"min_len": 10, "desc": "Plan Departamental"},
        "plan_municipal": {"min_len": 10, "desc": "Plan Municipal"}
    }
    
    info_fields = {
        "sector": {"min_len": 3, "desc": "Sector del proyecto"},
        "programa": {"min_len": 5, "desc": "Programa"},
        "subprograma": {"min_len": 5, "desc": "Subprograma"},
        "poblacion_beneficiada": {"min_len": 5, "desc": "Población beneficiada"},
        "indicador_producto": {"min_len": 10, "desc": "Indicador de producto"},
        "meta_producto": {"min_len": 1, "desc": "Meta del producto"}
    }
    
    def check_field(field_name, rules, severity):
        value = data.get(field_name, "")
        value_str = str(value) if value else ""
        
        # Check if empty
        if not value or value == "" or value == "N/A" or value == "Por definir":
            issues.append((field_name, severity, f"⛔ '{rules['desc']}' está vacío o no definido"))
            return
        
        # Check minimum length
        min_len = rules.get("min_len", 1)
        if len(value_str) < min_len:
            issues.append((field_name, "warning", f"⚠️ '{rules['desc']}' muy corto ({len(value_str)}/{min_len} caracteres)"))
        
        # Check if numeric field
        if rules.get("numeric"):
            clean = value_str.replace(".", "").replace(",", "").replace("$", "").replace(" ", "")
            if not clean.isdigit():
                issues.append((field_name, "warning", f"⚠️ '{rules['desc']}' debe ser numérico"))
        
        # Detect placeholder/fake data
        fake_patterns = ["ejemplo", "xxx", "test", "prueba", "sample", "lorem", "placeholder"]
        for pattern in fake_patterns:
            if pattern.lower() in value_str.lower():
                issues.append((field_name, "critical", f"⛔ '{field_name}' contiene datos de ejemplo/prueba - PROHIBIDO"))
                break
    
    # Check all fields
    for field, rules in critical_fields.items():
        check_field(field, rules, "critical")
    
    for field, rules in warning_fields.items():
        check_field(field, rules, "warning")
    
    for field, rules in info_fields.items():
        check_field(field, rules, "info")
    
    return issues

def render_validation_panel(issues: list, doc_type: str) -> bool:
    """
    Render validation panel with issues and skip option.
    Returns True if user wants to proceed, False if they want to fix.
    """
    if not issues:
        return True
    
    critical_count = len([i for i in issues if i[1] == "critical"])
    warning_count = len([i for i in issues if i[1] == "warning"])
    
    with st.expander(f"🔍 Validación Pre-Generación ({len(issues)} observaciones)", expanded=True):
        if critical_count > 0:
            st.error(f"⛔ {critical_count} campo(s) crítico(s) faltante(s)")
        if warning_count > 0:
            st.warning(f"⚠️ {warning_count} campo(s) recomendado(s) faltante(s)")
        
        # Group by severity
        for severity, emoji, color in [("critical", "⛔", "red"), ("warning", "⚠️", "orange"), ("info", "ℹ️", "blue")]:
            severity_issues = [i for i in issues if i[1] == severity]
            if severity_issues:
                st.markdown(f"**{emoji} {'Críticos' if severity == 'critical' else 'Recomendados' if severity == 'warning' else 'Opcionales'}:**")
                for field, _, message in severity_issues:
                    st.caption(message)
        
        st.markdown("---")
        st.caption("💡 Puede completar los campos faltantes o continuar con la generación.")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("⏭️ Saltar Validación y Continuar", key=f"skip_validation_{doc_type}", use_container_width=True):
                return True
        with col2:
            st.button("✏️ Volver a Editar", key=f"edit_form_{doc_type}", use_container_width=True, disabled=True)
    
    return False

# Check authentication BEFORE setting page config
is_authenticated, user_role = check_authentication()

if not is_authenticated:
    login_page()
    st.stop()

# --- Page Configuration (only after authentication) ---
st.set_page_config(
    page_title=APP_TITLE,
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Custom CSS for Professional Look ---
st.markdown("""
<style>
    /* Main styling */
    .main-header {
        font-size: 1.8rem;
        font-weight: 600;
        color: #1a1a2e;
        margin-bottom: 0.3rem;
    }
    .sub-header {
        font-size: 0.95rem;
        color: #555;
        margin-bottom: 1.5rem;
    }
    .section-header {
        font-size: 1.1rem;
        font-weight: 600;
        color: #1a1a2e;
        border-left: 3px solid #4CAF50;
        padding-left: 0.8rem;
        margin-top: 1.5rem;
        margin-bottom: 0.8rem;
    }
    
    /* Model badge */
    .model-badge {
        display: inline-block;
        padding: 4px 12px;
        border-radius: 20px;
        font-size: 0.8rem;
        font-weight: 500;
        margin: 2px;
    }
    .model-groq { background: #f0f4ff; color: #3b5998; border: 1px solid #3b5998; }
    .model-gemini { background: #fff3e0; color: #e65100; border: 1px solid #e65100; }
    .model-openai { background: #e8f5e9; color: #2e7d32; border: 1px solid #2e7d32; }
    .model-anthropic { background: #fce4ec; color: #c2185b; border: 1px solid #c2185b; }
    
    /* Success/info boxes */
    .success-box {
        background-color: #e8f5e9;
        border-left: 4px solid #4CAF50;
        border-radius: 4px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .info-box {
        background-color: #e3f2fd;
        border-left: 4px solid #2196F3;
        border-radius: 4px;
        padding: 1rem;
        margin: 1rem 0;
    }
    
    /* Sidebar styling */
    .sidebar-section {
        padding: 0.5rem 0;
        border-bottom: 1px solid #eee;
        margin-bottom: 0.5rem;
    }
    .sidebar-title {
        font-weight: 600;
        color: #333;
        margin-bottom: 0.5rem;
</style>
""", unsafe_allow_html=True)

# --- Session State Initialization (Comprehensive) ---
def init_session_state():
    """Initialize all required session state variables safely"""
    defaults = {
        # Core generation state
        'generated_content': None,
        'generated_file': None,
        'extracted_data': {},
        
        # Edit mode state
        'previous_document': None,
        'edit_instructions_text': "",
        'selected_edit_pages': [],
        'additional_edit_files': [],
        'start_edit_process': False,
        'edit_mode_selected': None,
        
        # Generation history (new feature)
        'generation_history': [],
        'last_generation_time': None,
        
        # Error tracking
        'last_error': None,
        'error_count': 0,
        
        # User preferences
        'preferred_model': 'groq',
        'auto_download': False,
        
        # App version for debugging
        'app_version': '2.5.0 (Sanitized)'
    }
    
    for key, default_value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default_value

# Initialize on app load
init_session_state()

# --- API Status Check ---
def check_api_status():
    """Check which APIs are configured and working"""
    status = {}
    for provider_id, config in LLM_PROVIDERS.items():
        api_key = get_secret(config["env_key"])
        status[provider_id] = {
            "configured": bool(api_key),
            "key_preview": f"{api_key[:8]}..." if api_key and len(api_key) > 8 else "Not set"
        }
    return status


def render_data_upload_option(doc_type: str, key_prefix: str) -> dict:
    """
    Render the data file upload option for auto-filling forms
    
    Args:
        doc_type: Document type for extraction mapping
        key_prefix: Unique key prefix for Streamlit widgets
        
    Returns:
        Dictionary with extracted data (empty if none)
    """
    with st.expander("📄 Cargar archivo con datos del proyecto (opcional)", expanded=False):
        st.info("Suba un archivo PDF, DOCX o XLSX con información del proyecto para auto-llenar el formulario.")
        
        data_file = st.file_uploader(
            "Archivo de datos",
            type=["pdf", "docx", "xlsx"],
            key=f"{key_prefix}_data_file",
            help="El sistema extraerá automáticamente los datos del proyecto"
        )
        
        # Optional context field for updates
        user_context = st.text_area(
            "📝 Contexto Adicional (Opcional)",
            placeholder="Ej: 'Este documento es para ACTUALIZAR el proyecto existente BPIN 2024001234. Mantener los datos de presupuesto pero actualizar las fechas.'",
            height=80,
            key=f"{key_prefix}_user_context",
            help="Use este campo para indicar si es actualización, qué campos priorizar, o cualquier instrucción especial para la IA."
        )
        
        if data_file:
            if st.button("🔍 Extraer Datos con IA", key=f"{key_prefix}_extract_btn"):
                with st.spinner("Extrayendo datos con Groq Llama..."):
                    try:
                        # Use Groq Llama for extraction (30k TPM limit - much higher)
                        try:
                            llm = get_llm("groq_llama")
                        except Exception as llm_err:
                            st.warning(f"No se pudo iniciar IA de extracción: {llm_err}. Usando patrones.")
                            llm = None
                        
                        # Pass user context to extraction
                        extracted = extract_data_from_upload(data_file, doc_type, llm, user_context=user_context)
                        
                        if extracted and not extracted.get("error"):
                            st.session_state.extracted_data[doc_type] = extracted
                            
                            # Store raw JSON for display
                            st.session_state[f"{key_prefix}_raw_json"] = extracted
                            
                            # IMMEDIATELY apply to widget keys for auto-fill
                            FIELD_TO_WIDGET = {
                                "municipio": ["mga_municipio", "uni_municipio", "cert_municipio", "dts_municipio"],
                                "entidad": ["mga_entidad", "uni_entidad", "cert_entidad", "dts_entidad"],
                                "bpin": ["mga_bpin", "uni_bpin", "cert_bpin", "dts_bpin"],
                                "responsable": ["mga_responsable", "uni_responsable", "cert_responsable", "dts_responsable"],
                                "cargo": ["mga_cargo", "uni_cargo", "cert_cargo", "dts_cargo"],
                                "departamento": ["mga_depto", "uni_depto", "cert_depto", "dts_depto"],
                                "nombre_proyecto": ["mga_proyecto", "uni_proyecto", "cert_proyecto", "dts_proyecto"],
                                "valor_total": ["mga_valor", "uni_valor", "cert_valor", "dts_valor"],
                                "plan_nacional": ["mga_plan_nacional", "uni_pnd"],
                                "plan_departamental": ["mga_plan_depto", "uni_pd"],
                                "plan_municipal": ["mga_plan_mun", "uni_pm"],
                                "objeto": ["uni_objeto"],
                                "necesidad": ["uni_necesidad"],
                                "alcance": ["uni_alcance"],
                                "sector": ["uni_sector"],
                                "alcalde": ["uni_alcalde", "cert_alcalde"],
                                "programa": ["uni_programa", "dts_programa"],
                                "subprograma": ["uni_subprograma", "dts_subprograma"],
                            }
                            
                            filled_count = 0
                            for field_name, widget_keys in FIELD_TO_WIDGET.items():
                                if field_name in extracted and extracted[field_name]:
                                    value = str(extracted[field_name])
                                    for widget_key in widget_keys:
                                        st.session_state[widget_key] = value
                                    filled_count += 1
                            
                            st.success(f"✅ IA extrajo {len(extracted)} campos. {filled_count} campos aplicados al formulario.")
                            st.rerun()
                        else:
                            st.warning("No se pudieron extraer datos del documento. Complete el formulario manualmente.")
                    except Exception as e:
                        st.error(f"Error al extraer datos: {str(e)}")
        
        # Show editable JSON if extraction was done
        raw_json_key = f"{key_prefix}_raw_json"
        if raw_json_key in st.session_state and st.session_state[raw_json_key]:
            import json
            raw_data = st.session_state[raw_json_key]
            
            st.markdown("### 📋 Datos Extraídos por IA (JSON Editable)")
            st.info("Puede copiar valores de aquí o editar el JSON y volver a cargar.")
            
            # Define known form fields that get auto-filled
            known_form_fields = {
                "municipio", "departamento", "entidad", "bpin", "nombre_proyecto",
                "valor_total", "duracion", "responsable", "cargo", "alcalde",
                "objeto", "necesidad", "modalidad", "fuente_financiacion",
                "sector", "poblacion_beneficiada", "alcance", "programa",
                "plan_nacional", "plan_departamental", "plan_municipal",
                "indicador_producto", "meta_producto", "subprograma"
            }
            
            # Separate used and unused data
            used_data = {k: v for k, v in raw_data.items() if k in known_form_fields and k != "context_dump"}
            unused_data = {k: v for k, v in raw_data.items() if k not in known_form_fields and k not in ["context_dump", "user_context"]}
            
            # Show used data (goes to form)
            json_str = json.dumps(used_data, indent=2, ensure_ascii=False)
            edited_json = st.text_area(
                "JSON de Datos Extraídos (Auto-llenado)",
                value=json_str,
                height=200,
                key=f"{key_prefix}_json_edit"
            )
            
            # Show unused data (won't auto-fill but user can copy)
            if unused_data:
                st.markdown("### 📦 Datos Adicionales Extraídos (No Auto-llenados)")
                st.warning("⚠️ Estos datos no se utilizan en todos los campos del formulario. Revísalos y, si encuentras algún dato útil, complétalo manualmente.")
                unused_json = json.dumps(unused_data, indent=2, ensure_ascii=False)
                st.text_area(
                    "Datos Adicionales",
                    value=unused_json,
                    height=150,
                    key=f"{key_prefix}_unused_json",
                    disabled=False
                )
            
            # Show raw context dump for reference
            if raw_data.get("context_dump"):
                with st.expander("📄 Texto Completo Extraído del Documento"):
                    st.text_area(
                        "Contenido del documento",
                        value=raw_data["context_dump"][:5000] + ("..." if len(raw_data.get("context_dump", "")) > 5000 else ""),
                        height=200,
                        key=f"{key_prefix}_context_dump_display",
                        disabled=True
                    )
            
            # Button to apply edited JSON
            if st.button("🔄 Aplicar JSON Editado", key=f"{key_prefix}_apply_json"):
                try:
                    parsed = json.loads(edited_json)
                    # Merge with existing data, keeping context_dump
                    merged = {**st.session_state.extracted_data.get(doc_type, {}), **parsed}
                    st.session_state.extracted_data[doc_type] = merged
                    
                    # EXPLICIT field-to-widget key mapping (same as sidebar)
                    FIELD_TO_WIDGET = {
                        "municipio": ["mga_municipio", "uni_municipio", "cert_municipio", "dts_municipio"],
                        "entidad": ["mga_entidad", "uni_entidad", "cert_entidad", "dts_entidad"],
                        "bpin": ["mga_bpin", "uni_bpin", "cert_bpin", "dts_bpin"],
                        "responsable": ["mga_responsable", "uni_responsable", "cert_responsable", "dts_responsable"],
                        "cargo": ["mga_cargo", "uni_cargo", "cert_cargo", "dts_cargo"],
                        "departamento": ["mga_depto", "uni_depto", "cert_depto", "dts_depto"],
                        "nombre_proyecto": ["mga_proyecto", "uni_proyecto", "cert_proyecto", "dts_proyecto"],
                        "valor_total": ["mga_valor", "uni_valor", "cert_valor", "dts_valor"],
                        "plan_nacional": ["mga_plan_nacional", "uni_pnd"],
                        "plan_departamental": ["mga_plan_depto", "uni_pd"],
                        "plan_municipal": ["mga_plan_mun", "uni_pm"],
                        "objeto": ["uni_objeto"],
                        "necesidad": ["uni_necesidad"],
                        "alcance": ["uni_alcance"],
                        "sector": ["uni_sector"],
                        "alcalde": ["uni_alcalde", "cert_alcalde"],
                        "programa": ["uni_programa", "dts_programa"],
                        "subprograma": ["uni_subprograma", "dts_subprograma"],
                    }
                    
                    # Apply to widget keys
                    for field_name, widget_keys in FIELD_TO_WIDGET.items():
                        if field_name in parsed and parsed[field_name]:
                            value = str(parsed[field_name])
                            for widget_key in widget_keys:
                                st.session_state[widget_key] = value
                    
                    st.success(f"✅ JSON aplicado! {len(parsed)} campos actualizados.")
                    st.rerun()
                except json.JSONDecodeError as je:
                    st.error(f"JSON inválido: {je}")
    
    return st.session_state.extracted_data.get(doc_type, {})


def get_model_options():
    """Get available model options based on configured API keys"""
    options = {}
    
    # Check for Groq (priority)
    if get_secret("GROQ_API_KEY"):
        options["groq"] = "Groq - Llama (Rápido)"
    
    # Check for Gemini
    if get_secret("GOOGLE_API_KEY"):
        options["gemini"] = "Google Gemini"
    
    # Check for OpenAI
    if get_secret("OPENAI_API_KEY"):
        options["openai"] = "OpenAI GPT-4"
    
    # Check for Anthropic
    if get_secret("ANTHROPIC_API_KEY"):
        options["anthropic"] = "Anthropic Claude"
    
    if not options:
        options["groq"] = "Groq (Configure API Key)"
    
    return options


def render_sidebar():
    """Render the sidebar with model selection, mode, and customization controls"""
    with st.sidebar:
        # Logo/Title
        st.markdown("### 🛠️ MGA Agent v2.5")
        st.caption("Generador de Documentos")
        
        # User info and logout
        user_role = st.session_state.get('user_role', 'user')
        role_emoji = "👑" if user_role == "admin" else "👤"
        st.markdown(f"**{role_emoji} Rol: {user_role.title()}**")
        
        if GROQ_API_KEY:
            safe_key = GROQ_API_KEY[:10] + "..." if len(GROQ_API_KEY) > 10 else "Short/Invalid"
            st.caption(f"🔑 Key Loaded: `{safe_key}`")
            
            # Connection Test Button
            if st.button("⚡ Probar Conexión", key="test_conn_btn", use_container_width=True):
                with st.spinner("Probando APIs..."):
                    # Try Groq first
                    try:
                        from config import get_llm
                        llm = get_llm("groq")
                        res = llm.invoke("Hello")
                        st.success(f"✅ Groq OK: {res.content[:20]}...")
                    except Exception as e:
                        st.error(f"❌ Groq Error: {str(e)[:100]}")
                        
                        # Try Gemini as fallback
                        try:
                            llm_gemini = get_llm("gemini")
                            res_g = llm_gemini.invoke("Hello")
                            st.success(f"✅ Gemini OK: {res_g.content[:20]}...")
                            st.info("💡 Considere usar Gemini como modelo principal")
                        except Exception as ge:
                            st.error(f"❌ Gemini Error: {str(ge)[:100]}")
        else:
            st.error("🔑 No API Key Found!")
        
        # Show rate limit for users
        can_generate, limit = check_rate_limit()
        used = st.session_state.get('generation_count_today', 0)
        if user_role != "admin":
            st.caption(f"📊 Usos hoy: {used}/{limit}")
            if not can_generate:
                st.error("⚠️ Límite diario alcanzado")
        
        # Logout button
        if st.button("🚪 Cerrar Sesión", key="logout_btn", use_container_width=True):
            st.session_state.authenticated = False
            st.session_state.user_role = None
            st.rerun()
        
        st.markdown("---")
        
        # ══════════════════════════════════════
        # MODE SELECTION
        # ══════════════════════════════════════
        st.markdown("**📋 Modo de Generación**")
        
        generation_mode = st.radio(
            "Seleccione modo:",
            options=["crear_nuevo", "actualizar_existente", "generar_desde_mga"],
            format_func=lambda x: {
                "crear_nuevo": "🆕 Crear desde cero",
                "actualizar_existente": "✏️ Actualizar existente",
                "generar_desde_mga": "📄 Generar documentos desde MGA"
            }.get(x, x),
            help="Crear nuevo: Genera MGA completo. Actualizar: Modifica existente. Desde MGA: Genera Análisis Sector/Estudios Previos.",
            label_visibility="collapsed"
        )
        st.session_state.generation_mode = generation_mode
        
        # ══════════════════════════════════════
        # UPDATE MODE CONTROLS
        # ══════════════════════════════════════
        if generation_mode == "actualizar_existente":
            st.success("📝 Modo Edición Activado")
            
            # 1. PROMPT (First, as requested)
            st.markdown("**1. Instrucciones de Edición**")
            edit_prompt = st.text_area(
                "Instrucciones",
                placeholder="Ej: Actualizar los montos a la vigencia 2026 y cambiar el nombre del responsable...",
                height=120,
                key="edit_prompt_input",
                help="Describa qué cambios desea que la IA realice en el documento.",
                label_visibility="collapsed"
            )
            # Update session state immediately
            if edit_prompt != st.session_state.get("edit_instructions_text", ""):
                 st.session_state.edit_instructions_text = edit_prompt

            # 2. FILE UPLOAD (Second)
            st.markdown("**2. Documento a Editar**")
            prev_doc = st.file_uploader(
                "Subir MGA anterior",
                type=["pdf", "docx"],
                help="Suba el documento MGA (PDF o Word) que desea actualizar",
                key="prev_doc_upload",
                label_visibility="collapsed"
            )
            
            if prev_doc:
                st.session_state.previous_document = prev_doc
                file_ext = prev_doc.name.split('.')[-1].lower()
                
                # Page Count Logic (PDF & DOCX)
                page_count = 0
                try:
                    if file_ext == 'pdf':
                        import fitz
                        prev_doc.seek(0)
                        doc = fitz.open(stream=prev_doc.read(), filetype="pdf")
                        page_count = doc.page_count
                        prev_doc.seek(0)
                    elif file_ext == 'docx':
                        import docx
                        prev_doc.seek(0)
                        doc = docx.Document(prev_doc)
                        # Estimate pages for DOCX (not exact, but helpful context)
                        # Or just count paragraphs/sections if pages aren't available
                        # Using core properties if available, otherwise just confirm content
                        try:
                            page_count = doc.core_properties.revision # sometimes used as proxy or just show "Detectado"
                        except:
                            page_count = "?"
                            
                        # Better approach involves rendering, but for now we just acknowledge it
                        if page_count == "?" or page_count == 0:
                             st.info(f"📄 Documento Word cargado")
                        else:
                             st.info(f"📄 Documento cargado (Rev: {page_count})")
                             
                        # Reset seek
                        prev_doc.seek(0)
                        
                    if file_ext == 'pdf' and page_count > 0:
                        st.info(f"📄 Documento PDF: {page_count} páginas")
                        
                except Exception as e:
                    st.warning(f"📄 Archivo cargado (no se pudo contar páginas: {str(e)})")
            
            # Page Selection (with "Select All" option)
            if prev_doc and page_count > 0:
                 st.caption("📑 Seleccione páginas a editar")
                 
                 # Select All checkbox
                 select_all = st.checkbox(
                     "Seleccionar todas las páginas",
                     value=True,
                     key="select_all_pages",
                     help="Marque para editar todo el documento"
                 )
                 
                 if select_all:
                     st.session_state.selected_edit_pages = []  # Empty means all pages
                     st.info(f"📄 Se editarán las {page_count} páginas del documento")
                 else:
                     selected_pages = st.multiselect(
                         "Páginas Específicas",
                         options=list(range(1, page_count + 1)),
                         default=[],
                         key="selected_pages_edit",
                         help="Seleccione las páginas que desea editar"
                     )
                     st.session_state.selected_edit_pages = selected_pages
                     if selected_pages:
                         st.info(f"📄 Se editarán las páginas: {', '.join(map(str, selected_pages))}")
            
            st.markdown("---")
            st.caption("ℹ️ La IA analizará el documento y aplicará sus instrucciones.")

        # ══════════════════════════════════════
        # GENERATE FROM MGA MODE
        # ══════════════════════════════════════
        elif generation_mode == "generar_desde_mga":
            st.success("📄 Generar Documentos desde MGA")
            
            st.markdown("**1. Suba el MGA existente**")
            mga_source = st.file_uploader(
                "MGA fuente",
                type=["pdf", "docx"],
                help="Suba el documento MGA del cual extraer datos",
                key="mga_source_upload",
                label_visibility="collapsed"
            )
            
            if mga_source:
                st.session_state.mga_source_file = mga_source
                st.info(f"✅ Archivo cargado: {mga_source.name}")
            
            st.markdown("**2. Seleccione documentos a generar**")
            docs_to_generate = st.multiselect(
                "Documentos",
                options=["analisis_sector", "estudios_previos"],
                format_func=lambda x: "📊 Análisis del Sector" if x == "analisis_sector" else "📋 Estudios Previos",
                default=[],
                key="docs_from_mga",
                label_visibility="collapsed"
            )
            st.session_state.docs_to_generate_from_mga = docs_to_generate
            
            if docs_to_generate:
                st.info(f"Se generarán {len(docs_to_generate)} documento(s)")
            
            st.markdown("---")
            st.caption("ℹ️ La IA extraerá datos del MGA y generará los documentos seleccionados.")

        
        # ══════════════════════════════════════
        # MODEL SELECTION
        # ══════════════════════════════════════
        st.markdown("**🤖 Modelo de IA**")
        
        model_options = get_model_options()
        selected_model = st.selectbox(
            "Modelo",
            options=list(model_options.keys()),
            format_func=lambda x: model_options[x],
            help="Seleccione el modelo de IA",
            label_visibility="collapsed"
        )
        
        # ══════════════════════════════════════
        # DOCUMENT TYPE (Only for New Mode)
        # ══════════════════════════════════════
        if generation_mode == "crear_nuevo":
            st.markdown("---")
            st.markdown("**📑 Plantilla**")
            template = st.selectbox(
                "Plantilla",
                options=["estandar", "simplificado", "completo"],
                format_func=lambda x: x.title(),
                label_visibility="collapsed"
            )
            st.session_state.selected_template = template
        
        st.markdown("---")
        st.caption("Versión 2.2 | © 2026")
        
        return selected_model


def render_estudios_previos_form():
    """Render input form for Estudios Previos"""
    
    # Data extraction option
    extracted = render_data_upload_option("estudios_previos", "ep")
    
    st.markdown('<p class="section-header">Datos del Proyecto</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        municipio = st.text_input("Municipio *", value=extracted.get("municipio", ""), placeholder="Ej: San Pablo")
        departamento = st.text_input("Departamento *", value=extracted.get("departamento", ""), placeholder="Ej: Bolívar")
        entidad = st.text_input("Entidad Contratante", value=extracted.get("entidad", ""), placeholder="Ej: Alcaldía Municipal")
        bpin = st.text_input("Código BPIN", placeholder="Ej: 2024000001367")
    
    with col2:
        tipo_proyecto = st.text_input("Tipo de Proyecto *", placeholder="Ej: Convenio Interadministrativo")
        valor_total = st.text_input("Valor Total (COP) *", value=extracted.get("presupuesto", ""), placeholder="Ej: 70000000")
        plazo = st.number_input("Plazo de Ejecución (días)", min_value=1, value=90)
        fuente = st.text_input("Fuente de Financiación", placeholder="Ej: Recursos Propios")
    
    st.markdown('<p class="section-header">Descripción del Proyecto</p>', unsafe_allow_html=True)
    
    necesidad = st.text_area(
        "Descripción de la Necesidad *",
        placeholder="Describa el problema o necesidad que origina el proyecto...",
        height=100
    )
    
    objeto = st.text_area(
        "Objeto del Convenio/Contrato *",
        placeholder="Describa el objeto del contrato...",
        height=80
    )
    
    alcance = st.text_area(
        "Alcance y Actividades",
        placeholder="Liste las actividades principales a desarrollar...",
        height=100
    )
    
    st.markdown('<p class="section-header">Presupuesto</p>', unsafe_allow_html=True)
    
    rubros = st.text_area(
        "Rubros Presupuestales",
        placeholder="Ej:\n- Honorarios profesionales: $50,000,000\n- Gastos operativos: $10,000,000\n- Otros: $10,000,000",
        height=100
    )
    
    st.markdown('<p class="section-header">Obligaciones (Secciones 4 y 5)</p>', unsafe_allow_html=True)
    
    obligaciones_municipio = st.text_area(
        "Obligaciones del Municipio *",
        placeholder="Liste las obligaciones del municipio. Use • para viñetas.\nEj:\n• Proveer la información cartográfica requerida\n• Facilitar el acceso del equipo técnico\n• Garantizar la disponibilidad presupuestal",
        height=120
    )
    
    obligaciones_contratista = st.text_area(
        "Obligaciones del Contratista/Empresa *",
        placeholder="Liste las obligaciones del contratista. Use • para viñetas.\nEj:\n• Ejecutar las actividades descritas en el alcance\n• Presentar informes técnicos parciales y finales\n• Cumplir con las normas de seguridad",
        height=120
    )
    
    st.markdown('<p class="section-header">CDP - Certificado de Disponibilidad Presupuestal (Sección 7)</p>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        cdp_codigo = st.text_input("Código CDP", placeholder="Ej: 20260108-3")
        cdp_rubro = st.text_input("Rubro", placeholder="Ej: 2.3.2.02.02.009.45.25 – Servicio de Asistencia Técnica")
    with col2:
        cdp_fecha = st.text_input("Fecha CDP", placeholder="Ej: 08/01/2026")
        cdp_fuente = st.text_input("Fuente", placeholder="Ej: SGP – Propósito General – Libre Inversión")
    with col3:
        cdp_valor = st.text_input("Valor CDP", placeholder="Ej: $17.400.000,00")
    
    st.markdown('<p class="section-header">Membrete/Letterhead</p>', unsafe_allow_html=True)
    
    letterhead_file = st.file_uploader(
        "Subir plantilla con membrete (.docx)",
        type=["docx"],
        help="Suba un archivo .docx con el encabezado y pie de página de la alcaldía. El documento generado usará este membrete."
    )
    
    st.markdown('<p class="section-header">Responsable (Sección 12)</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        responsable = st.text_input("Nombre del Responsable *", placeholder="Ej: Carlos Augusto Gil Delgado")
    with col2:
        cargo = st.text_input("Cargo", value="Secretario de Planeación Municipal")
    
    return {
        "municipio": municipio,
        "departamento": departamento,
        "entidad": entidad,
        "bpin": bpin,
        "tipo_proyecto": tipo_proyecto,
        "valor_total": valor_total,
        "plazo": str(plazo),
        "fuente_financiacion": fuente,
        "necesidad": necesidad,
        "objeto": objeto,
        "alcance": alcance,
        "rubros": rubros,
        "obligaciones_municipio": obligaciones_municipio,
        "obligaciones_contratista": obligaciones_contratista,
        "responsable": responsable,
        "cargo": cargo,
        "lugar": f"Municipio de {municipio}, {departamento}",
        "cdp_data": {
            "cdp": cdp_codigo,
            "fecha": cdp_fecha,
            "rubro": cdp_rubro,
            "fuente": cdp_fuente,
            "valor": cdp_valor
        },
        "letterhead_file": letterhead_file
    }


def render_analisis_sector_form():
    """Render input form for Análisis del Sector"""
    
    # Data extraction option
    extracted = render_data_upload_option("analisis_sector", "as")
    
    st.markdown('<p class="section-header">Datos del Contrato</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        numero_contrato = st.text_input("Número de Contrato", placeholder="Ej: 001-2025")
        modalidad = st.selectbox("Modalidad", [
            "Convenio Interadministrativo",
            "Contratación Directa",
            "Licitación Pública",
            "Selección Abreviada",
            "Concurso de Méritos"
        ])
        municipio = st.text_input("Municipio *", placeholder="Ej: San Pablo")
        departamento = st.text_input("Departamento *", placeholder="Ej: Bolívar")
    
    with col2:
        entidad = st.text_input("Entidad *", placeholder="Ej: Alcaldía Municipal")
        nombre_proyecto = st.text_input("Nombre del Proyecto *", placeholder="Ej: Actualización PSMV")
        bpin = st.text_input("Código BPIN", placeholder="Ej: 2024000001367")
        valor_total = st.text_input("Valor Total (COP) *", placeholder="Ej: 70000000")
    
    st.markdown('<p class="section-header">Información del Sector</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        sector_options = OFFICIAL_SECTORS
        sector = st.selectbox("Sector *", sector_options, key="as_sector")
        codigo_ciiu = st.text_input("Código CIIU", placeholder="Ej: 7110 - Actividades de arquitectura e ingeniería")
    
    with col2:
        codigos_unspsc = st.text_input("Códigos UNSPSC", placeholder="Ej: 77101600, 77101700")
        duracion = st.number_input("Duración (días)", min_value=1, value=90)
    
    objeto = st.text_area(
        "Objeto del Contrato *",
        placeholder="Describa el objeto del contrato...",
        height=80
    )
    
    plan_desarrollo = st.text_input(
        "Plan de Desarrollo Relacionado",
        placeholder="Ej: Plan de Desarrollo Municipal 2024-2027 'San Pablo Mejor'"
    )
    
    fuente = st.text_input("Fuente de Financiación", placeholder="Ej: Recursos Propios del Municipio")
    
    st.markdown('<p class="section-header">Responsable</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        responsable = st.text_input("Nombre del Responsable *", placeholder="Ej: Carlos Augusto Gil Delgado")
    with col2:
        cargo = st.text_input("Cargo", value="Secretario de Planeación Municipal")
    
    st.markdown('<p class="section-header">Membrete/Letterhead</p>', unsafe_allow_html=True)
    
    letterhead_file = st.file_uploader(
        "Subir plantilla con membrete (.docx)",
        type=["docx"],
        help="Suba un archivo .docx con el encabezado y pie de página de la alcaldía.",
        key="analisis_letterhead"
    )
    
    return {
        "numero_contrato": numero_contrato,
        "modalidad": modalidad,
        "municipio": municipio,
        "departamento": departamento,
        "entidad": entidad,
        "nombre_proyecto": nombre_proyecto,
        "bpin": bpin,
        "valor_total": valor_total,
        "sector": sector,
        "codigo_ciiu": codigo_ciiu,
        "codigos_unspsc": codigos_unspsc,
        "duracion": str(duracion),
        "objeto": objeto,
        "plan_desarrollo": plan_desarrollo,
        "fuente_financiacion": fuente,
        "responsable": responsable,
        "cargo": cargo,
        "ano": str(datetime.now().year),
        "letterhead_file": letterhead_file
    }


def render_dts_form():
    """Render input form for DTS (Documento Técnico de Soporte)"""
    
    # Data extraction option
    extracted = render_data_upload_option("dts", "dts")
    
    st.markdown('<p class="section-header">Datos del Proyecto</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        municipio = st.text_input("Municipio *", placeholder="Ej: San Pablo", key="dts_municipio")
        departamento = st.text_input("Departamento *", placeholder="Ej: Bolívar", key="dts_depto")
        entidad = st.text_input("Entidad *", placeholder="Ej: Alcaldía Municipal", key="dts_entidad")
        bpin = st.text_input("Código BPIN", placeholder="Ej: 2024000001367", key="dts_bpin")
    
    with col2:
        nombre_proyecto = st.text_input("Nombre del Proyecto *", placeholder="Ej: Subsidio de Servicios Públicos", key="dts_proyecto")
        valor_total = st.text_input("Valor Total (COP) *", placeholder="Ej: 1200532612", key="dts_valor")
        duracion = st.number_input("Duración (días)", min_value=1, value=365, key="dts_duracion")
    
    st.markdown('<p class="section-header">Planes de Desarrollo</p>', unsafe_allow_html=True)
    
    programa = st.text_input(
        "Programa",
        placeholder="Ej: 4002 - Usuarios beneficiados con subsidios al consumo del servicio",
        key="dts_programa"
    )
    
    subprograma = st.text_input(
        "Subprograma",
        placeholder="Ej: Agua de vida con Identidad Bolivarense",
        key="dts_subprograma"
    )
    
    st.markdown('<p class="section-header">Objeto y Descripción</p>', unsafe_allow_html=True)
    
    objeto = st.text_area(
        "Objeto del Proyecto *",
        placeholder="Describa el objeto del proyecto de subsidio...",
        height=80,
        key="dts_objeto"
    )
    
    descripcion_problema = st.text_area(
        "Descripción del Problema",
        placeholder="Describa la situación actual y el problema a resolver...",
        height=100,
        key="dts_problema"
    )
    
    st.markdown('<p class="section-header">Membrete/Letterhead</p>', unsafe_allow_html=True)
    
    letterhead_file = st.file_uploader(
        "Subir plantilla con membrete (.docx)",
        type=["docx"],
        help="Suba un archivo .docx con el encabezado y pie de página de la alcaldía.",
        key="dts_letterhead"
    )
    
    st.markdown('<p class="section-header">Responsable</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        responsable = st.text_input("Nombre del Responsable *", placeholder="Ej: Carlos Augusto Gil Delgado", key="dts_responsable")
    with col2:
        cargo = st.text_input("Cargo", value="Secretario de Planeación Municipal", key="dts_cargo")
    
    return {
        "municipio": municipio,
        "departamento": departamento,
        "entidad": entidad,
        "bpin": bpin,
        "nombre_proyecto": nombre_proyecto,
        "valor_total": valor_total,
        "duracion": str(duracion),
        "programa": programa,
        "subprograma": subprograma,
        "objeto": objeto,
        "descripcion_problema": descripcion_problema,
        "responsable": responsable,
        "cargo": cargo,
        "letterhead_file": letterhead_file
    }


def render_certificaciones_form():
    """Render input form for Presentación y Certificaciones"""
    
    # Data extraction option
    extracted = render_data_upload_option("certificaciones", "cert")
    
    st.markdown('<p class="section-header">Datos del Proyecto</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        municipio = st.text_input("Municipio *", placeholder="Ej: San Pablo", key="cert_municipio")
        departamento = st.text_input("Departamento *", placeholder="Ej: Bolívar", key="cert_depto")
        entidad = st.text_input("Entidad *", placeholder="Ej: Alcaldía Municipal", key="cert_entidad")
        bpin = st.text_input("Código BPIN", placeholder="Ej: 2024000001367", key="cert_bpin")
    
    with col2:
        nombre_proyecto = st.text_input("Nombre del Proyecto *", placeholder="Ej: Subsidio de Servicios Públicos", key="cert_proyecto")
        valor_total = st.text_input("Valor Total (COP) *", placeholder="Ej: 1662485076", key="cert_valor")
        alcalde = st.text_input("Nombre del Alcalde *", placeholder="Ej: Jair Acevedo Cavadía", key="cert_alcalde")
    
    st.markdown('<p class="section-header">Plan de Desarrollo</p>', unsafe_allow_html=True)
    
    plan_desarrollo = st.text_input(
        "Plan de Desarrollo Municipal",
        placeholder="Ej: San Pablo Mejor 2024-2027",
        key="cert_plan"
    )
    
    st.markdown('<p class="section-header">Membrete/Letterhead</p>', unsafe_allow_html=True)
    
    letterhead_file = st.file_uploader(
        "Subir plantilla con membrete (.docx)",
        type=["docx"],
        help="Suba un archivo .docx con el encabezado y pie de página de la alcaldía.",
        key="cert_letterhead"
    )
    
    st.markdown('<p class="section-header">Responsable (Secretario de Planeación)</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        responsable = st.text_input("Nombre del Responsable *", placeholder="Ej: Carlos Augusto Gil Delgado", key="cert_responsable")
    with col2:
        cargo = st.text_input("Cargo", value="Secretario de Planeación Municipal", key="cert_cargo")
    
    return {
        "municipio": municipio,
        "departamento": departamento,
        "entidad": entidad,
        "bpin": bpin,
        "nombre_proyecto": nombre_proyecto,
        "valor_total": valor_total,
        "alcalde": alcalde,
        "plan_desarrollo": plan_desarrollo,
        "responsable": responsable,
        "cargo": cargo,
        "letterhead_file": letterhead_file
    }


def render_mga_subsidios_form():
    """
    Simplified MGA form - Upload files + minimal info → AI generates MGA
    """
    
    st.markdown("""
    <div style="background: linear-gradient(135deg, #0D47A1, #1976D2); padding: 20px; border-radius: 10px; margin-bottom: 20px;">
        <h3 style="color: white; margin: 0;">📋 Generar MGA desde Documentos</h3>
        <p style="color: rgba(255,255,255,0.8); margin: 5px 0 0 0;">Suba POAI + Plan de Desarrollo → La IA genera el MGA completo</p>
    </div>
    """, unsafe_allow_html=True)
    
    # ============================================================
    # FILE UPLOADS - Main Data Sources
    # ============================================================
    st.markdown("### 📁 Archivos de Entrada (Obligatorios)")
    
    col1, col2 = st.columns(2)
    
    with col1:
        poai_file = st.file_uploader(
            "📊 **POAI** (Excel .xlsx) *",
            type=["xlsx", "xls"],
            help="Plan Operativo Anual de Inversiones - archivo principal del proyecto",
            key="mga_poai_file"
        )
        
    with col2:
        dev_plan_file = st.file_uploader(
            "📋 **Plan de Desarrollo** (PDF) *",
            type=["pdf"],
            help="Plan de Desarrollo Municipal o Departamental en PDF",
            key="mga_devplan_file"
        )
    
    # Optional additional file
    basic_info_file = st.file_uploader(
        "📝 Información Adicional (Opcional - PDF/Word)",
        type=["pdf", "docx"],
        help="Otros documentos con información relevante del proyecto",
        key="mga_basicinfo_file"
    )
    
    # Extract data from uploaded files
    context_dump = ""
    extracted_summary = []
    
    # Process POAI (XLSX) - IMPROVED EXTRACTION
    extracted_poai_codes = []  # Store extracted program codes
    extracted_poai_data = {}   # Store additional extracted fields
    poai_critical_section = "" # HIGH PRIORITY section for codes - placed FIRST
    
    if poai_file:
        try:
            import pandas as pd
            xlsx = pd.ExcelFile(poai_file)
            poai_text = ""
            
            # Column patterns for different data types
            CODE_PATTERNS = ['código programa', 'codigo programa', 'código presupuestal', 'codigo presupuestal']
            BPIN_PATTERNS = ['bpin', 'codigo bpin', 'código bpin']
            PROJECT_NAME_PATTERNS = ['producto mga', 'nombre proyecto', 'proyecto', 'producto']
            SECTOR_PATTERNS = ['sector', 'cod sector']
            VALUE_PATTERNS = ['total', 'recursos', 'valor', 'presupuesto']
            
            print(f"[POAI DEBUG] Processing {len(xlsx.sheet_names)} sheets: {xlsx.sheet_names}")
            
            for sheet_name in xlsx.sheet_names:
                # Detect Header Row dynamically
                header_idx = 0
                try:
                    df_preview = pd.read_excel(xlsx, sheet_name=sheet_name, header=None, nrows=20)
                    for idx, row in df_preview.iterrows():
                        row_str = " ".join([str(x).lower() for x in row.values])
                        if "código" in row_str and "programa" in row_str:
                            header_idx = idx
                            print(f"[POAI DEBUG] Found likely header at row {header_idx} in '{sheet_name}'")
                            break
                except Exception as e:
                    print(f"[POAI DEBUG] Error scanning headers: {e}")
                
                # Load dataframe with detected header row
                df = pd.read_excel(xlsx, sheet_name=sheet_name, header=header_idx)
                
                # --- VISIBLE DEBUGGING FOR USER ---
                st.write(f"🔍 **Diagnóstico POAI (Hoja: {sheet_name})**")
                st.write(f"- Fila de encabezado detectada: **{header_idx + 1}**")
                st.code(f"Columnas: {list(df.columns)[:5]}...")
                
                poai_text += f"\n=== Hoja: {sheet_name} ===\n"
                poai_text += df.to_string(index=False)[:4000]
                
                # Log all columns for debugging
                print(f"[POAI DEBUG] Sheet '{sheet_name}' columns: {list(df.columns)}")
                
                # Normalize column names for matching (handle NEWLINES and multiple spaces)
                def norm_col(c):
                    return " ".join(str(c).lower().replace('\n', ' ').split())
                
                # Log column names for debugging
                print(f"[POAI DEBUG] Sheet '{sheet_name}' columns (normalized): {[norm_col(c) for c in df.columns]}")
                
                # === 1. EXTRACT PROGRAM CODES ===
                # Broad search: any column containing 'código' AND ('programa' OR 'presupuestal')
                # NOW USES NORMALIZED COLUMN NAMES to catch 'Código\nPrograma'
                code_columns = [col for col in df.columns if 
                    any(pattern in norm_col(col) for pattern in CODE_PATTERNS)]
                
                print(f"[POAI DEBUG] Code columns found (pattern match): {code_columns}")
                
                # Fallback: columns with just 'código' that have numeric values
                if not code_columns:
                    code_columns = [col for col in df.columns if 
                        ('código' in norm_col(col) or 'codigo' in norm_col(col)) and
                        df[col].dropna().apply(lambda x: str(x).replace('.0', '').replace('.', '').isdigit()).any()]
                    print(f"[POAI DEBUG] Code columns found (fallback): {code_columns}")
                
                # --- FALLBACK 2: BRUTE FORCE SEARCH (If still no codes) ---
                if not code_columns:
                    st.warning("⚠️ No se detectaron columnas de código por nombre. Activando búsqueda profunda de valores...")
                    print("[POAI DEBUG] Starting Brute Force Search for codes...")
                    # Scan first 20 columns
                    for col in df.columns[:20]:
                        # Check sample values
                        sample = df[col].dropna().astype(str).tolist()
                        code_like_count = 0
                        for val in sample[:50]: # Check first 50 rows
                            v = val.replace('.0', '').strip()
                            # Check if it looks like a program code (2-5 digits, not a year like 2024-2030)
                            if v.isdigit() and 2 <= len(v) <= 5:
                                if not (len(v) == 4 and v.startswith("202")): # Avoid years like 2024, 2025
                                    code_like_count += 1
                        
                        if code_like_count > 0:
                            print(f"[POAI DEBUG] Brute force found potential column: {col} ({code_like_count} matches)")
                            code_columns.append(col)
                
                # --- FALLBACK 3: EXTREME BRUTE FORCE (Scan EVERYTHING) ---
                if not code_columns and not extracted_poai_codes:
                    st.warning("⚠️ Búsqueda profunda falló. Iniciando ESCANEO TOTAL de la hoja...")
                    print("[POAI DEBUG] Starting EXTREME Brute Force (Full Sheet Scan)...")
                    
                    # Flatten the entire dataframe to a list of strings
                    all_values = df.astype(str).values.flatten()
                    
                    # Find any 4-digit number that isn't a year (2020-2030)
                    potential_codes = []
                    for val in all_values[:5000]: # Check first 5000 cells
                        v = val.replace('.0', '').strip()
                        if v.isdigit() and len(v) == 4:
                            # Exclude years 2020-2030
                            if not (v.startswith("202") or v == "2030"):
                                potential_codes.append(v)
                    
                    # Get unique codes
                    potential_codes = sorted(list(set(potential_codes)))
                    
                    if potential_codes:
                        print(f"[POAI DEBUG] EXTREME SCAN FOUND: {potential_codes}")
                        st.success(f"✅ Códigos encontrados por escaneo total: {', '.join(potential_codes[:5])}")
                        # Add them directly to extracted codes
                        extracted_poai_codes.extend(potential_codes)
                    else:
                        print("[POAI DEBUG] EXTREME SCAN FOUND NOTHING.")

                # Name/description columns for program names
                name_columns = [col for col in df.columns if 
                    'programa' in norm_col(col) and 
                    'presupuestal' in norm_col(col) and
                    'código' not in norm_col(col) and 
                    'codigo' not in norm_col(col)]
                
                print(f"[POAI DEBUG] Name columns found: {name_columns}")
                
                # Debug: Show which columns were detected
                if code_columns:
                    st.info(f"🔍 Columnas de código encontradas en '{sheet_name}': {code_columns}")
                
                # Extract actual codes from the POAI
                for code_col in code_columns:
                    codes = df[code_col].dropna().astype(str).unique()
                    print(f"[POAI DEBUG] Raw codes in column '{code_col}': {list(codes)[:10]}")
                    
                    for code in codes:
                        # Clean the code - remove .0 if it's a float-like string
                        clean_code = str(code).replace('.0', '') if str(code).endswith('.0') else str(code)
                        clean_code = clean_code.strip()
                        
                        print(f"[POAI DEBUG] Processing code: '{code}' -> clean: '{clean_code}'")
                        
                        # Validate: must be 2-5 digits
                        if clean_code and clean_code != 'nan' and len(clean_code) >= 2 and len(clean_code) <= 5 and clean_code.isdigit():
                            print(f"[POAI DEBUG] Code '{clean_code}' PASSED validation")
                            
                            # Try to find matching program name
                            name_found = ""
                            if name_columns:
                                for name_col in name_columns:
                                    mask = df[code_col].astype(str).str.replace('.0', '') == clean_code
                                    names = df.loc[mask, name_col].dropna().unique()
                                    for name in names:
                                        if name and str(name) != 'nan':
                                            name_found = str(name).strip()
                                            break
                                    if name_found:
                                        break
                            
                            # Add code (with or without name)
                            if name_found:
                                full_code = f"{clean_code} - {name_found}"
                            else:
                                full_code = clean_code
                            
                            # Avoid duplicates
                            if clean_code not in [c.split(' - ')[0] for c in extracted_poai_codes]:
                                extracted_poai_codes.append(full_code)
                                print(f"[POAI DEBUG] ✅ EXTRACTED CODE: {full_code}")
                        else:
                            print(f"[POAI DEBUG] Code '{clean_code}' FAILED validation (len={len(clean_code)}, isdigit={clean_code.isdigit() if clean_code else 'N/A'})")
                
                # === 2. EXTRACT BPIN ===
                for col in df.columns:
                    if any(pattern in norm_col(col) for pattern in BPIN_PATTERNS):
                        bpin_values = df[col].dropna().astype(str).unique()
                        for bpin in bpin_values:
                            clean_bpin = str(bpin).replace('.0', '') if str(bpin).endswith('.0') else str(bpin)
                            if clean_bpin and len(clean_bpin) >= 8 and clean_bpin not in ['nan', 'NaN']:
                                extracted_poai_data['bpin'] = clean_bpin
                                break
                
                # === 3. EXTRACT PROJECT NAME ===
                for col in df.columns:
                    if any(pattern in norm_col(col) for pattern in PROJECT_NAME_PATTERNS):
                        names = df[col].dropna().astype(str).unique()
                        for name in names:
                            if name and len(str(name)) > 10 and str(name) not in ['nan', 'NaN']:
                                extracted_poai_data['nombre_proyecto'] = str(name).strip()
                                break
                        if 'nombre_proyecto' in extracted_poai_data:
                            break
                
                # === 4. EXTRACT SECTOR ===
                for col in df.columns:
                    if any(pattern in norm_col(col) for pattern in SECTOR_PATTERNS):
                        sectors = df[col].dropna().astype(str).unique()
                        for sector in sectors:
                            if sector and len(str(sector)) > 3 and str(sector) not in ['nan', 'NaN']:
                                extracted_poai_data['sector'] = str(sector).strip()
                                break
                        if 'sector' in extracted_poai_data:
                            break
                            
                # === 5. EXTRACT TOTAL VALUE ===
                for col in df.columns:
                    if any(pattern in norm_col(col) for pattern in VALUE_PATTERNS):
                        # Try to find numeric values
                        try:
                            values = pd.to_numeric(df[col], errors='coerce').dropna()
                            if not values.empty:
                                total_val = values.sum() if values.sum() > 0 else values.max()
                                if total_val > 1000:  # Reasonable minimum for a budget
                                    extracted_poai_data['valor_total'] = f"{total_val:,.0f}"
                                    break
                        except:
                            pass
            
            # === BUILD CONTEXT - CODES GO FIRST (CRITICAL) ===
            print(f"[POAI DEBUG] Total extracted codes: {len(extracted_poai_codes)}")
            print(f"[POAI DEBUG] Extracted codes list: {extracted_poai_codes}")
            
            if extracted_poai_codes:
                # Allow user to SELECT the correct project
                st.markdown("---")
                st.success(f"✅ Se encontraron **{len(extracted_poai_codes)}** proyectos en el POAI.")
                
                # Default to index 0, but allow selection
                selected_code = st.selectbox(
                    "📌 **Seleccione el Código del Proyecto para generar el MGA:**", 
                    options=extracted_poai_codes,
                    index=0,
                    help="El POAI contiene múltiples proyectos. Seleccione el que corresponde al MGA que desea generar."
                )
                
                # Use ONLY the selected code in the critical section
                poai_critical_section = f"""
╔══════════════════════════════════════════════════════════════════╗
║  🚨🚨🚨 CÓDIGO OBJETIVO DEL PROYECTO - ¡USAR SOLO ESTE! 🚨🚨🚨      ║
╠══════════════════════════════════════════════════════════════════╣
  🚨 CÓDIGO SELECCIONADO: {selected_code}
╠══════════════════════════════════════════════════════════════════╣
║  ❌ NO uses ningún otro código del archivo (401, 406, etc.)      ║
║  ✅ ESTE es el único código válido para este documento.          ║
╚══════════════════════════════════════════════════════════════════╝
"""
                print(f"[POAI DEBUG] ✅ User selected code: {selected_code}")
                # We do NOT show the list in st.success anymore to avoid clutter, the selectbox is enough
            else:
                print("[POAI DEBUG] ⚠️ NO CODES EXTRACTED - AI will fabricate codes!")
            
            # Add other extracted data to critical section
            if extracted_poai_data:
                poai_critical_section += "\n=== DATOS CLAVE EXTRAÍDOS DEL POAI ===\n"
                if 'bpin' in extracted_poai_data:
                    poai_critical_section += f"  BPIN: {extracted_poai_data['bpin']}\n"
                if 'nombre_proyecto' in extracted_poai_data:
                    poai_critical_section += f"  PROYECTO: {extracted_poai_data['nombre_proyecto']}\n"
                if 'sector' in extracted_poai_data:
                    poai_critical_section += f"  SECTOR: {extracted_poai_data['sector']}\n"
                if 'valor_total' in extracted_poai_data:
                    poai_critical_section += f"  VALOR TOTAL: ${extracted_poai_data['valor_total']} COP\n"
                
                # Show extracted data in UI
                st.info(f"📊 Datos extraídos: {', '.join([f'{k}: {v[:30]}...' if len(str(v)) > 30 else f'{k}: {v}' for k, v in extracted_poai_data.items()])}")
            
            # CRITICAL: Put extracted codes FIRST, then raw POAI data
            context_dump = poai_critical_section + f"\n\n=== DATOS COMPLETOS DEL POAI ===\n{poai_text[:10000]}" + context_dump
            extracted_summary.append(f"✅ POAI: {len(xlsx.sheet_names)} hojas, {len(extracted_poai_codes)} códigos")
            
            # Final debug: show what's being sent to AI
            print(f"[POAI DEBUG] Context dump starts with (first 500 chars):\n{context_dump[:500]}")
        except Exception as e:
            st.error(f"❌ Error procesando POAI: {e}")
    
    # Process Development Plan (PDF) - Use cheap model summarization
    dev_plan_summary = {}
    if dev_plan_file:
        try:
            from extractors.document_data_extractor import summarize_development_plan
            
            # Progressive loading indicator
            with st.spinner("📋 Resumiendo Plan de Desarrollo con IA... (esto ahorra tokens)"):
                dev_plan_summary = summarize_development_plan(dev_plan_file)
            
            if dev_plan_summary.get("success"):
                raw_len = dev_plan_summary.get("raw_text_length", 0)
                sum_len = dev_plan_summary.get("summary_length", 0)
                reduction = int((1 - sum_len / max(raw_len, 1)) * 100) if raw_len > 0 else 0
                
                extracted_summary.append(f"✅ Plan: {raw_len:,}→{sum_len:,} chars ({reduction}% reducción)")
                
                # Show summary preview in expander
                with st.expander("📊 Resumen del Plan de Desarrollo (generado por IA)", expanded=False):
                    if dev_plan_summary.get("resumen_global"):
                        st.write(f"**Resumen:** {dev_plan_summary['resumen_global']}")
                    
                    datos_prog = dev_plan_summary.get("datos_programa", {})
                    if datos_prog.get("codigos_programa"):
                        st.write(f"**Programas:** {', '.join(datos_prog['codigos_programa'])}")
                    if datos_prog.get("metas"):
                        st.write(f"**Metas:** {', '.join(datos_prog['metas'][:3])}")
            else:
                # Fallback to basic extraction if summarization fails
                st.warning(f"⚠️ Resumen falló: {dev_plan_summary.get('error', 'Unknown')}. Usando extracción básica.")
                import fitz
                dev_plan_file.seek(0)
                pdf = fitz.open(stream=dev_plan_file.read(), filetype="pdf")
                dev_text = ""
                for page in pdf:
                    dev_text += page.get_text()
                context_dump += f"\n\n=== PLAN DE DESARROLLO ===\n{dev_text[:20000]}"
                extracted_summary.append(f"✅ Plan: {len(dev_text):,} chars (básico)")
                
        except Exception as e:
            st.error(f"❌ Error procesando Plan: {e}")
    
    # Process Additional Info
    if basic_info_file:
        try:
            if basic_info_file.name.endswith('.pdf'):
                import fitz
                pdf = fitz.open(stream=basic_info_file.read(), filetype="pdf")
                basic_text = ""
                for page in pdf:
                    basic_text += page.get_text()[:3000]
                context_dump += f"\n\n=== INFORMACIÓN ADICIONAL ===\n{basic_text[:8000]}"
            else:
                import docx
                doc = docx.Document(basic_info_file)
                basic_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])[:8000]
                context_dump += f"\n\n=== INFORMACIÓN ADICIONAL ===\n{basic_text}"
            extracted_summary.append("✅ Info adicional cargada")
        except Exception as e:
            st.warning(f"⚠️ Error info adicional: {e}")
    
    # Show extraction status
    if extracted_summary:
        st.success(" | ".join(extracted_summary))
    
    # ============================================================
    # MINIMAL REQUIRED FIELDS - Only essentials
    # ============================================================
    st.markdown("---")
    st.markdown("### 📝 Datos Básicos (Mínimos Requeridos)")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        municipio = st.text_input("Municipio *", placeholder="Ej: San Pablo", key="mga_municipio")
        departamento = st.text_input("Departamento *", placeholder="Ej: Bolívar", key="mga_depto")
    
    with col2:
        responsable = st.text_input("Responsable *", placeholder="Ej: Carlos Gil", key="mga_responsable")
        cargo = st.text_input("Cargo", value="Secretario de Planeación", key="mga_cargo")
    
    with col3:
        sector = st.selectbox("Sector", OFFICIAL_SECTORS, key="mga_sector")
    
    # Letterhead
    st.markdown("---")
    letterhead_file = st.file_uploader(
        "📄 Membrete/Letterhead (.docx) - Opcional",
        type=["docx"],
        help="Plantilla con encabezado y pie de página institucional",
        key="mga_letterhead"
    )
    
    # Info box
    if not poai_file or not dev_plan_file:
        st.warning("⚠️ **Suba POAI y Plan de Desarrollo** para generar el MGA. La IA extraerá todos los datos necesarios de estos documentos.")
    else:
        st.info("✅ **Listo para generar.** La IA analizará los documentos y creará el MGA completo.")
    
    # Plan de Desarrollo details (optional manual input)
    with st.expander("📋 Datos Plan de Desarrollo (Opcional)", expanded=False):
        st.caption("La IA extrae estos datos automáticamente del Plan de Desarrollo, pero puede ingresarlos manualmente si lo desea.")
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Plan Departamental**")
            plan_depto_nombre = st.text_input("Nombre del Plan", key="pd_depto_nombre", placeholder="Ej: Plan de Desarrollo Departamental 2024-2027")
            plan_depto_estrategia = st.text_input("Estrategia", key="pd_depto_estrategia", placeholder="Ej: Desarrollo territorial sostenible")
            plan_depto_programa = st.text_input("Programa", key="pd_depto_programa", placeholder="Ej: 4501 - Mejoramiento de vías")
        
        with col2:
            st.markdown("**Plan Municipal**")
            plan_mun_nombre = st.text_input("Nombre del Plan", key="pd_mun_nombre", placeholder="Ej: Plan de Desarrollo Municipal 2024-2027")
            plan_mun_estrategia = st.text_input("Estrategia", key="pd_mun_estrategia", placeholder="Ej: Infraestructura para el progreso")
            plan_mun_programa = st.text_input("Programa", key="pd_mun_programa", placeholder="Ej: 4502 - Vías terciarias")
    
    return {
        "municipio": municipio,
        "departamento": departamento,
        "entidad": f"Alcaldía de {municipio}" if municipio else "",
        "bpin": "",  # AI extracts from POAI
        "identificador": "",  # AI extracts from POAI
        "nombre_proyecto": "",  # AI extracts from POAI
        "valor_total": "",  # AI extracts from POAI
        "duracion": "365",
        "fecha_creacion": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
        "plan_nacional": "(2022-2026) Colombia Potencia Mundial de la Vida",
        "plan_departamental": {
            "nombre": plan_depto_nombre,
            "estrategia": plan_depto_estrategia,
            "programa": plan_depto_programa
        } if plan_depto_nombre or plan_depto_estrategia or plan_depto_programa else "",
        "plan_municipal": {
            "nombre": plan_mun_nombre,
            "estrategia": plan_mun_estrategia,
            "programa": plan_mun_programa
        } if plan_mun_nombre or plan_mun_estrategia or plan_mun_programa else "",
        "responsable": responsable,
        "cargo": cargo,
        "sector": sector,
        "letterhead_file": letterhead_file,
        "context_dump": context_dump,  # POAI + additional file content
        "dev_plan_summary": dev_plan_summary  # Structured summary from cheap model
    }


def render_unified_form():
    """Render unified input form for generating ALL documents at once"""
    
    # Data extraction option
    extracted = render_data_upload_option("mga_subsidios", "unified")
    
    # 1. INFORMACIÓN GENERAL
    st.markdown('<p class="section-header">1. Información General del Proyecto (Todos)</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        municipio = st.text_input("Municipio *", value=extracted.get("municipio", ""), placeholder="Ej: San Pablo", key="uni_municipio")
        departamento = st.text_input("Departamento *", value=extracted.get("departamento", ""), placeholder="Ej: Bolívar", key="uni_depto")
        entidad = st.text_input("Entidad *", value=extracted.get("entidad", ""), placeholder="Ej: Alcaldía Municipal", key="uni_entidad")
        bpin = st.text_input("Código BPIN", value=extracted.get("bpin", ""), placeholder="Ej: 2024000001367", key="uni_bpin")
    
    with col2:
        nombre_proyecto = st.text_input("Nombre del Proyecto *", value=extracted.get("nombre_proyecto", ""), placeholder="Ej: Construcción de...", key="uni_proyecto")
        valor_total = st.text_input("Valor Total (COP) *", value=extracted.get("valor_total", ""), placeholder="Ej: 100000000", key="uni_valor")
        duracion = st.number_input("Duración (días)", min_value=1, value=int(extracted.get("duracion", "90")) if str(extracted.get("duracion", "90")).isdigit() else 90, key="uni_duracion")
        fecha_creacion = datetime.now().strftime("%d/%m/%Y") # Auto-generated

    # 2. RESPONSABLES
    st.markdown('<p class="section-header">2. Responsables y Autoridades</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        responsable = st.text_input("Responsable del Proyecto *", value=extracted.get("responsable", ""), placeholder="Ej: Juan Pérez", key="uni_responsable")
        cargo = st.text_input("Cargo del Responsable", value=extracted.get("cargo", "Secretario de Planeación"), key="uni_cargo")
    with col2:
        alcalde = st.text_input("Nombre del Alcalde (Certificaciones)", value=extracted.get("alcalde", ""), placeholder="Ej: María Rodríguez", key="uni_alcalde")

    # 3. CONTRATACIÓN Y JURÍDICA
    st.markdown('<p class="section-header">3. Detalles de Contratación (Estudios Previos / Sector)</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        obj_contractual = st.text_area("Objeto Contractual", value=extracted.get("objeto", ""), height=100, key="uni_objeto")
        modalidad = st.selectbox("Modalidad de Selección", ["Contratación Directa", "Licitación Pública", "Selección Abreviada", "Concurso de Méritos", "Mínima Cuantía", "Convenio Interadministrativo"], key="uni_modalidad")
        numero_contrato = st.text_input("Número de Contrato (Análisis Sector)", placeholder="Ej: 001-2025 (Si aplica)", key="uni_num_contrato")
    with col2:
        necesidad = st.text_area("Necesidad / Justificación / Descripción Problema", value=extracted.get("descripcion", ""), height=100, key="uni_necesidad", help="Se usará para Estudios Previos y DTS")
        alcance = st.text_area("Alcance (Actividades Principales)", value=extracted.get("alcance", ""), height=100, key="uni_alcance", help="Lista de actividades principales para Estudio Previos")
        fuente = st.text_input("Fuente de Financiación", value="Recursos Propios", key="uni_fuente")
        cdp_codigo = st.text_input("Código CDP (Opcional)", placeholder="Ej: 20260108-3", key="uni_cdp")
        
    st.markdown('<p class="section-header">3.1 Presupuesto (Detalle)</p>', unsafe_allow_html=True)
    rubros = st.text_area("Rubros / Desglose Presupuestal", value="Honorarios: 60%\nGastos Operativos: 40%", height=80, key="uni_rubros", help="Desglose para Estudios Previos")

    # 4. TÉCNICO Y SECTORIAL
    st.markdown('<p class="section-header">4. Información Técnica y Sectorial</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        sector_options = OFFICIAL_SECTORS
        sector = st.selectbox("Sector Económico", sector_options, key="uni_sector")
        programa = st.text_input("Programa (DTS)", placeholder="Ej: 4002 - Usuarios beneficiados...", key="uni_programa")
    with col2:
        unspsc = st.text_input("Códigos UNSPSC", placeholder="Ej: 43210000, 72101500", key="uni_unspsc")
        codigo_ciiu = st.text_input("Código CIIU", placeholder="Ej: 7110", key="uni_ciiu")
        subprograma = st.text_input("Subprograma (DTS)", placeholder="Ej: Agua de vida...", key="uni_subprograma")
        
    # 5. CONTEXTO / DUMP
    st.markdown('<p class="section-header">5. Contexto del Documento (Dump Data)</p>', unsafe_allow_html=True)
    st.info("ℹ️ Texto extraído del archivo. Se usará como contexto adicional para generar secciones faltantes.")
    
    context_dump = st.text_area(
        "Texto Extraído / Contexto Adicional",
        value=extracted.get("context_dump", ""),
        height=200,
        placeholder="Aquí aparecerá el texto completo extraído del archivo de datos...",
        key="uni_context_dump"
    )

    st.markdown('<p class="section-header">6. Planes de Desarrollo (MGA / Certificaciones)</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        plan_nacional = st.text_input("Plan Nacional", value=extracted.get("plan_nacional", "(2022-2026) Colombia Potencia Mundial de la Vida"), key="uni_pnd")
        plan_municipal = st.text_input("Plan Municipal", value=extracted.get("plan_municipal", ""), placeholder="Ej: San Pablo Mejor 2024-2027", key="uni_pm")
    with col2:
        plan_departamental = st.text_input("Plan Departamental", value=extracted.get("plan_departamental", ""), placeholder="Ej: Bolívar Me Enamora", key="uni_pd")

    st.markdown('<p class="section-header">Opcional: Membrete</p>', unsafe_allow_html=True)
    letterhead_file = st.file_uploader("Subir plantilla (.docx)", type=["docx"], key="uni_letterhead")

    # Construct unified data object
    timestamp = datetime.now().strftime("%d/%m/%Y")
    
    return {
        # Common
        "municipio": municipio,
        "departamento": departamento,
        "entidad": entidad,
        "bpin": bpin,
        "nombre_proyecto": nombre_proyecto,
        "valor_total": valor_total,
        "duracion": str(duracion),
        "fecha": timestamp,
        "fecha_creacion": timestamp,
        "responsable": responsable,
        "cargo": cargo,
        "alcalde": alcalde, # Certificaciones specific
        
        # Estudios Previos
        "objeto": obj_contractual,
        "necesidad": necesidad,
        "alcance": alcance, # Added in previous step
        "rubros": rubros, # Added in previous step
        "modalidad": modalidad,
        "plazo": str(duracion),
        "fuente_financiacion": fuente,
        "cdp_data": {"cdp": cdp_codigo, "fecha": timestamp, "valor": valor_total}, # Synthesize CDP data
        "tipo_proyecto": modalidad,
        "lugar": f"{municipio}, {departamento}",
        
        # Analisis Sector
        "sector": sector,
        "codigos_unspsc": unspsc,
        "codigo_ciiu": codigo_ciiu,
        "numero_contrato": numero_contrato if numero_contrato else "POR DEFINIR",
        "plan_desarrollo": plan_municipal, # Map to specific field
        
        # DTS
        "descripcion_problema": necesidad, # Reuse
        "programa": programa,
        "subprograma": subprograma, # Added in previous step
        "poblacion_beneficiada": "Población general del municipio",
        
        # MGA
        "plan_nacional": plan_nacional,
        "plan_departamental": plan_departamental,
        "plan_municipal": plan_municipal,
        "identificador": bpin,
        
        # Files
        "letterhead_file": letterhead_file,
        "context_dump": context_dump, # Raw text for fallback
    }


def generate_document(doc_type: str, data: dict, model: str):
    """Generate document using selected model"""
    try:
        # Get LLM based on selection
        llm = get_llm(model)
        
        # Get section toggles from session state
        section_toggles = st.session_state.get('section_toggles', {})
        data['section_toggles'] = section_toggles
        
        # Get generation mode and edit instructions
        generation_mode = st.session_state.get('generation_mode', 'crear_nuevo')
        data['generation_mode'] = generation_mode
        
        # Get edit instructions if in update mode
        if generation_mode == 'actualizar_existente':
            edit_instructions = st.session_state.get('edit_instructions', '')
            data['edit_instructions'] = edit_instructions
        
        # Create generator based on document type
        if doc_type == "estudios_previos":
            generator = EstudiosPreviosGenerator(llm)
            result = generator.generate_complete(data)
            return result.get("documento_completo", ""), result.get("filepath")
        elif doc_type == "analisis_sector":
            generator = AnalisisSectorGenerator(llm)
            result = generator.generate_complete(data)
            return result.get("documento_completo", ""), result.get("filepath")
        elif doc_type == "dts":
            generator = DTSGenerator(llm)
            result = generator.generate_complete(data)
            return result.get("documento_completo", ""), result.get("filepath")
        elif doc_type == "certificaciones":
            generator = CertificacionesGenerator(llm)
            result = generator.generate_complete(data)
            return result.get("documento_completo", ""), result.get("filepath")
        else:  # mga_subsidios
            generator = MGASubsidiosGenerator(llm)
            result = generator.generate_complete(data)
            return result.get("documento_completo", ""), result.get("filepath")
        
    except Exception as e:
        st.error(f"Error al generar documento: {str(e)}")
        return None, None


def run_generation_logic(doc_type: str, data: dict, model: str):
    """
    Shared generation logic to be used by both main panel and sidebar buttons.
    Returns: result_data (dict or tuple depending on mode)
    """
    # Check for skip validation (if not already checked by caller)
    if 'skip_validation' not in st.session_state:
        st.session_state.skip_validation = False
    
    # Validation is assumed to be done or skipped by the caller/UI state by the time we click Generate
    # For MGA Subsidios, we check context_dump (from uploaded files) instead of nombre_proyecto
    project_name = data.get("nombre_proyecto", data.get("proyecto", ""))
    context_dump = data.get("context_dump", "")
    
    # Allow generation if we have context_dump (files uploaded) OR project name
    if not project_name and not context_dump:
        st.error("⚠️ Por favor suba los archivos (POAI + Plan) o ingrese el nombre del proyecto.")
        return None

    # Clear previous generation state
    st.session_state.generated_content = None
    st.session_state.generated_file = None
    
    # Progress feedback
    with st.spinner(f"Generando {doc_type} con {model}..."):
        if doc_type == "unified":
            generator = UnifiedGenerator()
            # Progress bar for unified generation
            progress_bar = st.progress(0, text="Iniciando generación paralela de 5 documentos...")
            result = generator.generate_all(data, model)
            progress_bar.progress(100, text="Generación Completada!")
            return result
        else:
            # Individual generation
            content, filepath = generate_document(doc_type, data, model)
            if content and filepath:
                # Save generation to session state for persistence
                st.session_state.generated_content = content
                st.session_state.generated_file = filepath
                
                # Track in generation history
                from datetime import datetime
                st.session_state.generation_history.append({
                    "type": doc_type,
                    "time": datetime.now().strftime("%H:%M:%S"),
                    "file": os.path.basename(filepath) if filepath else "N/A"
                })
                st.session_state.last_generation_time = datetime.now()
                return (content, filepath)
            return None


def render_sidebar_generation_controls(doc_type: str, data: dict, selected_model: str, validation_issues: list):
    """Render download controls in the sidebar (only when file is ready)"""
    with st.sidebar:
        # Download Button (if file is ready) - this is the ONLY thing shown
        if st.session_state.generated_file and os.path.exists(st.session_state.generated_file):
            st.markdown("---")
            st.markdown("### 📥 Descarga Rápida")
            file_path = st.session_state.generated_file
            file_name = os.path.basename(file_path)
            
            with open(file_path, "rb") as f:
                st.download_button(
                    label=f"⬇️ {file_name}",
                    data=f,
                    file_name=file_name,
                    mime="application/pdf" if file_path.endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="sidebar_download_btn",
                    use_container_width=True
                )



def main():
    """Main application"""
    # Render sidebar and get selected model
    selected_model = render_sidebar()
    
    # Main header
    st.markdown(f'<p class="main-header">{APP_TITLE}</p>', unsafe_allow_html=True)
    st.markdown(f'<p class="sub-header">{APP_DESCRIPTION}</p>', unsafe_allow_html=True)
    
    # ═══════════════════════════════════════════════════════════════
    # PROMINENT MODE BANNER - Clear visual differentiation
    # ═══════════════════════════════════════════════════════════════
    mode = st.session_state.get('generation_mode', 'crear_nuevo')
    
    if mode == "crear_nuevo":
        st.markdown("""
        <div style="background: linear-gradient(90deg, #2E7D32 0%, #43A047 100%); 
                    color: white; padding: 12px 20px; border-radius: 8px; margin-bottom: 16px;
                    font-size: 16px; font-weight: 600; text-align: center;">
            🆕 MODO: CREAR DOCUMENTO NUEVO
            <span style="font-weight: 400; font-size: 13px; display: block; margin-top: 4px;">
                Complete el formulario y genere un documento desde cero
            </span>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div style="background: linear-gradient(90deg, #1565C0 0%, #1976D2 100%); 
                    color: white; padding: 12px 20px; border-radius: 8px; margin-bottom: 16px;
                    font-size: 16px; font-weight: 600; text-align: center;">
            🔄 MODO: ACTUALIZAR DOCUMENTO EXISTENTE
            <span style="font-weight: 400; font-size: 13px; display: block; margin-top: 4px;">
                Suba un PDF/DOCX en el sidebar → Procesar → Los datos se llenarán automáticamente
            </span>
        </div>
        """, unsafe_allow_html=True)
    
    # Document type selection
    st.markdown("**Seleccionar Tipo de Documento**")
    
    doc_type = st.radio(
        "Tipo de documento a generar:",
        options=["unified", "estudios_previos", "analisis_sector", "dts", "certificaciones", "mga_subsidios"],
        format_func=lambda x: {
            "unified": "🚀 Generar Todo (Unified Mode)",
            "estudios_previos": "Estudios Previos", 
            "analisis_sector": "Análisis del Sector", 
            "dts": "DTS (Documento Técnico)",
            "certificaciones": "Certificaciones",
            "mga_subsidios": "MGA"
        }.get(x, x),
        horizontal=True,
        label_visibility="collapsed"
    )
    
    st.markdown("---")
    
    # ═══════════════════════════════════════════════════════════════
    # DEDICATED UPDATE MODE VIEW
    # ═══════════════════════════════════════════════════════════════
    if mode == "actualizar_existente":
        # Get session state values
        doc = st.session_state.get("previous_document")
        instr = st.session_state.get("edit_instructions_text")
        target_pages = st.session_state.get("selected_edit_pages", [])
        
        if not doc or not instr:
             st.info("👈 **Para comenzar:** Configure las instrucciones y cargue su documento en el panel lateral.")
             
             col1, col2 = st.columns(2)
             with col1:
                 st.markdown("""
                 #### Pasos para actualizar:
                 1. Escriba en el sidebar qué cambios desea realizar
                 2. Cargue el archivo PDF o Word original
                 3. Presione el botón de actualizar abajo
                 """)
             with col2:
                 st.markdown("""
                 #### Capacidades:
                 - Actualización de fechas y vigencias
                 - Ajuste de presupuestos (conservando tablas)
                 - Cambio de responsables
                 - Modificación de objetos y alcances
                 """)
             
             # Tips section
             with st.expander("💡 Tips para mejores resultados", expanded=False):
                 st.markdown("""
**Escriba instrucciones específicas:**

✅ **Bueno:** "Cambiar el nombre del responsable de 'Juan López' a 'María García' en todas las páginas"

❌ **Evitar:** "Actualizar responsable"

---

**Ejemplos de prompts efectivos:**

• "Cambiar la fecha de impresión de 2025 a 2026"
• "Reemplazar 'Alcaldía de Bogotá' por 'Alcaldía de Medellín'"
• "Actualizar el valor total de $100,000,000 a $150,000,000"
• "Cambiar el código BPIN de 2024000001234 a 2026000005678"

---

**Para evitar cambios no deseados:**

• Sea específico con el texto exacto a cambiar
• Indique qué NO debe cambiar si es necesario
• Use páginas específicas si solo necesita editar secciones
                 """)
        
        else:
             # Show ready status
             st.markdown("""
             <div style="background: linear-gradient(90deg, #1565C0 0%, #42A5F5 100%); 
                         color: white; padding: 16px 24px; border-radius: 10px; margin-bottom: 20px;">
                 <h3 style="margin: 0; color: white;">✅ Documento Listo para Editar</h3>
             </div>
             """, unsafe_allow_html=True)
             
             # Status cards
             col1, col2, col3 = st.columns(3)
             
             with col1:
                 st.metric(label="📄 Documento", value=doc.name[:20] + "..." if len(doc.name) > 20 else doc.name)
             
             with col2:
                 page_info = f"Páginas: {', '.join(map(str, target_pages))}" if target_pages else "Todas las páginas"
                 st.metric(label="📑 Objetivo", value=page_info[:20])
             
             with col3:
                 st.metric(label="📝 Instrucciones", value=f"{len(instr)} caracteres")
             
             # Instructions preview
             with st.expander("👁️ Ver instrucciones de edición", expanded=True):
                 st.info(instr)
             
             st.markdown("---")

    # Render appropriate form (Only for NEW mode usually, but we keep it for now as background data structure)
    if mode == "crear_nuevo":
        if doc_type == "unified":
            data = render_unified_form()
        elif doc_type == "estudios_previos":
            data = render_estudios_previos_form()
        elif doc_type == "analisis_sector":
            data = render_analisis_sector_form()
        elif doc_type == "dts":
            data = render_dts_form()
        elif doc_type == "certificaciones":
            data = render_certificaciones_form()
        else:  # mga_subsidios
            data = render_mga_subsidios_form()
    else:
        # For update mode, we just need basic data or empty data, the generator will use the uploaded file + instructions
        data = {
            "edit_mode": True,
            "instructions": st.session_state.get("edit_instructions_text", ""),
            "original_file": st.session_state.get("previous_document")
        }

    # ═══════════════════════════════════════════════════════════════
    # PRE-GENERATION VALIDATION PANEL (Only for Create New)
    # ═══════════════════════════════════════════════════════════════
    
    validation_issues = []
    if mode == "crear_nuevo":
        # Run validation on current data
        validation_issues = validate_form_data(data, doc_type)
        
        if validation_issues:
            st.markdown("### 🔍 Validación de Datos")
            
            critical_count = len([i for i in validation_issues if i[1] == "critical"])
            warning_count = len([i for i in validation_issues if i[1] == "warning"])
            info_count = len([i for i in validation_issues if i[1] == "info"])
            
            if critical_count > 0:
                st.error(f"⛔ {critical_count} problema(s) crítico(s) encontrado(s)")
            if warning_count > 0:
                st.warning(f"⚠️ {warning_count} campo(s) recomendado(s) faltante(s)")
            
            with st.expander("Ver detalles de validación", expanded=True):
                for field, severity, message in validation_issues:
                    if severity == "critical":
                        st.error(message)
                    elif severity == "warning":
                        st.warning(message)
                    else:
                        st.caption(message)
            
            st.caption("💡 Puede corregir los campos arriba o continuar de todos modos.")
        
        st.markdown("---")
    
    # Call sidebar generation controls
    render_sidebar_generation_controls(doc_type, data, selected_model, validation_issues)

    # Generate button (Main Panel)
    btn_label = "🚀 Generar Documento(s)" if mode == "crear_nuevo" else "🚀 Actualizar Documento"
    
    if st.button(btn_label, type="primary", use_container_width=True):
        # Check validation (only for new)
        has_fake_data = any("PROHIBIDO" in issue[2] for issue in validation_issues if len(issue) > 2)
        if mode == "crear_nuevo" and has_fake_data:
            st.error("⛔ No se puede generar con datos de ejemplo/prueba. Por favor use datos reales.")
        elif mode == "actualizar_existente" and (not st.session_state.get("previous_document") or not st.session_state.get("edit_instructions_text")):
             st.error("⚠️ Falta información: Por favor suba un documento e indique las instrucciones de edición en el menú lateral.")
        elif mode == "actualizar_existente":
            # === EDIT MODE LOGIC ===
            with st.spinner("🔄 Analizando documento y aplicando ediciones..."):
                try:
                    # Get LLM for AI analysis
                    llm = get_llm(selected_model)
                    
                    # Get file and instructions
                    uploaded_file = st.session_state.get("previous_document")
                    instructions = st.session_state.get("edit_instructions_text", "")
                    target_pages = st.session_state.get("selected_edit_pages", [])
                    
                    # Run the editor
                    result = edit_mga_document(
                        file=uploaded_file,
                        user_prompt=instructions,
                        llm=llm,
                        target_pages=target_pages if target_pages else None
                    )
                    
                    if result.get("success"):
                        st.success("✅ ¡Documento editado exitosamente!")
                        
                        # Show edits applied
                        edits_applied = result.get("edits_applied", [])
                        if edits_applied:
                            with st.expander("📋 Cambios realizados", expanded=True):
                                for edit in edits_applied:
                                    if edit.get("status") == "applied":
                                        st.markdown(f"✅ **Original:** `{edit.get('original', 'N/A')}`")
                                        st.markdown(f"   **Nuevo:** `{edit.get('new', 'N/A')}`")
                                    elif edit.get("status") == "not_found":
                                        st.warning(f"⚠️ No encontrado: `{edit.get('original', 'N/A')}`")
                                    elif edit.get("error"):
                                        st.error(f"❌ Error: {edit.get('error')}")
                        
                        # Summary
                        if result.get("summary"):
                            st.info(f"📝 Resumen: {result.get('summary')}")
                        
                        # Download button
                        edited_file = result.get("edited_file")
                        file_name = result.get("file_name", "MGA_editado.docx")
                        
                        if edited_file:
                            file_ext = file_name.split('.')[-1].lower()
                            mime_type = "application/pdf" if file_ext == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            
                            st.download_button(
                                label=f"⬇️ Descargar {file_name}",
                                data=edited_file,
                                file_name=file_name,
                                mime=mime_type,
                                key="download_edited_doc",
                                use_container_width=True
                            )
                    else:
                        st.error(f"❌ Error al editar: {result.get('error', 'Error desconocido')}")
                        
                except Exception as e:
                    st.error(f"❌ Error inesperado: {str(e)}")
                    import traceback
                    with st.expander("Ver detalles del error"):
                        st.code(traceback.format_exc())
        
        elif mode == "generar_desde_mga":
            # === GENERATE FROM EXISTING MGA LOGIC ===
            mga_file = st.session_state.get("mga_source_file")
            docs_to_gen = st.session_state.get("docs_to_generate_from_mga", [])
            
            if not mga_file:
                st.error("⚠️ Por favor suba un documento MGA en el menú lateral.")
            elif not docs_to_gen:
                st.error("⚠️ Por favor seleccione al menos un documento a generar.")
            else:
                with st.spinner("📄 Extrayendo datos del MGA y generando documentos..."):
                    try:
                        # Extract text from MGA
                        from extractors.document_data_extractor import summarize_development_plan
                        
                        # Get summary from MGA file
                        mga_summary = summarize_development_plan(mga_file)
                        
                        if not mga_summary.get("success"):
                            # Fallback: manual extraction
                            import fitz
                            mga_file.seek(0)
                            pdf = fitz.open(stream=mga_file.read(), filetype="pdf")
                            mga_text = ""
                            for page in pdf:
                                mga_text += page.get_text()
                            mga_file.seek(0)
                            
                            context_data = mga_text[:15000]
                        else:
                            context_data = str(mga_summary)
                        
                        # Get LLM
                        llm = get_llm(selected_model)
                        generated_files = []
                        
                        for doc_type_to_gen in docs_to_gen:
                            if doc_type_to_gen == "analisis_sector":
                                from generators.analisis_sector_generator import AnalisisSectorGenerator
                                generator = AnalisisSectorGenerator(llm)
                                result = generator.generate({
                                    "context_dump": context_data,
                                    "entidad": data.get("entidad", ""),
                                    "sector": data.get("sector", ""),
                                    "municipio": data.get("municipio", ""),
                                })
                            elif doc_type_to_gen == "estudios_previos":
                                from generators.estudios_previos_generator import EstudiosPreviosGenerator
                                generator = EstudiosPreviosGenerator(llm)
                                result = generator.generate({
                                    "context_dump": context_data,
                                    "entidad": data.get("entidad", ""),
                                    "municipio": data.get("municipio", ""),
                                })
                            
                            if result:
                                generated_files.append({
                                    "type": doc_type_to_gen,
                                    "file": result.get("file_path", ""),
                                    "status": "success" if result.get("success") else "error",
                                    "error": result.get("error")
                                })
                        
                        # Show results
                        if generated_files:
                            st.success(f"✅ Se generaron {len(generated_files)} documento(s) desde el MGA")
                            for gf in generated_files:
                                if gf["status"] == "success" and gf["file"]:
                                    fname = os.path.basename(gf["file"])
                                    with open(gf["file"], "rb") as f:
                                        st.download_button(
                                            label=f"⬇️ Descargar {gf['type'].replace('_', ' ').title()}",
                                            data=f.read(),
                                            file_name=fname,
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                        )
                                else:
                                    st.error(f"❌ Error en {gf['type']}: {gf.get('error', 'Unknown')}")
                        
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")
                        import traceback
                        with st.expander("Ver detalles"):
                            st.code(traceback.format_exc())
        
        else:
            # === NEW DOCUMENT GENERATION LOGIC ===
            result = run_generation_logic(doc_type, data, selected_model)
            
            if result:
                if doc_type == "unified":
                    # Unified result handling
                    if result["success"]:
                        st.success(f"✅ Se generaron exitosamente {len(result['results'])} documentos.")
                        
                        # Show individual results
                        for res in result["results"]:
                            if res["status"] == "success":
                                file_name = os.path.basename(res["file"])
                                st.write(f"📄 {res['type'].replace('_', ' ').title()}: **{file_name}**")
                            else:
                                st.error(f"❌ Error en {res['type']}: {res.get('error', 'Unknown')}")
                        
                        # Unified ZIP Download
                        if result.get("zip_file"):
                            with open(result["zip_file"], "rb") as f:
                                st.download_button(
                                    label="⬇️ Descargar TODOS los Documentos (ZIP)",
                                    data=f,
                                    file_name=os.path.basename(result["zip_file"]),
                                    mime="application/zip",
                                    key="unified_download"
                                )
                    else:
                        st.error("❌ Ocurrió un error al generar los documentos.")
                        if "error" in result:
                            st.error(f"Detalle: {result['error']}")
                        if "results" in result and result["results"]:
                             with st.expander("Ver detalles de errores"):
                                 st.json(result["results"])
                else:
                    # Individual result handling
                    content, filepath = result
                    st.success("✅ Documento generado exitosamente!")
                    
                    # Download button
                    with open(filepath, "rb") as f:
                        file_name = os.path.basename(filepath)
                        st.download_button(
                            label="⬇️ Descargar Documento (Word/PDF)",
                            data=f,
                            file_name=file_name,
                            mime="application/pdf" if filepath.endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    # Preview expander
                    with st.expander("Ver contenido generado", expanded=False):
                        st.markdown(content)
        
    # Reset button
    if st.button("Generar Nuevo Documento"):
        st.session_state.generated_content = None
        st.session_state.generated_file = None
        st.rerun()
    
    # ═══════════════════════════════════════════════════════════
    # FOOTER - Help, Diagnostics, History
    # ═══════════════════════════════════════════════════════════
    st.markdown("---")
    
    footer_col1, footer_col2, footer_col3 = st.columns(3)
    
    with footer_col1:
        with st.expander("❓ Ayuda Rápida"):
            st.markdown("""
            **Modos de Uso:**
            - 🆕 **Nuevo**: Crear documento desde cero
            - 🔄 **Actualizar**: Editar documento existente
            
            **Tipos de Documento:**
            - 📋 Estudios Previos
            - 📊 Análisis del Sector
            - 📄 DTS
            - ✅ Certificaciones
            - 📑 MGA Subsidios
            - 🔥 Modo Unificado (todos)
            
            **Tips:**
            - Use "Modo Unificado" para generar todos
            - Suba archivos de apoyo para mejor precisión
            - Revise siempre el documento generado
            """)
    
    with footer_col2:
        with st.expander("📊 Historial de Generación"):
            if st.session_state.generation_history:
                for item in st.session_state.generation_history[-5:]:  # Last 5
                    st.caption(f"⏰ {item['time']} - {item['type']}: {item['file']}")
            else:
                st.caption("No hay generaciones recientes")
    
    with footer_col3:
        with st.expander("🔧 Diagnóstico"):
            api_status = check_api_status()
            for provider, status in api_status.items():
                if status["configured"]:
                    st.caption(f"✅ {provider}: {status['key_preview']}")
                else:
                    st.caption(f"❌ {provider}: No configurado")
            st.caption(f"📱 Versión: {st.session_state.app_version}")
            if st.session_state.last_error:
                st.caption(f"⚠️ Último error: {st.session_state.last_error}")


if __name__ == "__main__":
    main()
