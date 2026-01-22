"""
Markdown to Word Converter - MGA Template Style
Each section: gray header + bordered content (one continuous border)
Tables inside sections share the section border (no extra boxing)
<br> converts to line breaks
"""

import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Colors
HEADER_COLOR = "BFBFBF"  # Gray
TABLE_HEADER_COLOR = "BFBFBF"
BORDER_COLOR = "000000"


class MarkdownToWordConverter:
    """Converts markdown to Word matching MGA template"""
    
    def __init__(self, doc: Document):
        self.doc = doc
    
    def add_formatted_content(self, markdown_text: str):
        """Parse markdown and add formatted content"""
        # Pre-process: convert <br> tags to newlines
        markdown_text = re.sub(r'<br\s*/?>', '\n', markdown_text, flags=re.IGNORECASE)
        
        lines = markdown_text.split('\n')
        i = 0
        
        # Collect sections
        sections = []
        current_section = None
        current_content = []
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Skip empty at start
            if not stripped and current_section is None:
                i += 1
                continue
            
            # Check for section heading
            is_section = False
            section_title = None
            
            if stripped.startswith('##'):
                section_title = stripped.lstrip('#').strip()
                is_section = True
            elif re.match(r'^\d+\.?\s+[A-ZÁÉÍÓÚÑ]', stripped):
                section_title = stripped
                is_section = True
            
            if is_section:
                if current_section is not None:
                    sections.append((current_section, current_content))
                current_section = section_title
                current_content = []
                i += 1
                continue
            
            # Skip horizontal rules
            if stripped in ['---', '***', '___']:
                i += 1
                continue
            
            # Add to content
            if current_section is not None:
                current_content.append(line)
            else:
                if stripped:
                    self._add_direct_para(stripped)
            
            i += 1
        
        # Save last section
        if current_section is not None:
            sections.append((current_section, current_content))
        
        # Render sections
        for title, content_lines in sections:
            self._add_section(title, content_lines)
    
    def _add_direct_para(self, text: str):
        """Add paragraph without box"""
        para = self.doc.add_paragraph()
        self._add_text_with_formatting(para, text)
        para.paragraph_format.space_after = Pt(6)
    
    def _add_section(self, title: str, content_lines: list):
        """Add section with gray header + content in one bordered box"""
        # Create 2-row table (header + content)
        table = self.doc.add_table(rows=2, cols=1)
        self._set_table_borders(table)
        
        # Remove spacing between rows
        table.autofit = True
        
        # Header row
        header_cell = table.rows[0].cells[0]
        header_cell.text = self._clean_text(title).upper()
        self._shade_cell(header_cell, HEADER_COLOR)
        for para in header_cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        
        # Content row - no extra spacing
        content_cell = table.rows[1].cells[0]
        content_cell.paragraphs[0].clear()
        
        # Render content
        self._render_content(content_cell, content_lines)
        
        # Small space after section
        spacing_para = self.doc.add_paragraph()
        spacing_para.paragraph_format.space_after = Pt(6)
    
    def _render_content(self, cell, lines: list):
        """Render content in cell - tables, lists, paragraphs"""
        i = 0
        is_first = True
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            if not stripped:
                i += 1
                continue
            
            # Check for table
            if stripped.startswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._add_table_in_cell(cell, table_lines, is_first)
                is_first = False
                continue
            
            # Get or create paragraph
            if is_first:
                para = cell.paragraphs[0]
                is_first = False
            else:
                para = cell.add_paragraph()
            
            # Check for bullet
            if stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('• '):
                para.paragraph_format.left_indent = Cm(0.5)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.space_before = Pt(1)
                self._add_text_with_formatting(para, "• " + stripped[2:])
                i += 1
                continue
            
            # Check for sub-bullet
            if stripped.startswith('o ') or stripped.startswith('  ') and '-' in stripped[:5]:
                para.paragraph_format.left_indent = Cm(1.0)
                para.paragraph_format.space_after = Pt(1)
                clean = re.sub(r'^[o\s\-•]+', '', stripped)
                self._add_text_with_formatting(para, "○ " + clean)
                i += 1
                continue
            
            # Check for numbered list
            num_match = re.match(r'^(\d+)\.\s*(.+)', stripped)
            if num_match:
                para.paragraph_format.left_indent = Cm(0.5)
                para.paragraph_format.space_after = Pt(1)
                self._add_text_with_formatting(para, f"{num_match.group(1)}. {num_match.group(2)}")
                i += 1
                continue
            
            # Regular paragraph
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.space_before = Pt(1)
            self._add_text_with_formatting(para, stripped)
            i += 1
    
    def _add_table_in_cell(self, parent_cell, table_lines: list, is_first: bool):
        """Add table in cell - NO extra outer border, just grid lines"""
        if len(table_lines) < 2:
            return
        
        # Parse header
        headers = [c.strip() for c in table_lines[0].split('|') if c.strip()]
        
        # Skip separator
        start = 1
        if len(table_lines) > 1 and '---' in table_lines[1]:
            start = 2
        
        # Parse rows
        rows = []
        for line in table_lines[start:]:
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                rows.append(cells)
        
        if not headers:
            return
        
        num_cols = len(headers)
        num_rows = len(rows) + 1
        
        # Create table - use Table Grid style which has borders
        tbl = parent_cell.add_table(rows=num_rows, cols=num_cols)
        tbl.style = 'Table Grid'
        tbl.autofit = True
        
        # Header row with gray
        for idx, h in enumerate(headers):
            if idx < num_cols:
                c = tbl.rows[0].cells[idx]
                # Handle line breaks in cell text
                c.text = ""
                para = c.paragraphs[0]
                self._add_text_with_formatting(para, self._clean_text(h))
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                self._shade_cell(c, TABLE_HEADER_COLOR)
        
        # Data rows
        for row_idx, row_data in enumerate(rows):
            for col_idx, text in enumerate(row_data):
                if col_idx < num_cols:
                    c = tbl.rows[row_idx + 1].cells[col_idx]
                    c.text = ""
                    para = c.paragraphs[0]
                    self._add_text_with_formatting(para, self._clean_text(text))
                    for run in para.runs:
                        run.font.size = Pt(9)
    
    def _add_text_with_formatting(self, para, text: str):
        """Add text with bold formatting and line breaks"""
        # Handle <br> as line break
        text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
        
        # Split by \n for line breaks
        parts = text.split('\n')
        
        for idx, part in enumerate(parts):
            # Handle **bold**
            pattern = r'(\*\*(.+?)\*\*|([^*]+))'
            for match in re.finditer(pattern, part):
                m = match.group(0)
                if m.startswith('**') and m.endswith('**'):
                    run = para.add_run(m[2:-2])
                    run.bold = True
                    run.font.size = Pt(10)
                else:
                    run = para.add_run(m)
                    run.font.size = Pt(10)
            
            # Add line break if not last part
            if idx < len(parts) - 1:
                run = para.add_run()
                run.add_break()
    
    def _shade_cell(self, cell, color: str):
        """Add background color"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _set_table_borders(self, table):
        """Set black borders on table"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        tblBorders = OxmlElement('w:tblBorders')
        for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), BORDER_COLOR)
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
    
    def _clean_text(self, text: str) -> str:
        """Clean markdown from text"""
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        text = re.sub(r'^#+\s*', '', text)
        return text.strip()


def add_signature_table(doc: Document, data: dict):
    """Add signature section"""
    table = doc.add_table(rows=2, cols=1)
    
    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), BORDER_COLOR)
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    
    # Header
    h = table.rows[0].cells[0]
    h.text = "RESPONSABLES"
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), HEADER_COLOR)
    shading.set(qn('w:val'), 'clear')
    h._tc.get_or_add_tcPr().append(shading)
    for para in h.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)
    
    # Content
    c = table.rows[1].cells[0]
    c.paragraphs[0].clear()
    
    sig = c.add_table(rows=1, cols=2)
    
    left = sig.rows[0].cells[0]
    lp = left.paragraphs[0]
    r1 = lp.add_run(f"NOMBRE: {data.get('responsable', '').upper()}")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.add_break()
    r2 = lp.add_run(f"CARGO: {data.get('cargo', 'Secretario de Planeación Municipal')}")
    r2.bold = True
    r2.font.size = Pt(10)
    
    right = sig.rows[0].cells[1]
    rp = right.paragraphs[0]
    rf = rp.add_run("FIRMA:")
    rf.bold = True
    rf.font.size = Pt(10)
