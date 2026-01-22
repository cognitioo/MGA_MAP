"""
MGA Subsidios Document Builder
Generates Word documents for complete MGA Subsidios (24 pages)
"""

import os
import re
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MGASubsidiosBuilder:
    """Builder for MGA Subsidios Word documents (24 pages)"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.doc = None
    
    def _safe_list(self, value, default=None):
        """Safely convert a value to a list, returning default if not iterable"""
        if default is None:
            default = []
        if isinstance(value, list):
            return value
        return default
    
    def _safe_str(self, value, default: str = "") -> str:
        """Safely convert value to string (handles dict, list, None)"""
        import json
        if value is None:
            return default
        if isinstance(value, str):
            return value
        if isinstance(value, dict):
            return json.dumps(value, ensure_ascii=False)
        if isinstance(value, list):
            return ", ".join(str(v) for v in value)
        return str(value)
    
    def build(self, data: dict, ai_content: dict, letterhead_file=None) -> str:
        """Build the MGA Subsidios document"""
        self._has_letterhead = letterhead_file is not None
        self._data = data  # Store data for access in page methods
        
        if letterhead_file:
            self.doc = self._load_template(letterhead_file)
        else:
            self.doc = Document()
        
        self._apply_styles()
        
        # Page 1 - Datos Básicos (ROBUST DATA MAPPING)
        page1_content = ai_content.get("pagina_1_datos_basicos", {})
        
        # Merge with form data to ensure no empty fields
        # If AI returns empty/missing, use the form data directly
        page1_content["nombre"] = page1_content.get("nombre") or data.get("nombre_proyecto") or ""
        page1_content["nombre_proyecto"] = data.get("nombre_proyecto") or ""
        page1_content["codigo_bpin"] = page1_content.get("codigo_bpin") or data.get("bpin") or ""
        page1_content["identificador"] = page1_content.get("identificador") or data.get("identificador") or ""
        page1_content["formulador_ciudadano"] = page1_content.get("formulador_ciudadano") or data.get("responsable") or ""
        page1_content["formulador_oficial"] = page1_content.get("formulador_oficial") or data.get("responsable") or ""
        page1_content["fecha_creacion"] = page1_content.get("fecha_creacion") or data.get("fecha_creacion") or datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        
        # Sector is tricky, try to get from data if available, or keep AI's
        if not page1_content.get("sector") and data.get("sector"):
             page1_content["sector"] = data.get("sector")
             
        self._add_page_1_datos_basicos(page1_content)
        
        self.doc.add_page_break()
        
        # Page 2 - Plan de Desarrollo
        self._add_page_2_plan_desarrollo(ai_content.get("pagina_2_plan_desarrollo", {}))
        
        self.doc.add_page_break()
        
        # Page 3 - Problemática
        self._add_page_3_problematica(ai_content.get("pagina_3_problematica", {}))
        
        self.doc.add_page_break()
        
        # Page 4 - Causas y Efectos
        self._add_page_4_causas_efectos(ai_content.get("pagina_4_causas_efectos", {}))
        
        self.doc.add_page_break()
        
        # Page 5 - Participantes
        self._add_page_5_participantes(ai_content.get("pagina_5_participantes", {}))
        
        self.doc.add_page_break()
        
        # Pages 6-11
        self._add_pages_6_11(ai_content)
        
        self.doc.add_page_break()
        
        # Pages 12-16
        self._add_pages_12_16(ai_content)
        
        self.doc.add_page_break()
        
        # Pages 17-21
        self._add_pages_17_21(ai_content)
        
        self.doc.add_page_break()
        
        # Pages 22+ (Dynamic Indicators)
        self._add_pages_indicadores(ai_content)
        
        self.doc.add_page_break()
        
        # Pages 23+ (Dynamic Regionalization)
        self._add_pages_regionalizacion(ai_content)
        
        self.doc.add_page_break()
        
        # Page 24 - Focalización
        self._add_page_focalizacion(ai_content)
        
        # Save and return
        return self._save_document(data)
    
    def _load_template(self, letterhead_file):
        """Load template from uploaded file"""
        try:
            if hasattr(letterhead_file, 'read'):
                return Document(BytesIO(letterhead_file.read()))
            else:
                return Document(letterhead_file)
        except Exception:
            return Document()
    
    def _apply_styles(self):
        """Apply document styles and add page numbers in footer"""
        for section in self.doc.sections:
            if not self._has_letterhead:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
                section.top_margin = Cm(2.5)  # More space for header
                section.bottom_margin = Cm(2)
            
            # Add page numbers in footer (bottom left) - "Página X"
            footer = section.footer
            footer.is_linked_to_previous = False
            
            # Clear existing footer paragraphs
            for p in footer.paragraphs:
                p.clear()
            
            # Add page number paragraph
            if footer.paragraphs:
                p = footer.paragraphs[0]
            else:
                p = footer.add_paragraph()
            
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add "Página " text
            run_text = p.add_run("Página ")
            run_text.font.size = Pt(9)
            run_text.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add page number field with proper structure
            run = p.add_run()
            fld_char_begin = OxmlElement('w:fldChar')
            fld_char_begin.set(qn('w:fldCharType'), 'begin')
            run._r.append(fld_char_begin)
            
            run2 = p.add_run()
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = ' PAGE '
            run2._r.append(instr)
            
            run3 = p.add_run()
            fld_char_separate = OxmlElement('w:fldChar')
            fld_char_separate.set(qn('w:fldCharType'), 'separate')
            run3._r.append(fld_char_separate)
            
            run4 = p.add_run("1")  # Placeholder that will be replaced
            run4.font.size = Pt(9)
            run4.font.color.rgb = RGBColor(128, 128, 128)
            
            run5 = p.add_run()
            fld_char_end = OxmlElement('w:fldChar')
            fld_char_end.set(qn('w:fldCharType'), 'end')
            run5._r.append(fld_char_end)
    
    def _add_document_letterhead(self, nombre_proyecto: str = ""):
        """
        Add configurable document letterhead/header.
        If user has provided a letterhead template, this is skipped.
        Otherwise, creates a basic header with project title and date.
        """
        if self._has_letterhead:
            # User provided letterhead template - skip adding our own
            return
        
        # Create header table: [Logo placeholder | Project Title | Date]
        header_table = self.doc.add_table(rows=1, cols=3)
        header_table.autofit = True
        
        # Left cell - Logo placeholder (Departamento Nacional de Planeación)
        cell_logo = header_table.rows[0].cells[0]
        p_logo = cell_logo.paragraphs[0]
        run_logo = p_logo.add_run("Departamento\nNacional de Planeación")
        run_logo.font.size = Pt(8)
        run_logo.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
        run_logo.bold = True
        
        # Center cell - Project title (Blue text, no background)
        cell_title = header_table.rows[0].cells[1]
        p_title = cell_title.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if nombre_proyecto:
            run_title = p_title.add_run(nombre_proyecto)
            run_title.font.size = Pt(8)
            run_title.font.color.rgb = RGBColor(0, 112, 192)  # Client Blue
            # run_title.bold = True
        
        # Right cell - Date
        cell_date = header_table.rows[0].cells[2]
        p_date = cell_date.paragraphs[0]
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        now = datetime.now()
        # Format: Impreso el DD/MM/YYYY HH:MM:SS p.m.
        date_str = f"Impreso el {now.strftime('%d/%m/%Y %I:%M:%S %p').lower()}"
        run_date = p_date.add_run(date_str)
        run_date.font.size = Pt(7)
        run_date.font.color.rgb = RGBColor(100, 100, 100)
        
        # Add spacing after header
        self.doc.add_paragraph()
    
    def _add_header(self, section_name: str, subsection: str = ""):
        """Add standard MGA section header (right-aligned section indicator)"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if subsection:
            run = p.add_run(f"{section_name} / ")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 128, 128)  # Teal
            run2 = p.add_run(subsection)
            run2.font.size = Pt(10)
            run2.font.color.rgb = RGBColor(0, 128, 128)
            run2.bold = True
        else:
            run = p.add_run(section_name)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 128, 128)
            run.bold = True
    
    def _add_section_title(self, title: str, color_hex: str = "0099CC"):
        """Add section title in blue"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 153, 204)  # Blue
        run.bold = True
        p.paragraph_format.space_after = Pt(12)
    
    def _add_subsection_title(self, title: str):
        """Add subsection title"""
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 128, 128)  # Teal
        run.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    
    
    def _add_field(self, label: str, value: str):
        """Add a label-value field (label on top, value in gray box)"""
        # Label
        p = self.doc.add_paragraph()
        run_label = p.add_run(label)
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = RGBColor(112, 173, 71)  # Client Green (roughly)
        run_label.bold = True
        p.paragraph_format.space_after = Pt(2)
        
        # Value in gray box (Table 1x1)
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.rows[0].cells[0]
        self._set_cell_shading(cell, "F2F2F2")  # Light Gray
        
        p_val = cell.paragraphs[0]
        p_val.paragraph_format.space_before = Pt(0)
        p_val.paragraph_format.space_after = Pt(0)
        p_val.paragraph_format.line_spacing = 1.0
        # Reduce cell margins
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ['top', 'bottom']:
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), '0') # Slimmer box (0 margins)
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)
        
        run_val = p_val.add_run(value)
        run_val.font.size = Pt(10)
        
        # Add spacing after field
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    def _add_inline_field(self, label: str, value: str):
        """Add a label-value field inline (Label | Value in gray box)"""
        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = True
        
        # Cell 1: Label
        cell_label = table.rows[0].cells[0]
        p_label = cell_label.paragraphs[0]
        run_label = p_label.add_run(label)
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = RGBColor(112, 173, 71)  # Client Green
        run_label.bold = True
        
        # Cell 2: Value in Gray Box
        cell_val = table.rows[0].cells[1]
        self._set_cell_shading(cell_val, "F2F2F2")  # Light Gray
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_before = Pt(0)
        p_val.paragraph_format.space_after = Pt(0)
        p_val.paragraph_format.line_spacing = 1.0
        
        # Reduce margins
        tcPr = cell_val._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ['top', 'bottom']:
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), '0') # Slimmer box (0 margins)
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

        run_val = p_val.add_run(value)
        run_val.font.size = Pt(10)
        
        # Adjust column widths if possible, or just let autofit handle it
        # Docx tables are tricky with exact widths, autofit usually works for this layout
        
        # Add spacing
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    def _add_page_1_datos_basicos(self, content: dict):
        """Add Page 1 - Datos Básicos with proper table structure matching client template"""
        # Add letterhead/header first (if not using custom template)
        nombre_proyecto = content.get("nombre_proyecto", "") or content.get("nombre", "")
        self._add_document_letterhead(nombre_proyecto)
        
        # Section indicator (right-aligned)
        self._add_header("Datos básicos")
        
        # Main section title (centered, underlined)
        self._add_section_title("Datos básicos")
        self._add_subsection_title("01 - Datos básicos del proyecto")
        
        # Nombre field
        self._add_field("Nombre", content.get("nombre", ""))
        
        # Add some spacing
        self.doc.add_paragraph()
        
        # TWO-COLUMN TABLE: Tipología | Código BPIN
        table_tipo_bpin = self.doc.add_table(rows=2, cols=2)
        table_tipo_bpin.autofit = True
        
        # Row 0: Labels (Green)
        cell_tipo_label = table_tipo_bpin.rows[0].cells[0]
        cell_bpin_label = table_tipo_bpin.rows[0].cells[1]
        
        run_tipo = cell_tipo_label.paragraphs[0].add_run("Tipología")
        run_tipo.font.size = Pt(10)
        run_tipo.font.color.rgb = RGBColor(112, 173, 71)  # Client Green
        run_tipo.bold = True
        
        run_bpin = cell_bpin_label.paragraphs[0].add_run("Código BPIN")
        run_bpin.font.size = Pt(10)
        run_bpin.font.color.rgb = RGBColor(112, 173, 71)  # Client Green
        run_bpin.bold = True
        
        # Row 1: Values (Gray Box)
        cell_tipo_val = table_tipo_bpin.rows[1].cells[0]
        cell_bpin_val = table_tipo_bpin.rows[1].cells[1]
        
        self._set_cell_shading(cell_tipo_val, "F2F2F2")
        self._set_cell_shading(cell_bpin_val, "F2F2F2")
        
        # Make slimmer
        for cell in [cell_tipo_val, cell_bpin_val]:
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'bottom']:
                node = OxmlElement(f'w:{side}')
                node.set(qn('w:w'), '0')
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)
        
        tipo_text = cell_tipo_val.paragraphs[0].add_run(content.get("tipologia", "A - PIIP - Bienes y Servicios"))
        tipo_text.font.size = Pt(10)
        
        bpin_text = cell_bpin_val.paragraphs[0].add_run(content.get("codigo_bpin", ""))
        bpin_text.font.size = Pt(10)
        
        # Add some spacing
        self.doc.add_paragraph()
        
        # Sector field
        self._add_field("Sector", content.get("sector", ""))
        
        # Add some spacing
        self.doc.add_paragraph()
        
        # TWO-COLUMN: Es Proyecto Tipo | Fecha creación works differently in client screenshot
        # Client screenshot: "Es Proyecto Tipo: [No]" (Gray box) ... "Fecha creación: [Date]" (Gray box)
        # Let's use a 1-row table with 4 cells: Label | Value | Label | Value
        
        table_tipo_fecha = self.doc.add_table(rows=1, cols=4)
        table_tipo_fecha.autofit = True
        
        # Cell 0: Label "Es Proyecto Tipo"
        run_es = table_tipo_fecha.rows[0].cells[0].paragraphs[0].add_run("Es Proyecto Tipo:")
        run_es.font.size = Pt(10)
        run_es.font.color.rgb = RGBColor(112, 173, 71)
        run_es.bold = True
        
        # Cell 1: Value "No" (Gray Box)
        cell_es_val = table_tipo_fecha.rows[0].cells[1]
        self._set_cell_shading(cell_es_val, "F2F2F2")
        cell_es_val.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell_es_val.paragraphs[0].paragraph_format.space_after = Pt(0)
        run_es_val = cell_es_val.paragraphs[0].add_run(content.get("es_proyecto_tipo", "No"))
        run_es_val.font.size = Pt(10)
        
        # Cell 2: Label "Fecha creación"
        run_fe = table_tipo_fecha.rows[0].cells[2].paragraphs[0].add_run("Fecha creación:")
        run_fe.font.size = Pt(10)
        run_fe.font.color.rgb = RGBColor(112, 173, 71)
        run_fe.bold = True
        table_tipo_fecha.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Cell 3: Value Date (Gray Box)
        cell_fe_val = table_tipo_fecha.rows[0].cells[3]
        self._set_cell_shading(cell_fe_val, "F2F2F2")
        cell_fe_val.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell_fe_val.paragraphs[0].paragraph_format.space_after = Pt(0)
        run_fe_val = cell_fe_val.paragraphs[0].add_run(content.get("fecha_creacion", datetime.now().strftime("%d/%m/%Y %H:%M:%S")))
        run_fe_val.font.size = Pt(10)
        
        # Add some spacing
        self.doc.add_paragraph()
        
        # Identificador - inline layout
        self._add_inline_field("Identificador:", content.get("identificador", ""))
        
        # Formulador Ciudadano - inline layout
        self._add_inline_field("Formulador Ciudadano:", content.get("formulador_ciudadano", ""))
        
        # Formulador Oficial - inline layout
        self._add_inline_field("Formulador Oficial:", content.get("formulador_oficial", ""))
    
    def _add_horizontal_line(self):
        """Add a thin horizontal separator line"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # Create a simple border line using a paragraph border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    def _add_page_2_plan_desarrollo(self, content: dict):
        """Add Page 2 - Plan de Desarrollo"""
        self._add_header("Identificación", "Plan de desarrollo")
        
        self._add_section_title("Contribución a la política pública")
        
        # Plan Nacional
        self._add_subsection_title("01 - Contribución al Plan Nacional de Desarrollo")
        
        plan_nacional = content.get("plan_nacional", {})
        self._add_field("Plan", plan_nacional.get("nombre", ""))
        self._add_field("Programa", plan_nacional.get("programa", ""))
        
        # Table
        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        
        headers = ["Transformación", "Pilar", "Catalizador", "Componente"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            self._set_cell_shading(cell, "0099CC")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8)
        
        table.rows[1].cells[0].text = plan_nacional.get("transformacion", "")
        table.rows[1].cells[1].text = plan_nacional.get("pilar", "")
        table.rows[1].cells[2].text = plan_nacional.get("catalizador", "")
        table.rows[1].cells[3].text = plan_nacional.get("componente", "")
        
        self.doc.add_paragraph()
        
        # Plan Departamental
        self._add_subsection_title("02 - Plan de Desarrollo Departamental o Sectorial")
        plan_depto = content.get("plan_departamental", {})
        self._add_field("Plan de Desarrollo Departamental o Sectorial", plan_depto.get("nombre", ""))
        self._add_field("Estrategia del Plan de Desarrollo Departamental o Sectorial", plan_depto.get("estrategia", ""))
        self._add_field("Programa del Plan Desarrollo Departamental o Sectorial", plan_depto.get("programa", ""))
        
        # Plan Municipal
        self._add_subsection_title("03 - Plan de Desarrollo Distrital o Municipal")
        plan_mun = content.get("plan_municipal", {})
        self._add_field("Plan de Desarrollo Distrital o Municipal", plan_mun.get("nombre", ""))
        self._add_field("Estrategia del Plan de Desarrollo Distrital o Municipal", plan_mun.get("estrategia", ""))
        self._add_field("Programa del Plan desarrollo Distrital o Municipal", plan_mun.get("programa", ""))
        
        # Grupos étnicos
        self._add_subsection_title("04 - Instrumentos de planeación de grupos étnicos")
        self._add_field("Tipo de entidad", content.get("instrumentos_grupos_etnicos", "No aplica"))
    
    def _add_page_3_problematica(self, content: dict):
        """Add Page 3 - Problemática"""
        self._add_header("Identificación", "Problemática")
        
        self._add_section_title("Identificación y descripción del problema")
        
        self._add_subsection_title("Problema central")
        p = self.doc.add_paragraph()
        p.add_run(self._safe_str(content.get("problema_central", "")))
        
        self._add_subsection_title("Descripción de la situación existente con respecto al problema")
        desc = self._safe_str(content.get("descripcion_situacion", "")).replace("\\n", "\n")
        for para in desc.split("\n\n"):
            p = self.doc.add_paragraph()
            p.add_run(para)
            p.paragraph_format.space_after = Pt(6)
        
        self._add_subsection_title("Magnitud actual del problema – indicadores de referencia")
        mag = self._safe_str(content.get("magnitud_problema", "")).replace("\\n", "\n")
        for para in mag.split("\n\n"):
            p = self.doc.add_paragraph()
            p.add_run(para)
            p.paragraph_format.space_after = Pt(6)
    
    def _add_page_4_causas_efectos(self, content: dict):
        """Add Page 4 - Causas y Efectos"""
        self._add_header("Identificación", "Problemática")
        
        # Causas
        self._add_subsection_title("01 - Causas que generan el problema")
        
        causas_table = self.doc.add_table(rows=1, cols=2)
        causas_table.style = 'Table Grid'
        
        causas_table.rows[0].cells[0].text = "Causas directas"
        causas_table.rows[0].cells[1].text = "Causas indirectas"
        self._set_cell_shading(causas_table.rows[0].cells[0], "0099CC")
        self._set_cell_shading(causas_table.rows[0].cells[1], "0099CC")
        for cell in causas_table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        causas_directas = self._safe_list(content.get("causas_directas", []))
        causas_indirectas = self._safe_list(content.get("causas_indirectas", []))
        
        max_rows = max(len(causas_directas), len(causas_indirectas))
        for i in range(max_rows):
            row = causas_table.add_row()
            if i < len(causas_directas):
                c = causas_directas[i]
                row.cells[0].text = f"{c.get('numero', '')}. {c.get('causa', '')}"
            if i < len(causas_indirectas):
                c = causas_indirectas[i]
                row.cells[1].text = f"{c.get('numero', '')} {c.get('causa', '')}"
        
        self.doc.add_paragraph()
        
        # Efectos
        self._add_subsection_title("02 - Efectos generados por el problema")
        
        efectos_table = self.doc.add_table(rows=1, cols=2)
        efectos_table.style = 'Table Grid'
        
        efectos_table.rows[0].cells[0].text = "Efectos directos"
        efectos_table.rows[0].cells[1].text = "Efectos indirectos"
        self._set_cell_shading(efectos_table.rows[0].cells[0], "0099CC")
        self._set_cell_shading(efectos_table.rows[0].cells[1], "0099CC")
        for cell in efectos_table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        efectos_directos = self._safe_list(content.get("efectos_directos", []))
        efectos_indirectos = self._safe_list(content.get("efectos_indirectos", []))
        
        max_rows = max(len(efectos_directos), len(efectos_indirectos))
        for i in range(max_rows):
            row = efectos_table.add_row()
            if i < len(efectos_directos):
                e = efectos_directos[i]
                row.cells[0].text = f"{e.get('numero', '')}. {e.get('efecto', '')}"
            if i < len(efectos_indirectos):
                e = efectos_indirectos[i]
                row.cells[1].text = f"{e.get('numero', '')} {e.get('efecto', '')}"
    
    def _add_page_5_participantes(self, content: dict):
        """Add Page 5 - Participantes"""
        self._add_header("Identificación", "Participantes")
        
        self._add_section_title("Identificación y análisis de participantes")
        
        self._add_subsection_title("01 - Identificación de los participantes")
        
        # Participants table
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        table.rows[0].cells[0].text = "Participante"
        table.rows[0].cells[1].text = "Contribución o Gestión"
        self._set_cell_shading(table.rows[0].cells[0], "0099CC")
        self._set_cell_shading(table.rows[0].cells[1], "0099CC")
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        for part in self._safe_list(content.get("participantes", [])):
            row = table.add_row()
            
            # Left cell - participant details
            cell_text = f"Actor: {part.get('actor', '')}\n"
            cell_text += f"Entidad: {part.get('entidad', '')}\n"
            cell_text += f"Posición: {part.get('posicion', '')}\n"
            cell_text += f"Intereses o Expectativas: {part.get('intereses', '')}"
            row.cells[0].text = cell_text
            
            # Right cell - contribution
            row.cells[1].text = part.get("contribucion", "")
        
        self.doc.add_paragraph()
        
        # Análisis
        self._add_subsection_title("02 - Análisis de los participantes")
        analisis = self._safe_str(content.get("analisis_participantes", "")).replace("\\n", "\n")
        for para in analisis.split("\n\n"):
            p = self.doc.add_paragraph()
            p.add_run(para)
            p.paragraph_format.space_after = Pt(6)
    
    def _add_page_6_poblacion(self, content: dict):
        """Add Page 6 - Población"""
        self._add_header("Identificación", "Población")
        
        # Población afectada
        self._add_subsection_title("01 - Población afectada por el problema")
        
        pob_afectada = content.get("poblacion_afectada", {})
        self._add_field("Tipo de población", pob_afectada.get("tipo", "Personas"))
        self._add_field("Número", str(pob_afectada.get("numero", "")))
        self._add_field("Fuente de la información", pob_afectada.get("fuente", ""))
        
        # Extract location data (handle both nested and flat structure)
        # Fallback to self._data (form input) if AI doesn't provide location
        loc_afectada = pob_afectada.get("localizacion", {})
        region_afectada = loc_afectada.get("region", "") or pob_afectada.get("region", "") or "Caribe"
        dept_afectada = loc_afectada.get("departamento", "") or pob_afectada.get("departamento", "") or self._data.get("departamento", "")
        mun_afectada = loc_afectada.get("municipio", "") or pob_afectada.get("municipio", "") or self._data.get("municipio", "")
        tipo_agrup_afectada = loc_afectada.get("tipo_agrupacion", "") or pob_afectada.get("tipo_agrupacion", "") or "Urbana"
        agrup_afectada = loc_afectada.get("agrupacion", "") or pob_afectada.get("agrupacion", "") or self._data.get("municipio", "")
        
        # Localización table
        self._add_subsection_title("Localización")
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        table.rows[0].cells[0].text = "Ubicación general"
        table.rows[0].cells[1].text = "Localización específica/Otro tipo de entidad étnica"
        self._set_cell_shading(table.rows[0].cells[0], "0099CC")
        self._set_cell_shading(table.rows[0].cells[1], "0099CC")
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        table.rows[1].cells[0].text = f"Región: {region_afectada}"
        table.rows[2].cells[0].text = f"Departamento: {dept_afectada}"
        table.rows[3].cells[0].text = f"Municipio: {mun_afectada}"
        table.rows[4].cells[0].text = f"Tipo de Agrupación: {tipo_agrup_afectada}\nAgrupación: {agrup_afectada}"
        
        self.doc.add_paragraph()
        
        # Población objetivo
        self._add_subsection_title("02 - Población objetivo de la intervención")
        
        pob_objetivo = content.get("poblacion_objetivo", {})
        self._add_field("Tipo de población", pob_objetivo.get("tipo", "Personas"))
        self._add_field("Número", str(pob_objetivo.get("numero", "")))
        self._add_field("Fuente de la información", pob_objetivo.get("fuente", ""))
        
        # Extract location data (handle both nested and flat structure)
        # Fallback to self._data (form input) if AI doesn't provide location
        loc_objetivo = pob_objetivo.get("localizacion", {})
        region_objetivo = loc_objetivo.get("region", "") or pob_objetivo.get("region", "") or "Caribe"
        dept_objetivo = loc_objetivo.get("departamento", "") or pob_objetivo.get("departamento", "") or self._data.get("departamento", "")
        mun_objetivo = loc_objetivo.get("municipio", "") or pob_objetivo.get("municipio", "") or self._data.get("municipio", "")
        tipo_agrup_objetivo = loc_objetivo.get("tipo_agrupacion", "") or pob_objetivo.get("tipo_agrupacion", "") or "Urbana"
        agrup_objetivo = loc_objetivo.get("agrupacion", "") or pob_objetivo.get("agrupacion", "") or self._data.get("municipio", "")
        
        # Localización table for objetivo
        self._add_subsection_title("Localización")
        table2 = self.doc.add_table(rows=5, cols=3)
        table2.style = 'Table Grid'
        
        table2.rows[0].cells[0].text = "Ubicación general"
        table2.rows[0].cells[1].text = "Localización específica/Otro tipo de entidad étnica"
        table2.rows[0].cells[2].text = "Nombre del consejo comunitario"
        for i in range(3):
            self._set_cell_shading(table2.rows[0].cells[i], "0099CC")
            for run in table2.rows[0].cells[i].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        table2.rows[1].cells[0].text = f"Región: {region_objetivo}"
        table2.rows[2].cells[0].text = f"Departamento: {dept_objetivo}"
        table2.rows[3].cells[0].text = f"Municipio: {mun_objetivo}"
        table2.rows[4].cells[0].text = f"Tipo de Agrupación: {tipo_agrup_objetivo}\nAgrupación: {agrup_objetivo}"
    
    def _add_page_7_objetivos(self, content: dict):
        """Add Page 7 - Objetivos"""
        self._add_header("Identificación", "Objetivos")
        
        self._add_section_title("Objetivos específicos")
        
        self._add_subsection_title("01 - Objetivo general e indicadores de seguimiento")
        
        self._add_field("Problema central", content.get("problema_central", ""))
        self._add_field("Objetivo general – Propósito", content.get("objetivo_general", ""))
        
        # Indicadores table
        self._add_subsection_title("Indicadores para medir el objetivo general")
        
        indicadores = content.get("indicadores", [])
        if indicadores:
            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            table.rows[0].cells[0].text = "Indicador objetivo"
            table.rows[0].cells[1].text = "Descripción"
            table.rows[0].cells[2].text = "Fuente de verificación"
            for cell in table.rows[0].cells:
                self._set_cell_shading(cell, "0099CC")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            for ind in indicadores:
                row = table.add_row()
                row.cells[0].text = ind.get("nombre", "")
                row.cells[1].text = f"Medido a través de: {ind.get('medido', '')}\nMeta: {ind.get('meta', '')}\nTipo de fuente: {ind.get('tipo_fuente', '')}"
                row.cells[2].text = ind.get("fuente_verificacion", "")
        
        self.doc.add_paragraph()
        
        # Relación causas-objetivos
        self._add_subsection_title("02 - Relaciones entre las causas y objetivos")
        
        relaciones = content.get("relacion_causas_objetivos", [])
        if relaciones:
            table2 = self.doc.add_table(rows=1, cols=2)
            table2.style = 'Table Grid'
            
            table2.rows[0].cells[0].text = "Causa relacionada"
            table2.rows[0].cells[1].text = "Objetivos específicos"
            for cell in table2.rows[0].cells:
                self._set_cell_shading(cell, "0099CC")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            for rel in relaciones:
                row = table2.add_row()
                row.cells[0].text = rel.get("causa", "")
                row.cells[1].text = rel.get("objetivo", "")
        
        self.doc.add_paragraph()
        
        # Alternativas
        self._add_section_title("Alternativas de la solución")
        self._add_subsection_title("01 - Alternativas de la solución")
        
        alternativas = content.get("alternativas", [])
        if alternativas:
            table3 = self.doc.add_table(rows=1, cols=3)
            table3.style = 'Table Grid'
            
            table3.rows[0].cells[0].text = "Nombre de la alternativa"
            table3.rows[0].cells[1].text = "Se evaluará con esta herramienta"
            table3.rows[0].cells[2].text = "Estado"
            for cell in table3.rows[0].cells:
                self._set_cell_shading(cell, "0099CC")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            for alt in alternativas:
                row = table3.add_row()
                row.cells[0].text = alt.get("nombre", "")
                row.cells[1].text = alt.get("evaluacion", "")
                row.cells[2].text = alt.get("estado", "")
        
        # Evaluaciones
        self._add_subsection_title("Evaluaciones a realizar")
        evaluaciones = content.get("evaluaciones", {})
        
        table4 = self.doc.add_table(rows=3, cols=2)
        table4.style = 'Table Grid'
        table4.rows[0].cells[0].text = "Rentabilidad:"
        table4.rows[0].cells[1].text = evaluaciones.get("rentabilidad", "Si")
        table4.rows[1].cells[0].text = "Costo - Eficiencia y Costo mínimo:"
        table4.rows[1].cells[1].text = evaluaciones.get("costo_eficiencia", "Si")
        table4.rows[2].cells[0].text = "Evaluación multicriterio:"
        table4.rows[2].cells[1].text = evaluaciones.get("multicriterio", "No")
    
    def _add_estudio_necesidades_servicio(self, servicio_data: dict, servicio_nombre: str):
        """Add estudio de necesidades for a service"""
        self._add_header("Identificación", "Alternativas")
        
        # Use dynamic project name instead of hardcoded subsidios text
        nombre_proyecto = self._data.get("nombre_proyecto", "")
        p = self.doc.add_paragraph()
        run = p.add_run(f"Alternativa 1. {nombre_proyecto}")
        run.bold = True
        run.font.size = Pt(10)
        
        self._add_section_title("Estudio de necesidades")
        self._add_subsection_title("01 - Bien o servicio")
        
        self._add_field("Bien o servicio", servicio_data.get("bien_servicio", ""))
        self._add_field("Medido a través de", servicio_data.get("medido", "Número"))
        self._add_field("Descripción", servicio_data.get("descripcion", ""))
        self._add_field("Descripción de la Demanda", servicio_data.get("descripcion_demanda", ""))
        self._add_field("Descripción de la Oferta", servicio_data.get("descripcion_oferta", ""))
        
        # Oferta/Demanda table
        tabla_od = servicio_data.get("tabla_oferta_demanda", [])
        if tabla_od:
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.autofit = False
            
            # Set column widths for better display
            col_widths = [Inches(1.2), Inches(1.5), Inches(1.5), Inches(1.5)]
            for i, width in enumerate(col_widths):
                for cell in table.columns[i].cells:
                    cell.width = width
            
            headers = ["Año", "Oferta", "Demanda", "Déficit"]
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                cell.width = col_widths[i]
                self._set_cell_shading(cell, "0099CC")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for idx, item in enumerate(tabla_od):
                row = table.add_row()
                row.cells[0].text = str(item.get("ano", ""))
                row.cells[1].text = str(item.get("oferta", ""))
                row.cells[2].text = str(item.get("demanda", ""))
                row.cells[3].text = str(item.get("deficit", ""))
                
                # Alternating row colors (light blue for odd rows, white for even)
                if idx % 2 == 0:
                    for cell in row.cells:
                        self._set_cell_shading(cell, "E8F4F8")  # Light blue
                
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def _add_pages_6_11(self, content: dict):
        """Add pages 6-11"""
        # Page 6 - Población
        self._add_page_6_poblacion(content.get("pagina_6_poblacion", {}))
        
        self.doc.add_page_break()
        
        # Page 7 - Objetivos
        self._add_page_7_objetivos(content.get("pagina_7_objetivos", {}))
        
        self.doc.add_page_break()
        
        # Pages 8-11 - Estudio de Necesidades
        estudio = content.get("pagina_8_9_10_11_estudio_necesidades", {})
        
        # Use servicio_principal (main service) - this matches the prompt output
        servicio_principal = estudio.get("servicio_principal", {})
        if servicio_principal:
            self._add_estudio_necesidades_servicio(servicio_principal, "Principal")
    
    def _add_page_12_analisis_tecnico(self, content: dict):
        """Add Page 12 - Análisis Técnico"""
        self._add_header("Preparación", "Análisis técnico")
        
        # Use dynamic project name
        nombre_proyecto = self._data.get("nombre_proyecto", "")
        p = self.doc.add_paragraph()
        run = p.add_run(f"Alternativa: {nombre_proyecto}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        self._add_section_title("Análisis técnico de la alternativa")
        self._add_subsection_title("01 - Análisis técnico de la alternativa")
        
        # Green label for subsection
        p = self.doc.add_paragraph()
        run = p.add_run("Análisis técnico de la alternativa")
        run.font.color.rgb = RGBColor(0, 128, 0)
        run.font.size = Pt(11)
        
        # Main description
        p = self.doc.add_paragraph()
        p.add_run("La alternativa corresponde:")
        
        # Get the technical analysis content from AI
        descripcion = content.get("analisis_tecnico", "")
        if descripcion:
            p = self.doc.add_paragraph()
            p.add_run(descripcion)
        
        # Functions list (if provided by AI)
        funciones = content.get("funciones", [])
        if funciones:
            for funcion in funciones:
                p = self.doc.add_paragraph()
                p.add_run(f"- {funcion}")
    
    def _add_page_13_localizacion(self, content: dict):
        """Add Page 13 - Localización"""
        self._add_header("Preparación", "Localización")
        
        # Add Alternativa line with dynamic project name
        nombre_proyecto = self._data.get("nombre_proyecto", "")
        p = self.doc.add_paragraph()
        run = p.add_run(f"Alternativa: {nombre_proyecto}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        self._add_section_title("Localización de la alternativa")
        self._add_subsection_title("01 - Localización de la alternativa")
        
        ubicacion = content.get("ubicacion", {})
        
        # Fallback to self._data for location values
        region = ubicacion.get('region', '') or self._data.get('region', 'Caribe')
        departamento = ubicacion.get('departamento', '') or self._data.get('departamento', '')
        municipio = ubicacion.get('municipio', '') or self._data.get('municipio', '')
        
        table = self.doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        
        table.rows[0].cells[0].text = "Ubicación general"
        table.rows[0].cells[1].text = "Ubicación específica"
        for cell in table.rows[0].cells:
            self._set_cell_shading(cell, "0099CC")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Fill general location (column 0) with fallback to input data
        table.rows[1].cells[0].text = f"Región: {region}"
        table.rows[2].cells[0].text = f"Departamento: {departamento}"
        table.rows[3].cells[0].text = f"Municipio: {municipio}"
        table.rows[4].cells[0].text = f"Tipo de Agrupación:"
        table.rows[5].cells[0].text = f"Agrupación:"
        table.rows[6].cells[0].text = f"Latitud:"
        table.rows[7].cells[0].text = f"Longitud:"
        
        self.doc.add_paragraph()
        
        # Factores analizados
        self._add_subsection_title("02 - Factores analizados")
        
        # Default factors if not provided by AI
        factores = content.get("factores_analizados", [])
        if not factores:
            factores = [
                "Cercanía a la población objetivo",
                "Cercanía de fuentes de abastecimiento",
                "Disponibilidad y costo de mano de obra",
                "Factores ambientales",
                "Otros"
            ]
        
        # Display factors as simple list
        p = self.doc.add_paragraph()
        p.add_run(",\n".join(factores))
    
    def _add_page_14_cadena_valor(self, content: dict):
        """Add Page 14 - Cadena de Valor"""
        self._add_header("Preparación", "Cadena de valor")
        
        # Add Alternativa line with dynamic project name
        nombre_proyecto = self._data.get("nombre_proyecto", "")
        p = self.doc.add_paragraph()
        run = p.add_run(f"Alternativa: {nombre_proyecto}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        self._add_section_title("Cadena de valor de la alternativa")
        
        # Costo total aligned right
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        costo_total = content.get('costo_total', '') or self._data.get('valor_total', '')
        run = p.add_run(f"Costo total de la alternativa: $ {costo_total}")
        run.bold = True
        
        # Support multiple objectives OR top-level productos
        objetivos = content.get("objetivos", [])
        
        # Check if prompt generated productos at top level (not under objetivos)
        top_level_productos = content.get("productos", [])
        top_level_actividades = content.get("actividades", [])
        
        if not objetivos and (top_level_productos or content.get("objetivo_general")):
            # Create synthetic objetivo from top-level data
            objetivos = [{
                "numero": "1",
                "descripcion": content.get("objetivo_general", nombre_proyecto),
                "costo": costo_total,
                "productos": top_level_productos
            }]
        elif not objetivos:
            # Fallback to single objetivo_especifico format
            obj = content.get("objetivo_especifico", {})
            if obj:
                objetivos = [obj]
        
        for obj in objetivos:
            # Objetivo header
            num = obj.get('numero', '1')
            costo_obj = obj.get('costo', costo_total)
            self._add_subsection_title(f"{num} - Objetivo específico {num}   Costo: $ {costo_obj}")
            
            p = self.doc.add_paragraph()
            p.add_run(obj.get("descripcion", ""))
            
            self.doc.add_paragraph()
            
            # Products for this objective (try obj.productos, then top_level, then fallback)
            productos = obj.get("productos", [])
            if not productos:
                productos = top_level_productos
            if not productos:
                # Fallback to single producto from content
                single_prod = content.get("producto", {})
                if single_prod:
                    productos = [single_prod]
            
            # Create ONE table for all products in this objective
            if productos:
                table = self.doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Header row
                table.rows[0].cells[0].text = "Producto"
                table.rows[0].cells[1].text = "Actividad y/o Entregable"
                for cell in table.rows[0].cells:
                    self._set_cell_shading(cell, "0099CC")
                    for run in cell.paragraphs[0].runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
                
                for idx, producto in enumerate(productos):
                    row = table.add_row()
                    
                    # Add light blue shading to product cells for visual separation
                    self._set_cell_shading(row.cells[0], "E8F4F8")
                    self._set_cell_shading(row.cells[1], "FFFFFF")
                    
                    # Left cell - producto with all details (matching client format)
                    codigo = producto.get('codigo', f'1.{idx+1}')
                    nombre = producto.get('nombre', '')
                    complemento = producto.get('complemento', '') or producto.get('complemento_info', '')
                    
                    prod_text = f"{codigo} {nombre}\n"
                    if complemento:
                        prod_text += f"    {complemento}  (Producto principal del proyecto)\n\n"
                    else:
                        prod_text += "\n"
                    prod_text += f"Complemento:\n"
                    prod_text += f"Medido a través de: {producto.get('medido', '') or producto.get('unidad', '')}\n"
                    prod_text += f"Cantidad: {producto.get('cantidad', '')}\n"
                    prod_text += f"Costo: $ {producto.get('costo', '')}\n"
                    prod_text += f"Etapa: {producto.get('etapa', 'Inversión')}\n"
                    prod_text += f"Localización:\n"
                    prod_text += f"Número de Personas: {producto.get('personas', '') or producto.get('num_personas', '')}\n"
                    prod_text += f"Acumulativo o no: {producto.get('acumulativo', 'No Acumulativo')}\n"
                    prod_text += f"Población Beneficiaria: {producto.get('poblacion_beneficiaria', '')}"
                    row.cells[0].text = prod_text
                    
                    # Right cell - actividades for this product
                    prod_actividades = producto.get("actividades", [])
                    if not prod_actividades:
                        prod_actividades = top_level_actividades or content.get("actividades", [])
                    
                    act_text = ""
                    for act in prod_actividades:
                        act_codigo = act.get('codigo', '')
                        act_nombre = act.get('nombre', '') or act.get('descripcion', '')
                        act_text += f"{act_codigo} {act_nombre}\n\n"
                        act_text += f"Costo: $ {act.get('costo', '')}\n"
                        act_text += f"Etapa: {act.get('etapa', 'Inversión')}\n\n"
                    row.cells[1].text = act_text
                
                self.doc.add_paragraph()
    
    def _add_page_15_actividades_detalle(self, content: dict):
        """Add Page 15 - Actividades Detalle"""
        self._add_header("Preparación", "Cadena de valor")
        
        actividades = content.get("actividades_periodo", [])
        
        for act in actividades:
            p = self.doc.add_paragraph()
            # Use dynamic activity description instead of hardcoded subsidios text
            act_desc = act.get('descripcion', act.get('nombre', ''))
            run = p.add_run(f"Actividad {act.get('codigo', '')} {act_desc}")
            run.bold = True
            run.font.size = Pt(10)
            
            # Periodo table
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            table.rows[0].cells[0].text = "Periodo"
            table.rows[0].cells[1].text = "Servicios domiciliarios"
            for cell in table.rows[0].cells:
                self._set_cell_shading(cell, "0099CC")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            for periodo in self._safe_list(act.get("periodos", [])):
                row = table.add_row()
                row.cells[0].text = str(periodo.get("periodo", "1"))
                row.cells[1].text = periodo.get("valor", "")
            
            row_total = table.add_row()
            row_total.cells[0].text = "Total"
            self._set_cell_shading(row_total.cells[0], "7FC8D8")
            row_total.cells[1].text = act.get("total", "")
            
            self.doc.add_paragraph()
            
            # Total table
            table2 = self.doc.add_table(rows=2, cols=2)
            table2.style = 'Table Grid'
            table2.rows[0].cells[0].text = "Periodo"
            table2.rows[0].cells[1].text = "Total"
            for cell in table2.rows[0].cells:
                self._set_cell_shading(cell, "0099CC")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            table2.rows[1].cells[0].text = "1"
            table2.rows[1].cells[1].text = act.get("total", "")
            
            self.doc.add_paragraph()
    
    def _add_page_16_riesgos(self, content: dict):
        """Add Page 16 - Riesgos with 6-column structure including hierarchy level"""
        self._add_header("Preparación", "Riesgos")
        
        self._add_section_title("Análisis de riesgos alternativa")
        self._add_subsection_title("01 - Análisis de riesgo")
        
        # Riesgos table - 6 columns with level indicator
        table = self.doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Set column widths (in inches) - total ~7 inches for letter size with margins
        # Level column narrower, others adjusted for content
        col_widths = [Inches(0.7), Inches(0.9), Inches(1.5), Inches(1.0), Inches(1.2), Inches(1.7)]
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = width
        
        # Header row
        headers = ["", "Tipo de riesgo", "Descripción del riesgo", "Probabilidad e impacto", "Efectos", "Medidas de mitigación"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.width = col_widths[i]
            self._set_cell_shading(cell, "0099CC")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8)
                run.font.bold = True
        
        riesgos = content.get("riesgos", [])
        current_level = None
        level_start_row = None
        
        for idx, riesgo in enumerate(riesgos):
            row = table.add_row()
            
            # Get level for this risk
            nivel = riesgo.get('nivel', '')
            
            # Add alternating row shading for visual separation
            if idx % 2 == 0:
                for i, cell in enumerate(row.cells):
                    if i > 0:  # Don't shade level column here
                        self._set_cell_shading(cell, "E8F4F8")  # Light blue
            
            # Level column (first column) - teal/green shading
            row.cells[0].text = nivel
            self._set_cell_shading(row.cells[0], "7FC8A8")  # Teal/green color
            for para in row.cells[0].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(7)
                    run.font.bold = True
            
            # Type column
            row.cells[1].text = riesgo.get('tipo', 'Administrativos')
            
            # Description column
            desc = riesgo.get('descripcion', '')
            row.cells[2].text = desc.replace('\\n', '\n')
            
            # Probability and impact column
            prob = riesgo.get('probabilidad', '')
            imp = riesgo.get('impacto', '')
            row.cells[3].text = f"Probabilidad:\n{prob}\n\nImpacto: {imp}"
            
            # Effects column
            row.cells[4].text = riesgo.get('efectos', '')
            
            # Mitigation column
            row.cells[5].text = riesgo.get('mitigacion', '')
            
            # Apply consistent font size to all cells
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
    
    def _add_pages_12_16(self, content: dict):
        """Add pages 12-16"""
        # Page 12 - Análisis Técnico (includes localization info)
        self._add_page_12_analisis_tecnico(content.get("pagina_12_analisis_tecnico", {}))
        
        self.doc.add_page_break()
        
        # Page 13/14 - Cadena de Valor (this is the main products/activities table)
        # Prompt generates this as pagina_13_cadena_valor
        cadena_data = content.get("pagina_13_cadena_valor") or content.get("pagina_14_cadena_valor", {})
        self._add_page_14_cadena_valor(cadena_data)
        
        self.doc.add_page_break()
        
        # Page 15/16 - Riesgos (comes AFTER Cadena de Valor)
        # Prompt generates this as pagina_14_riesgos
        self._add_page_16_riesgos(content.get("pagina_14_riesgos") or content.get("pagina_16_riesgos", {}))

    
    def _add_page_17_riesgos_continuacion(self, content: dict):
        """Add Page 17 - Riesgos Continuation"""
        self._add_header("Preparación", "Riesgos")
        
        riesgos = content.get("riesgos_adicionales", [])
        
        for riesgo in riesgos:
            # Risk description
            p = self.doc.add_paragraph()
            p.add_run(riesgo.get("descripcion_actividad", ""))
            p = self.doc.add_paragraph()
            p.add_run(riesgo.get("descripcion_riesgo", ""))
            
            # Add risk table with 4 columns
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            headers = ["Tipo", "Probabilidad/Impacto", "Efectos", "Mitigación"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                self._set_cell_shading(table.rows[0].cells[i], "0099CC")
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(8)
            
            # Data row
            row = table.add_row()
            row.cells[0].text = riesgo.get("tipo", "")
            row.cells[1].text = f"Probabilidad: {riesgo.get('probabilidad', '')}\n\nImpacto: {riesgo.get('impacto', '')}"
            row.cells[2].text = riesgo.get("efectos", "")
            row.cells[3].text = riesgo.get("mitigacion", "")
            
            self.doc.add_paragraph()
    
    def _add_page_18_19_ingresos_beneficios(self, content: dict):
        """Add Pages 18-19 - Ingresos y Beneficios with multiple benefit items"""
        self._add_header("Preparación", "Ingresos y beneficios")
        
        self._add_section_title("Ingresos y beneficios alternativa")
        self._add_subsection_title("01 - Ingresos y beneficios")
        
        # Handle multiple benefit items
        beneficios = content.get("beneficios", [])
        
        # Fallback to old structure if no beneficios list
        if not beneficios:
            beneficios = [{
                "titulo": content.get("descripcion", ""),
                "tipo": content.get("tipo", "Beneficios"),
                "medido": content.get("medido", "Número"),
                "bien_producido": content.get("bien_producido", "Otros"),
                "razon_precio_cuenta": content.get("razon_precio_cuenta", "1.0"),
                "descripcion_cantidad": content.get("descripcion_cantidad", ""),
                "descripcion_valor_unitario": content.get("descripcion_valor_unitario", ""),
                "tabla_periodos": content.get("tabla_periodos", [])
            }]
        
        for beneficio in beneficios:
            # Benefit title
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            p.add_run(beneficio.get("titulo", ""))
            
            # Fields without labels in gray boxes (inline style)
            p = self.doc.add_paragraph()
            p.add_run(f"Tipo: ").bold = True
            p.add_run(beneficio.get("tipo", "Beneficios"))
            
            p = self.doc.add_paragraph()
            p.add_run(f"Medido a través de: ").bold = True
            p.add_run(beneficio.get("medido", "Número"))
            
            p = self.doc.add_paragraph()
            p.add_run(f"Bien producido: ").bold = True
            p.add_run(beneficio.get("bien_producido", ""))
            
            p = self.doc.add_paragraph()
            p.add_run(f"Razón Precio Cuenta (RPC): ").bold = True
            p.add_run(beneficio.get("razon_precio_cuenta", ""))
            
            p = self.doc.add_paragraph()
            p.add_run(f"Descripción Cantidad: ").bold = True
            p.add_run(beneficio.get("descripcion_cantidad", ""))
            
            p = self.doc.add_paragraph()
            p.add_run(f"Descripción Valor Unitario: ").bold = True
            p.add_run(beneficio.get("descripcion_valor_unitario", ""))
            
            # Periodos table for this benefit
            tabla_periodos = beneficio.get("tabla_periodos", [])
            if tabla_periodos:
                table = self.doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                headers = ["Periodo", "Cantidad", "Valor unitario", "Valor total"]
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                    self._set_cell_shading(table.rows[0].cells[i], "0099CC")
                    for run in table.rows[0].cells[i].paragraphs[0].runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(9)
                
                for item in tabla_periodos:
                    row = table.add_row()
                    row.cells[0].text = str(item.get("periodo", ""))
                    row.cells[1].text = str(item.get("cantidad", ""))
                    row.cells[2].text = str(item.get("valor_unitario", ""))
                    row.cells[3].text = str(item.get("valor_total", ""))
            
            self.doc.add_paragraph()
        
        # Totales section
        self._add_subsection_title("02 - Totales")
        
        tabla_totales = content.get("tabla_totales", [])
        if tabla_totales:
            table2 = self.doc.add_table(rows=1, cols=3)
            table2.style = 'Table Grid'
            
            headers = ["Periodo", "Total beneficios", "Total"]
            for i, h in enumerate(headers):
                table2.rows[0].cells[i].text = h
                self._set_cell_shading(table2.rows[0].cells[i], "0099CC")
                for run in table2.rows[0].cells[i].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            for item in tabla_totales:
                row = table2.add_row()
                row.cells[0].text = str(item.get("periodo", ""))
                row.cells[1].text = str(item.get("total_beneficios", ""))
                row.cells[2].text = str(item.get("total", ""))
    
    def _add_page_20_flujo_economico(self, content: dict):
        """Add Page 20 - Flujo Económico with 10-column table"""
        self._add_header("Evaluación", "Flujo Económico")
        
        p = self.doc.add_paragraph()
        run = p.add_run(content.get("alternativa", "Alternativa 1"))
        run.bold = True
        
        self._add_section_title("Flujo")
        self._add_subsection_title("01 - Flujo Económico")
        
        flujo = content.get("flujo", [])
        if flujo:
            table = self.doc.add_table(rows=1, cols=10)
            table.style = 'Table Grid'
            
            headers = ["P", "Beneficios e ingresos (+)", "Créditos(+)", "Costos de preinversión (-)", 
                      "Costos de inversión (-)", "Costos de operación (-)", "Amortización (-)", 
                      "Intereses de los créditos (-)", "Valor de salvamento (+)", "Flujo Neto"]
            
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                self._set_cell_shading(table.rows[0].cells[i], "0099CC")
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(7)
            
            for item in flujo:
                row = table.add_row()
                row.cells[0].text = str(item.get("p", ""))
                row.cells[1].text = str(item.get("beneficios", ""))
                row.cells[2].text = str(item.get("creditos", "$0,0"))
                row.cells[3].text = str(item.get("costos_preinversion", "$0,0"))
                row.cells[4].text = str(item.get("costos_inversion", ""))
                row.cells[5].text = str(item.get("costos_operacion", "$0,0"))
                row.cells[6].text = str(item.get("amortizacion", "$0,0"))
                row.cells[7].text = str(item.get("intereses", "$0,0"))
                row.cells[8].text = str(item.get("valor_salvamento", "$0,0"))
                row.cells[9].text = str(item.get("flujo_neto", ""))
    
    def _add_page_21_indicadores_decision(self, content: dict):
        """Add Page 21 - Indicadores y Decisión"""
        self._add_header("Evaluación", "Indicadores y decisión")
        
        self._add_section_title("Indicadores y decisión")
        
        # Evaluación económica
        self._add_subsection_title("01 - Evaluación económica")
        
        eval_eco = content.get("evaluacion_economica", {})
        
        # Complex header table
        table = self.doc.add_table(rows=3, cols=6)
        table.style = 'Table Grid'
        
        # Header row 1
        table.rows[0].cells[0].text = "Indicadores de rentabilidad"
        self._set_cell_shading(table.rows[0].cells[0], "0099CC")
        table.rows[0].cells[3].text = "Indicadores de costo-eficiencia"
        self._set_cell_shading(table.rows[0].cells[3], "0099CC")
        table.rows[0].cells[4].text = "Indicadores de costo mínimo"
        self._set_cell_shading(table.rows[0].cells[4], "0099CC")
        
        # Header row 2
        headers2 = ["Valor Presente Neto (VPN)", "Tasa Interna de Retorno (TIR)", 
                   "Relación Costo Beneficio (RCB)", "Costo por beneficiario",
                   "Valor presente de los costos", "Costo Anual Equivalente (CAE)"]
        for i, h in enumerate(headers2):
            table.rows[1].cells[i].text = h
            self._set_cell_shading(table.rows[1].cells[i], "7FC8D8")
            for run in table.rows[1].cells[i].paragraphs[0].runs:
                run.font.size = Pt(8)
        
        # Values row
        table.rows[2].cells[0].text = eval_eco.get("vpn", "")
        table.rows[2].cells[1].text = eval_eco.get("tir", "")
        table.rows[2].cells[2].text = eval_eco.get("rcb", "")
        table.rows[2].cells[3].text = eval_eco.get("costo_beneficiario", "")
        table.rows[2].cells[4].text = eval_eco.get("valor_presente_costos", "")
        table.rows[2].cells[5].text = eval_eco.get("cae", "")
        
        self.doc.add_paragraph()
        
        # Costo por capacidad
        self._add_subsection_title("Costo por capacidad")
        
        costo_cap = content.get("costo_capacidad", {})
        productos = costo_cap.get("productos", [])
        
        # Create table with header + rows for each product
        num_rows = max(2, len(productos) + 1)  # At least 2 rows (header + 1 data)
        table2 = self.doc.add_table(rows=1, cols=2)
        table2.style = 'Table Grid'
        table2.rows[0].cells[0].text = "Producto"
        table2.rows[0].cells[1].text = "Costo unitario (valor presente)"
        for cell in table2.rows[0].cells:
            self._set_cell_shading(cell, "0099CC")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Add product rows
        if productos:
            for prod in productos:
                row = table2.add_row()
                row.cells[0].text = prod.get("nombre", "")
                row.cells[1].text = prod.get("costo", "")
        else:
            # Fallback to old single product format
            row = table2.add_row()
            row.cells[0].text = costo_cap.get("producto", "")
            row.cells[1].text = costo_cap.get("costo_unitario", "")
        
        self.doc.add_paragraph()
        
        # Decisión
        self._add_subsection_title("03 - Decisión")
        
        decision = content.get("decision", {})
        self._add_field("Alternativa", decision.get("alternativa", ""))
        
        self.doc.add_paragraph()
        
        # Alcance
        self._add_subsection_title("04 - Alcance")
        
        p = self.doc.add_paragraph()
        p.add_run(content.get("alcance", ""))
    
    def _add_pages_17_21(self, content: dict):
        """Add pages 17-21"""
        # Page 17 - Riesgos Continuation
        self._add_page_17_riesgos_continuacion(content.get("pagina_17_riesgos_continuacion", {}))
        
        self.doc.add_page_break()
        
        # Pages 18-19 - Ingresos y Beneficios
        self._add_page_18_19_ingresos_beneficios(content.get("pagina_18_19_ingresos_beneficios", {}))
        
        self.doc.add_page_break()
        
        # Page 20 - Flujo Económico
        self._add_page_20_flujo_economico(content.get("pagina_20_flujo_economico", {}))
        
        self.doc.add_page_break()
        
        # Page 21 - Indicadores y Decisión
        self._add_page_21_indicadores_decision(content.get("pagina_21_indicadores_decision", {}))
    
    def _add_pages_indicadores(self, content: dict):
        """Add dynamic pages for Indicadores de producto (one per product)"""
        # Get list of indicators
        indicadores = content.get("indicadores_producto", [])
        
        # Fallback for old structure if AI returns single object
        if not indicadores:
            old_single = content.get("pagina_22_indicadores_producto")
            if old_single:
                indicadores = [old_single]
        
        # If still empty, add at least one placeholder
        if not indicadores:
             indicadores = [{
                 "objetivo": {"numero": "1", "descripcion": ""},
                 "producto": {"nombre": "", "codigo": "1.1"},
                 "indicador": {"nombre": "", "codigo": "1.1.1"}
             }]

        for idx, ind_data in enumerate(indicadores):
            self._add_page_indicador(ind_data)
            # Add page break after each indicator page, except potentially the last if that's the end of doc
            # But usually it's cleaner to end with a break or let the save handle it.
            # We will add a break if it's not the last one, OR if we want to ensure separation.
            # Given MGA structure, usually each new page starts fresh.
            if idx < len(indicadores) - 1:
                self.doc.add_page_break()

    def _add_page_indicador(self, data: dict):
        """Render a single 'Indicadores de producto' page matching the client template"""
        
        # Header: Programación / Indicadores de producto
        self._add_header("Programación", "Indicadores de producto")
        
        # Spacing
        self.doc.add_paragraph()
        
        # 1. Producto Section
        # Title "Producto" in Green
        p = self.doc.add_paragraph()
        run = p.add_run("Producto")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(112, 173, 71) # Client Green
        run.bold = True
        
        # Gray Box for Product Name
        prod = data.get("producto", {})
        prod_text = f"{prod.get('codigo', '')} {prod.get('nombre', '')}"
        
        table_prod = self.doc.add_table(rows=1, cols=1)
        table_prod.autofit = True
        cell_prod = table_prod.rows[0].cells[0]
        self._set_cell_shading(cell_prod, "F2F2F2") # Light Gray
        
        p_prod_val = cell_prod.paragraphs[0]
        p_prod_val.paragraph_format.space_before = Pt(2)
        p_prod_val.paragraph_format.space_after = Pt(2)
        run_prod_val = p_prod_val.add_run(prod_text)
        run_prod_val.font.size = Pt(10)
        
        self.doc.add_paragraph()
        
        # 2. Indicador Section
        # Title "Indicador" in Green
        p = self.doc.add_paragraph()
        run = p.add_run("Indicador")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(112, 173, 71) # Client Green
        run.bold = True
        
        # Gray Box for Indicator Details
        ind = data.get("indicador", {})
        
        table_ind = self.doc.add_table(rows=1, cols=1)
        table_ind.autofit = True
        cell_ind = table_ind.rows[0].cells[0]
        self._set_cell_shading(cell_ind, "F2F2F2") # Light Gray
        
        # Content lines
        lines = []
        lines.append((f"{ind.get('codigo', '')} {ind.get('nombre', '')}", False)) # Indicator Name
        lines.append((f"Medido a través de: {ind.get('medido', '')}", True)) # Bold label prefix manually handled
        lines.append((f"Meta total: {ind.get('meta_total', '')}", True))
        lines.append((f"Fórmula: {ind.get('formula', '')}", True))
        lines.append((f"Es acumulativo: {ind.get('es_acumulativo', '')}", True))
        lines.append((f"Es Principal: {ind.get('es_principal', '')}", True))
        lines.append((f"Tipo de Fuente: {ind.get('tipo_fuente', '')}", True))
        lines.append((f"Fuente de Verificación: {ind.get('fuente_verificacion', '')}", True))
        
        p = cell_ind.paragraphs[0]
        for text, has_bold_prefix in lines:
            if has_bold_prefix and ":" in text:
                label, val = text.split(":", 1)
                run_l = p.add_run(label + ":")
                run_l.bold = True
                run_l.font.size = Pt(10)
                run_v = p.add_run(val)
                run_v.font.size = Pt(10)
            else:
                run = p.add_run(text)
                run.font.size = Pt(10)
            
            # Add newline (except last)
            p.add_run("\n")
            
        self.doc.add_paragraph()
        
        # 3. Programación de indicadores Section
        p = self.doc.add_paragraph()
        run = p.add_run("Programación de indicadores")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(112, 173, 71) # Client Green
        run.bold = True
        
        programacion = data.get("programacion_indicadores", [])
        if not programacion:
            # Add default if empty
            programacion = [{"periodo": "0", "meta": ind.get("meta_total", "0")}]
            
        # Table with max 2 periods per row block if we want to mimic the wide style, 
        # or just a simple list. The image shows:
        # | Periodo | Meta por periodo | Periodo | Meta por periodo |
        # So it's a 4-column table filling the width.
        
        table_prog = self.doc.add_table(rows=1, cols=4)
        table_prog.style = 'Table Grid'
        
        # Headers: Periodo | Meta | Periodo | Meta
        headers = ["Periodo", "Meta por periodo", "Periodo", "Meta por periodo"]
        for i, h in enumerate(headers):
            cell = table_prog.rows[0].cells[i]
            cell.text = h
            self._set_cell_shading(cell, "0099CC") # Client Blue
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)
                
        # Data Rows
        # Valid data items
        items = programacion
        
        # We process items in pairs (2 per row: 0,1 ; 2,3 etc)
        import math
        num_rows = math.ceil(len(items) / 2)
        
        for r in range(num_rows):
            row = table_prog.add_row()
            
            # First item of pair
            idx1 = r * 2
            if idx1 < len(items):
                item = items[idx1]
                row.cells[0].text = str(item.get("periodo", ""))
                row.cells[1].text = str(item.get("meta", ""))
                self._set_cell_shading(row.cells[0], "E8F4F8") # Light Blue for period? image shows blue header, light blue data
                self._set_cell_shading(row.cells[1], "E8F4F8")
            
            # Second item of pair
            idx2 = (r * 2) + 1
            if idx2 < len(items):
                item = items[idx2]
                row.cells[2].text = str(item.get("periodo", ""))
                row.cells[3].text = str(item.get("meta", ""))
                # Shade second pair too? Image isn't fully clear on alt rows vs blocks. 
                # Let's keep it clean or same shading.


    def _add_pages_regionalizacion(self, content: dict):
        """Add dynamic pages for Regionalizacion (one block per product)"""
        reg_productos = content.get("regionalizacion_productos", [])
        if not reg_productos:
            return

        self._add_header("Programación", "Regionalización")
        self._add_section_title("Regionalización")
        
        for idx, prod_data in enumerate(reg_productos):
            # Container paragraph for spacing
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(0)

            # Create ONE single table for the entire block to ensure alignment
            # Structure:
            # Row 0: Product Name (Gray) - Merged
            # Row 1: Location Headers (Blue) - Cols [0], [1], [2], [3], [4] (last two merged?? No, 5 cols)
            # Row 2: Location Data - Cols [0], [1], [2], [3], [4]
            # Row 3: Cost Headers (Blue) - Needs 6 cols.
            # MIXED COLUMNS ISSUE: Word tables are grid based. 
            # Solution: Use a 6-column grid.
            # Loc Headers: [0](1), [1](1), [2](2 merged), [3](1), [4](1) -> 6 grid units?
            # Let's simple use two separate tables but REMOVE the paragraph break between them to make them look merged.
            
            # --- Table 1: Product Header ---
            t1 = self.doc.add_table(rows=1, cols=1)
            t1.autofit = True
            cell = t1.rows[0].cells[0]
            self._set_cell_shading(cell, "E0E0E0") # Light Gray
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.add_run("Producto: ").bold = True
            p.add_run(prod_data.get("producto", ""))
            
            # --- Table 2: Location (5 cols) ---
            # To visually merge, we rely on 0 spacing.
            t2 = self.doc.add_table(rows=2, cols=5)
            t2.style = 'Table Grid'
            
            headers_loc = ["Región", "Departamento", "Municipio", "Tipo de Agrupación", "Agrupación"]
            for i, h in enumerate(headers_loc):
                cell = t2.rows[0].cells[i]
                cell.text = h
                self._set_cell_shading(cell, "0099CC") # Client Blue
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255,255,255)
                    run.font.bold = True
                    run.font.size = Pt(9)

            ubicacion = prod_data.get("ubicacion", {})
            t2.rows[1].cells[0].text = ubicacion.get("region", "Caribe")
            t2.rows[1].cells[1].text = ubicacion.get("departamento", "")
            t2.rows[1].cells[2].text = ubicacion.get("municipio", "")
            t2.rows[1].cells[3].text = ubicacion.get("tipo_agrupacion", "")
            t2.rows[1].cells[4].text = ubicacion.get("agrupacion", "")
            
            for cell in t2.rows[1].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

            # --- Table 3: Costs (6 cols) ---
            tabla_costos = prod_data.get("tabla_costos", [])
            if not tabla_costos:
                tabla_costos = [{"periodo": "0", "costo_total": "0"}]
                
            t3 = self.doc.add_table(rows=1 + len(tabla_costos), cols=6)
            t3.style = 'Table Grid'
            
            headers_cost = ["Periodo", "Costo Total", "Costo\nRegionalizado", "Meta Total", "Meta\nRegionalizada", "Beneficiarios"]
            for i, h in enumerate(headers_cost):
                cell = t3.rows[0].cells[i]
                cell.text = h
                self._set_cell_shading(cell, "8DB4E2") # Periwinkle Blue
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255,255,255)
                    run.font.bold = True
                    run.font.size = Pt(9)
            
            for r_idx, item in enumerate(tabla_costos):
                row = t3.rows[r_idx + 1]
                row.cells[0].text = str(item.get("periodo", ""))
                row.cells[1].text = str(item.get("costo_total", ""))
                row.cells[2].text = str(item.get("costo_regionalizado", ""))
                row.cells[3].text = str(item.get("meta_total", ""))
                row.cells[4].text = str(item.get("meta_regionalizada", ""))
                row.cells[5].text = str(item.get("beneficiarios", ""))
                
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)
            
            # Post-processing to "merge" them visually:
            # 1. Ensure columns widths are aligned roughly (Total width ~7 inches)
            # T2 (5 cols): 1.4 inches each
            # T3 (6 cols): 1.16 inches each
            # This mismatch is natural. The screenshot shows they are NOT aligned vertically.
            # The key to "looking like the screenshot" is NO SPACE between tables.
            
            pass 
            # Word automatically puts space unless we configure paragraph spacing tightly.
            # The python-docx library adds separate tables. 
            # We can try to reduce the paragraph formatting between them.

    def _add_page_focalizacion(self, content: dict):
        """Add Page 24 - Focalización (Advanced Report Style)"""
        focalizacion = content.get("focalizacion", [])
        
        if not focalizacion:
             focalizacion = [{"politica": "", "categoria": "", "subcategoria": "", "valor": "0"}]

        # --- Data Grouping Logic ---
        # Structure: { Policy: { Category: [Items...] } }
        grouped_data = {}
        for item in focalizacion:
            pol = item.get("politica", "General")
            cat = item.get("categoria", "General")
            if pol not in grouped_data:
                grouped_data[pol] = {}
            if cat not in grouped_data[pol]:
                grouped_data[pol][cat] = []
            grouped_data[pol][cat].append(item)

        self._add_header("Programación", "Focalización")
        self._add_section_title("Focalización")
        
        # Determine total rows needed
        # Header + Use a list to store row definitions first to calculate spans
        rows_to_render = [] # List of tuples/dicts representing row data
        
        # Iterate Grouped Data
        for policy, categories in grouped_data.items():
            policy_start_idx = len(rows_to_render)
            policy_total_val = 0.0
            
            for category, items in categories.items():
                cat_start_idx = len(rows_to_render)
                cat_total_val = 0.0
                
                # Items
                for item in items:
                    try:
                        val = float(str(item.get("valor", "0")).replace(".", "").replace(",", ".").replace("$", "").strip())
                    except:
                        val = 0.0
                    cat_total_val += val
                    
                    rows_to_render.append({
                        "type": "item",
                        "policy": policy,
                        "category": category,
                        "subcategory": item.get("subcategoria", ""),
                        "value": item.get("valor", "")
                    })
                
                # Category Total Row
                policy_total_val += cat_total_val
                rows_to_render.append({
                    "type": "total_cat",
                    "policy": policy,
                    "category": category, # Still part of this category block for merging
                    "subcategory": "Total Categoría",
                    "value": f"${cat_total_val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                })
                
            # Policy Total Row
            rows_to_render.append({
                "type": "total_pol",
                "policy": policy, # Still part of this policy block for merging
                "category": "TOTAL POLITICA TRANSVERSAL",
                "subcategory": "", 
                "value": f"${policy_total_val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            })

        # --- Table Rendering ---
        table = self.doc.add_table(rows=1 + len(rows_to_render), cols=4)
        table.style = 'Table Grid'
        
        # Headers
        headers = ["Política", "Categoría", "SubCategoría", "Valor"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            self._set_cell_shading(cell, "0099CC") # Client Blue
            self._set_cell_margins(cell, top=100, bottom=100) # Increased padding (approx 5pt)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(10)

        # Render Rows
        LIGHT_BLUE_BG = "E8F4F8" # Very light blue match
        
        for r_idx, row_data in enumerate(rows_to_render):
            table_row = table.rows[r_idx + 1] # +1 for header
            
            # Content
            # P/C handling is done via merging later, but set text for reference or first items
            if row_data["type"] == "total_pol":
                table_row.cells[1].text = row_data["category"] # "TOTAL POLITICA..."
                table_row.cells[3].text = row_data["value"]
                # Col 1 is merged (policy)
                # Col 2 spans Col 2-3? Or just text in Col 2. Reference shows "TOTAL..." in Cat column.
                
            elif row_data["type"] == "total_cat":
                table_row.cells[2].text = row_data["subcategory"] # "Total Categoría" uses SubCat column visually?
                # Actually screenshot shows "Total Categoría" in the THIRD column (SubCategoría)? 
                # Wait, "Total Categoría" text is aligned right next to value? 
                # Screenshot 2: "Total Categoría" is in SubCategoría column.
                table_row.cells[2].text = "Total Categoría"
                table_row.cells[3].text = row_data["value"]
                
            else: # Item
                table_row.cells[0].text = row_data["policy"]
                table_row.cells[1].text = row_data["category"]
                table_row.cells[2].text = row_data["subcategory"]
                table_row.cells[3].text = str(row_data["value"])
            
            # Styling (All rows light blue)
            for cell in table_row.cells:
                self._set_cell_shading(cell, LIGHT_BLUE_BG)
                self._set_cell_margins(cell, top=100, bottom=100) # Airy look
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
            
            # Formatting specifics
            if row_data["type"] == "total_cat":
                # Make "Total Categoría" bold?
                pass
            
            # Align Value
            table_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # --- Merging Logic ---
        # Scan through rows_to_render to find spans
        
        # Merge Policies
        start_idx = 0
        current_pol = rows_to_render[0]["policy"]
        for i in range(1, len(rows_to_render)):
            pol = rows_to_render[i]["policy"]
            if pol != current_pol:
                # End of block, merge previous
                self._merge_vertically(table, 0, start_idx + 1, i) # i is now header+i (end exclusive?) -> +1 for header
                # update
                current_pol = pol
                start_idx = i
        # Merge last block
        self._merge_vertically(table, 0, start_idx + 1, len(rows_to_render))

        # Merge Categories (Only within same policy)
        # Be careful not to merge "TOTAL POLITICA" row into category logic if it breaks visual
        start_idx = 0
        current_cat = rows_to_render[0]["category"]
        current_pol_for_cat = rows_to_render[0]["policy"]
        
        for i in range(1, len(rows_to_render)):
            row = rows_to_render[i]
            # Break merge if: Policy changes OR Category changes OR Type is total_pol
            if row["policy"] != current_pol_for_cat or row["category"] != current_cat or row["type"] == "total_pol":
                # Merge previous block
                # Don't merge if it's just 1 row
                if i - start_idx > 1:
                     self._merge_vertically(table, 1, start_idx + 1, i)
                
                # Reset
                start_idx = i
                current_cat = row["category"]
                current_pol_for_cat = row["policy"]
        # Last block
        if len(rows_to_render) - start_idx > 1:
             # Check if last row is total_pol, assume it breaks category content
             last_type = rows_to_render[-1]["type"]
             if last_type != "total_pol":
                self._merge_vertically(table, 1, start_idx + 1, len(rows_to_render))

    def _merge_vertically(self, table, col_idx, start_row_idx, end_row_idx):
        """Merge cells vertically from start_row_idx to end_row_idx (exclusive of end)"""
        # docx logic: set vMerge=restart on first, vMerge=continue on rest
        if end_row_idx <= start_row_idx:
            return
            
        cell_start = table.rows[start_row_idx].cells[col_idx]
        tc_start = cell_start._tc
        tcPr_start = tc_start.get_or_add_tcPr()
        vMerge_start = OxmlElement('w:vMerge')
        vMerge_start.set(qn('w:val'), 'restart')
        tcPr_start.append(vMerge_start)
        
        for i in range(start_row_idx + 1, end_row_idx):
            cell = table.rows[i].cells[col_idx]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vMerge = OxmlElement('w:vMerge')
            # attribute val defaults to 'continue' if omitted, but explicit is safer
            tcPr.append(vMerge)
            
            # Clear text in merged cells to avoid duplication display issues in some viewers
            cell.text = ""

    def _set_cell_margins(self, cell, top=0, bottom=0, start=0, end=0):
        """Set cell margins (padding) in Twips (1/20 pt)"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        
        if top > 0:
            node = OxmlElement('w:top')
            node.set(qn('w:w'), str(top))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        if bottom > 0:
            node = OxmlElement('w:bottom')
            node.set(qn('w:w'), str(bottom))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
            
        # Add to properties (replace existing if any? simplified append here)
        existing = tcPr.find(qn('w:tcMar'))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(tcMar)

    def _set_cell_shading(self, cell, color):
        """Set cell background color"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        tcPr.append(shading)
    
    def _save_document(self, data):
        """Save the document to file as PDF"""
        from docx2pdf import convert
        import platform
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Handle case where data is not a dict
        if isinstance(data, dict):
            municipio = data.get("municipio", "documento")
        else:
            municipio = "documento"
        
        # Remove invalid filename characters
        municipio = re.sub(r'[^a-zA-Z0-9_\-]', '', municipio.replace(" ", "_"))
        if not municipio:
            municipio = "documento"
        
        # First save as DOCX
        docx_filename = f"MGA_{municipio}_{timestamp}.docx"
        docx_filepath = os.path.join(self.output_dir, docx_filename)
        self.doc.save(docx_filepath)
        
        # Convert to PDF
        pdf_filename = f"MGA_{municipio}_{timestamp}.pdf"
        pdf_filepath = os.path.join(self.output_dir, pdf_filename)
        
        try:
            # Initialize COM for Windows (required for docx2pdf in threaded environments)
            if platform.system() == 'Windows':
                import pythoncom
                pythoncom.CoInitialize()
            
            convert(docx_filepath, pdf_filepath)
            
            # Uninitialize COM
            if platform.system() == 'Windows':
                pythoncom.CoUninitialize()
            
            # Remove temporary DOCX file
            os.remove(docx_filepath)
            return pdf_filepath
        except Exception as e:
            # If PDF conversion fails, return DOCX
            print(f"PDF conversion failed: {e}, returning DOCX")
            return docx_filepath

