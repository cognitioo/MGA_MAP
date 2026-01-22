"""
Direct Estudios Previos Document Builder
Builds Word documents directly matching the MGA template structure
- Continuous bordered sections (no gaps between sections)
- Tables inside sections share the outer border (no nested box effect)
- Tables use horizontal separators only, no vertical column dividers
"""

import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Colors
GRAY_HEADER = "BFBFBF"      # Light gray for section headers
DARK_HEADER = "404040"      # Dark gray for table column headers  
BLACK = "000000"
WHITE = "FFFFFF"


def sanitize_filename(filename: str) -> str:
    """Remove invalid characters from filename"""
    invalid_chars = r'[<>:"/\\|?*\t\n\r]'
    sanitized = re.sub(invalid_chars, '', filename)
    sanitized = re.sub(r'[\s_]+', '_', sanitized)
    return sanitized.strip('_')


class EstudiosPreviosDirectBuilder:
    """Builds Estudios Previos document directly with proper MGA template structure"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.doc = None
    
    def build(self, data: dict, ai_content: dict, letterhead_file=None) -> str:
        """
        Build the complete Estudios Previos document
        
        Args:
            data: Form data dictionary
            ai_content: AI-generated content dictionary
            letterhead_file: Optional uploaded .docx file with header/footer template
        """
        # Track if letterhead is used (to preserve margins in _apply_styles)
        self._has_letterhead = letterhead_file is not None
        
        # If letterhead provided, use it as base template (preserves images/logos)
        if letterhead_file:
            self.doc = self._load_template(letterhead_file)
        else:
            self.doc = Document()
        
        self._apply_styles()
        
        # 1. Main Title (NOT in a box)
        self._add_main_title(data)
        
        # 2. Header Table (DEPENDENCIA, FECHA, PROCESO)
        self._add_header_table(data)
        
        # 3. Subtitle
        self._add_subtitle()
        
        # 4. Section 1: MARCO LEGAL
        self._add_section_with_content(
            "1. MARCO LEGAL",
            ai_content.get("marco_legal", "")
        )
        
        # 5. Section 2: NECESIDAD QUE SATISFACE LA CONTRATACIÓN
        self._add_section_with_content(
            "2. NECESIDAD QUE SATISFACE LA CONTRATACIÓN",
            ai_content.get("necesidad", "")
        )
        
        # 6. Section 3: OBJETO Y ALCANCE (includes budget table)
        self._add_objeto_alcance_section(data, ai_content)
        
        # 7. Section 4: OBLIGACIONES DEL MUNICIPIO
        self._add_obligaciones_municipio_section(data, ai_content.get("obligaciones", {}))
        
        # 8. Section 5: OBLIGACIONES DEL CONTRATISTA
        self._add_obligaciones_contratista_section(data, ai_content.get("obligaciones", {}))
        
        # 9. Section 6: FUNDAMENTOS JURÍDICOS QUE SOPORTAN LA MODALIDAD DE SELECCIÓN
        self._add_section_with_content(
            "6. FUNDAMENTOS JURÍDICOS QUE SOPORTAN LA MODALIDAD DE SELECCIÓN",
            ai_content.get("fundamentos", "")
        )
        
        # 10. Section 7: ANÁLISIS QUE SOPORTA EL VALOR ESTIMADO Y FORMA DE PAGO
        self._add_section_with_content(
            "7. ANÁLISIS QUE SOPORTA EL VALOR ESTIMADO Y FORMA DE PAGO DEL CONTRATO",
            ai_content.get("analisis_valor", "")
        )
        
        # 10.1 CDP Table (inside Section 7)
        self._add_cdp_table(data.get("cdp_data", {}))
        
        # 11. Section 8: ANÁLISIS DE RIESGOS
        self._add_riesgos_section(ai_content.get("riesgos", []))
        
        # 12. Section 9: ANÁLISIS QUE SUSTENTA LA EXIGENCIA DE GARANTÍAS
        self._add_section_with_content(
            "9. ANÁLISIS QUE SUSTENTA LA EXIGENCIA DE GARANTÍAS",
            ai_content.get("garantias", "")
        )
        
        # 13. Section 10: PLAZO Y LUGAR DE EJECUCIÓN
        self._add_section_with_content(
            "10. PLAZO Y LUGAR DE EJECUCIÓN",
            ai_content.get("plazo_lugar", "")
        )
        
        # 14. Section 11: SUPERVISIÓN
        self._add_section_with_content(
            "11. SUPERVISIÓN",
            ai_content.get("supervision", "")
        )
        
        # 15. Section 12: RESPONSABLES
        self._add_responsables_section(data)
        
        # Save document
        bpin = sanitize_filename(data.get('bpin', 'DRAFT'))
        filename = f"Estudios_Previos_{bpin}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        self.doc.save(filepath)
        
        return filepath
    
    def _apply_styles(self):
        """Apply base document styles - tighter spacing (preserves letterhead margins)"""
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(8)  # Smaller default font
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        
        # Only set margins if no letterhead was loaded (letterhead margins set in _load_template)
        if not getattr(self, '_has_letterhead', False):
            for section in self.doc.sections:
                section.top_margin = Cm(1.5)
                section.bottom_margin = Cm(1.5)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
    
    def _load_template(self, letterhead_file):
        """
        Load the letterhead template as the base document.
        This preserves all images, logos, and graphics in headers/footers.
        Clears only the body content, keeping header/footer intact.
        """
        import io
        
        try:
            # Read the template document
            if hasattr(letterhead_file, 'read'):
                template_bytes = letterhead_file.read()
                letterhead_file.seek(0)  # Reset for potential re-read
                doc = Document(io.BytesIO(template_bytes))
            else:
                doc = Document(letterhead_file)
            
            # Clear all body content (keep headers/footers)
            for element in doc.element.body[:]:
                # Don't remove sectPr (section properties) as they contain header/footer refs
                if not element.tag.endswith('sectPr'):
                    doc.element.body.remove(element)
            
            # Set margins to leave room for header/footer graphics on ALL pages
            for section in doc.sections:
                section.different_first_page_header_footer = False
                # Set proper margins for content area (leaves room for letterhead)
                section.top_margin = Cm(4)        # Distance from top of page to content
                section.bottom_margin = Cm(3.5)   # Distance from bottom of page to content
                # header_distance = space from page edge to header
                # footer_distance = space from page edge to footer
                section.header_distance = Cm(1)   # Keep header at top
                section.footer_distance = Cm(1)   # Keep footer at bottom
            
            return doc
            
        except Exception as e:
            print(f"Warning: Could not load letterhead template: {e}")
            return Document()  # Return empty document as fallback
    
    def _add_main_title(self, data: dict):
        """Add main document title (centered, bold, BIG, NOT in a box)"""
        municipio = data.get('municipio', '').upper()
        depto = data.get('departamento', '').upper()
        
        title = f"ESTUDIOS PREVIOS PARA LA SUSCRIPCIÓN DEL CONVENIO INTERADMINISTRATIVO ENTRE EL MUNICIPIO DE {municipio} – {depto} Y EMACALA S.A.S E.S.P. PARA LA ACTUALIZACIÓN Y REVISIÓN DEL PLAN DE SANEAMIENTO Y MANEJO DE VERTIMIENTOS – PSMV"
        
        para = self.doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(12)  # BIGGER title to match client
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(24)  # Push down from letterhead header
        para.paragraph_format.space_after = Pt(6)  # Slight space after title
    
    def _add_header_table(self, data: dict):
        """Add header table (DEPENDENCIA, FECHA, PROCESO) with gray left column"""
        table = self.doc.add_table(rows=3, cols=2)
        table.autofit = True
        self._set_table_borders(table)
        
        # Spanish month names
        SPANISH_MONTHS = {
            1: "ENERO", 2: "FEBRERO", 3: "MARZO", 4: "ABRIL",
            5: "MAYO", 6: "JUNIO", 7: "JULIO", 8: "AGOSTO",
            9: "SEPTIEMBRE", 10: "OCTUBRE", 11: "NOVIEMBRE", 12: "DICIEMBRE"
        }
        now = datetime.now()
        spanish_date = f"{SPANISH_MONTHS[now.month]} DE {now.year}"
        
        rows_data = [
            ("DEPENDENCIA QUE PROYECTA", data.get("dependencia", "SECRETARÍA DE PLANEACIÓN")),
            ("FECHA", spanish_date),
            ("PROCESO", data.get("proceso", "CONTRATACIÓN DIRECTA"))
        ]
        
        for i, (label, value) in enumerate(rows_data):
            cell_label = table.rows[i].cells[0]
            cell_value = table.rows[i].cells[1]
            
            cell_label.text = label
            cell_value.text = value
            
            # Gray background on left column (label)
            self._shade_cell(cell_label, GRAY_HEADER)
            
            for para in cell_label.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(8)  # Smaller
            
            for para in cell_value.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(8)  # Smaller
        
        # No spacing after header table - flows directly
        # Removed paragraph spacing
    
    def _add_subtitle(self):
        """Add subtitle lines - very compact"""
        lines = [
            "DEPARTAMENTO NACIONAL DE PLANEACIÓN – METODOLOGÍA GENERAL AJUSTADA (MGA)",
            "ESTUDIOS PREVIOS – CONVENIO INTERADMINISTRATIVO"
        ]
        for line in lines:
            para = self.doc.add_paragraph()
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(1)  # Minimal
            run = para.add_run(line)
            run.bold = True
            run.font.size = Pt(8)  # Smaller
        
        # Tiny spacing before first section
        spacing = self.doc.add_paragraph()
        spacing.paragraph_format.space_after = Pt(1)
    
    def _add_encabezado_section(self, data: dict):
        """Add ENCABEZADO CON DATOS DEL CONTRATO section - simple two-column layout"""
        contract_data = [
            ("Entidad contratante", data.get("entidad", "Alcaldía Municipal de San Pablo")),
            ("Municipio", f"{data.get('municipio', '')} ({data.get('departamento', '')})"),
            ("Departamento", data.get("departamento", "")),
            ("Tipo de proyecto", data.get("tipo_proyecto", "Convenio Interadministrativo")),
            ("Código BPIN", data.get("bpin", "")),
            ("Objeto del convenio", data.get("objeto", "")),
            ("Responsable de la contratación", data.get("responsable", "")),
            ("Fuente de financiación", data.get("fuente_financiacion", "")),
            ("Valor total estimado", f"${data.get('valor_total', '0')} pesos"),
            ("Plazo de ejecución", f"{data.get('plazo', '90')} días calendario"),
            ("Lugar de ejecución", data.get("lugar", f"Municipio de {data.get('municipio', '')}, Departamento de {data.get('departamento', '')}"))
        ]
        
        # Create single table with header row + data rows
        num_rows = len(contract_data) + 1  # +1 for header
        table = self.doc.add_table(rows=num_rows, cols=2)
        self._set_table_borders(table)
        
        # Header row spans both columns (merge)
        header_cell_a = table.rows[0].cells[0]
        header_cell_b = table.rows[0].cells[1]
        header_cell_a.merge(header_cell_b)
        header_cell_a.text = "1. ENCABEZADO CON DATOS DEL CONTRATO"
        self._style_header_cell(header_cell_a)
        
        # Data rows
        for i, (label, value) in enumerate(contract_data):
            row = table.rows[i + 1]
            row.cells[0].text = label
            row.cells[1].text = str(value) if value else ""
            
            # Bold first column
            for para in row.cells[0].paragraphs:
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(8)  # Smaller
            
            for para in row.cells[1].paragraphs:
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(8)  # Smaller
        
        # NO spacing after - continuous sections
    
    def _add_section_with_content(self, title: str, content: str):
        """Add a section with gray header and text content"""
        table = self.doc.add_table(rows=2, cols=1)
        self._set_table_borders(table)
        
        # Header
        header_cell = table.rows[0].cells[0]
        header_cell.text = title
        self._style_header_cell(header_cell)
        
        # Content
        content_cell = table.rows[1].cells[0]
        content_cell.paragraphs[0].clear()
        self._add_formatted_text(content_cell, content)
        
        # NO spacing - continuous sections
    
    def _add_section_with_table(self, title: str, table_data: list, headers: list):
        """Add section with gray header and direct table content - flattened"""
        # Rows: 1 header + 1 col headers + N data
        num_cols = len(headers)
        num_data_rows = len(table_data) if table_data else 0
        total_rows = 2 + num_data_rows
        
        table = self.doc.add_table(rows=total_rows, cols=num_cols)
        self._set_outer_border_only(table)
        
        # Row 0: Section Header (merge all cols)
        for i in range(num_cols - 1):
            table.rows[0].cells[0].merge(table.rows[0].cells[i + 1])
        header_cell = table.rows[0].cells[0]
        header_cell.text = title
        self._style_header_cell(header_cell)
        # Add bottom border to header cell
        for cell in table.rows[0].cells:
            self._set_cell_bottom_border(cell)
            
        # Row 1: Column Headers - DARK background with WHITE text
        for i, h in enumerate(headers):
            cell = table.rows[1].cells[i]
            cell.text = h
            self._style_table_column_header(cell)
            self._set_cell_bottom_border(cell)
                    
        # Data Rows
        for row_idx, row_data in enumerate(table_data):
            # Parse data similar to before
            if isinstance(row_data, dict):
                values = list(row_data.values())
            else:
                values = row_data if isinstance(row_data, list) else [row_data]
            
            for col_idx, val in enumerate(values):
                if col_idx < num_cols:
                    # Note: +2 because Row 0 is title, Row 1 is headers
                    cell = table.rows[row_idx + 2].cells[col_idx]
                    cell.text = str(val) if val else ""
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(8)
    
    def _add_objeto_alcance_section(self, data: dict, ai_content: dict):
        """Add OBJETO Y ALCANCE section with text content AND budget table integrated"""
        objeto_text = ai_content.get("objeto_alcance", "")
        rubros = ai_content.get("presupuesto", [])
        # Defensive: ensure rubros is a list
        if not isinstance(rubros, list):
            rubros = []
        valor = data.get('valor_total', '0')
        bpin = data.get('bpin', '')
        
        # Calculate rows: header + content + value line + subtitle + column headers + data rows
        num_data_rows = len(rubros) if rubros else 0
        total_rows = 5 + num_data_rows
        
        table = self.doc.add_table(rows=total_rows, cols=4)
        self._set_outer_border_only(table)
        
        # Row 0: Section header (merge all cols)
        for i in range(3):
            table.rows[0].cells[0].merge(table.rows[0].cells[i + 1])
        table.rows[0].cells[0].text = "3. OBJETO Y ALCANCE"
        self._style_header_cell(table.rows[0].cells[0])
        self._set_cell_bottom_border(table.rows[0].cells[0])
        
        # Row 1: Object/Alcance text content (merge all cols)
        for i in range(3):
            table.rows[1].cells[0].merge(table.rows[1].cells[i + 1])
        self._add_formatted_text(table.rows[1].cells[0], objeto_text)
        
        # Row 2: Budget value line (merge all cols)
        for i in range(3):
            table.rows[2].cells[0].merge(table.rows[2].cells[i + 1])
        val_para = table.rows[2].cells[0].paragraphs[0]
        val_para.paragraph_format.space_after = Pt(2)
        val_run = val_para.add_run(f"Valor total estimado del convenio (MGA – BPIN {bpin}): ${valor} COP.")
        val_run.bold = True
        val_run.font.size = Pt(8)
        
        # Row 3: Subtitle for budget breakdown (merge all cols)
        for i in range(3):
            table.rows[3].cells[0].merge(table.rows[3].cells[i + 1])
        sub_para = table.rows[3].cells[0].paragraphs[0]
        sub_para.paragraph_format.space_after = Pt(0)
        sub_run = sub_para.add_run("Desglose presupuestal propuesto (precios constantes 2025)")
        sub_run.bold = True
        sub_run.font.size = Pt(8)
        
        # Row 4: Column headers - DARK background
        headers = ["Rubro", "Descripción", "%", "Valor (COP)"]
        for i, h in enumerate(headers):
            cell = table.rows[4].cells[i]
            cell.text = h
            self._style_table_column_header(cell)
            self._set_cell_bottom_border(cell)
        
        # Data rows
        for row_idx, rubro in enumerate(rubros):
            if isinstance(rubro, dict):
                data_row = table.rows[5 + row_idx]
                data_row.cells[0].text = rubro.get("nombre", "")
                data_row.cells[1].text = rubro.get("descripcion", "")
                data_row.cells[2].text = rubro.get("porcentaje", "")
                data_row.cells[3].text = rubro.get("valor", "")
                
                for cell in data_row.cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.space_after = Pt(0)
                        for run in para.runs:
                            run.font.size = Pt(8)
    
    def _add_obligaciones_municipio_section(self, data: dict, obligaciones: dict):
        """Add Section 4: OBLIGACIONES DEL MUNICIPIO"""
        municipio = data.get('municipio', 'Municipio')
        
        # Create table: header row + content row
        table = self.doc.add_table(rows=2, cols=1)
        self._set_outer_border_only(table)
        
        # Header row
        header_cell = table.rows[0].cells[0]
        header_cell.text = "4. OBLIGACIONES DEL MUNICIPIO"
        self._style_header_cell(header_cell)
        self._set_cell_bottom_border(header_cell)
        
        # Content row with formatted text
        content_cell = table.rows[1].cells[0]
        self._add_formatted_text(content_cell, obligaciones.get("municipio", ""))
    
    def _add_obligaciones_contratista_section(self, data: dict, obligaciones: dict):
        """Add Section 5: OBLIGACIONES DEL CONTRATISTA"""
        # Create table: header row + content row
        table = self.doc.add_table(rows=2, cols=1)
        self._set_outer_border_only(table)
        
        # Header row
        header_cell = table.rows[0].cells[0]
        header_cell.text = "5. OBLIGACIONES DEL CONTRATISTA"
        self._style_header_cell(header_cell)
        self._set_cell_bottom_border(header_cell)
        
        # Content row with formatted text
        content_cell = table.rows[1].cells[0]
        self._add_formatted_text(content_cell, obligaciones.get("empresa", ""))
    
    def _add_presupuesto_section(self, data: dict, rubros: list):
        """Add PRESUPUESTO ESTIMADO section - simplified structure"""
        valor = data.get('valor_total', '0')
        bpin = data.get('bpin', '')
        
        # Calculate number of rows: 1 header + 1 value text + 1 subtitle + 1 column headers + N data rows
        num_data_rows = len(rubros) if rubros else 0
        total_rows = 4 + num_data_rows  # header, value line, subtitle, column headers, data
        
        table = self.doc.add_table(rows=total_rows, cols=4)
        self._set_outer_border_only(table)
        
        # Row 0: Section header (merge all 4 cols)
        for i in range(3):
            table.rows[0].cells[0].merge(table.rows[0].cells[i + 1])
        table.rows[0].cells[0].text = "7. PRESUPUESTO ESTIMADO"
        self._style_header_cell(table.rows[0].cells[0])
        self._set_cell_bottom_border(table.rows[0].cells[0])
        
        # Row 1: Value statement (merge all 4 cols)
        for i in range(3):
            table.rows[1].cells[0].merge(table.rows[1].cells[i + 1])
        val_para = table.rows[1].cells[0].paragraphs[0]
        val_para.paragraph_format.space_after = Pt(0)
        val_run = val_para.add_run(f"Valor total estimado del convenio (MGA – BPIN {bpin}): ${valor} COP.")
        val_run.font.size = Pt(8)
        
        # Row 2: Subtitle (merge all 4 cols)
        for i in range(3):
            table.rows[2].cells[0].merge(table.rows[2].cells[i + 1])
        sub_para = table.rows[2].cells[0].paragraphs[0]
        sub_para.paragraph_format.space_after = Pt(0)
        sub_run = sub_para.add_run("Desglose presupuestal propuesto (precios constantes 2025)")
        sub_run.bold = True
        sub_run.font.size = Pt(8)
        
        # Row 3: Column headers - DARK background with WHITE text
        headers = ["Rubro", "Descripción", "%", "Valor (COP)"]
        for i, h in enumerate(headers):
            cell = table.rows[3].cells[i]
            cell.text = h
            self._style_table_column_header(cell)
            self._set_cell_bottom_border(cell)
        
        # Data rows
        for row_idx, rubro in enumerate(rubros):
            if isinstance(rubro, dict):
                data_row = table.rows[4 + row_idx]
                data_row.cells[0].text = rubro.get("nombre", "")
                data_row.cells[1].text = rubro.get("descripcion", "")
                data_row.cells[2].text = rubro.get("porcentaje", "")
                data_row.cells[3].text = rubro.get("valor", "")
                
                for cell in data_row.cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.space_after = Pt(0)
                        for run in para.runs:
                            run.font.size = Pt(8)
    
    def _add_riesgos_section(self, riesgos: list):
        """Add ANÁLISIS DE RIESGOS section - simplified structure"""
        # Defensive: ensure riesgos is a list
        if not isinstance(riesgos, list):
            riesgos = []
        # Rows: 1 header + 1 column headers + N data rows
        num_data_rows = len(riesgos) if riesgos else 0
        total_rows = 2 + num_data_rows
        
        table = self.doc.add_table(rows=total_rows, cols=4)
        self._set_outer_border_only(table)
        
        # Row 0: Section header (merge all 4 cols)
        for i in range(3):
            table.rows[0].cells[0].merge(table.rows[0].cells[i + 1])
        table.rows[0].cells[0].text = "8. ANÁLISIS DE RIESGOS"
        self._style_header_cell(table.rows[0].cells[0])
        self._set_cell_bottom_border(table.rows[0].cells[0])
        
        # Row 1: Column headers - DARK background with WHITE text
        headers = ["Riesgo identificado", "Descripción", "Probabilidad", "Medida de mitigación"]
        for i, h in enumerate(headers):
            cell = table.rows[1].cells[i]
            cell.text = h
            self._style_table_column_header(cell)
            self._set_cell_bottom_border(cell)
        
        # Data rows
        for row_idx, r in enumerate(riesgos):
            if isinstance(r, dict):
                data_row = table.rows[2 + row_idx]
                data_row.cells[0].text = r.get("riesgo", "")
                data_row.cells[1].text = r.get("descripcion", r.get("impacto", ""))
                data_row.cells[2].text = r.get("probabilidad", "")
                data_row.cells[3].text = r.get("mitigacion", "")
                
                for cell in data_row.cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.space_after = Pt(0)
                        for run in para.runs:
                            run.font.size = Pt(8)
    
    def _add_cdp_table(self, cdp_data: dict):
        """Add CDP (Certificado de Disponibilidad Presupuestal) table in Section 7
        Columns: CDP, Fecha, Rubro, Fuente, Valor
        """
        # Use provided data or defaults
        cdp = cdp_data.get("cdp", "")
        fecha = cdp_data.get("fecha", "")
        rubro = cdp_data.get("rubro", "")
        fuente = cdp_data.get("fuente", "")
        valor = cdp_data.get("valor", "")
        
        # Create table: 1 header row + 1 data row
        table = self.doc.add_table(rows=2, cols=5)
        self._set_table_full_borders(table)
        
        # Header row with gray background
        headers = ["CDP", "Fecha", "Rubro", "Fuente", "Valor"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            self._shade_cell(cell, GRAY_HEADER)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(8)
        
        # Data row
        data_row = table.rows[1]
        data_row.cells[0].text = cdp
        data_row.cells[1].text = fecha
        data_row.cells[2].text = rubro
        data_row.cells[3].text = fuente
        data_row.cells[4].text = valor
        
        for cell in data_row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(8)
    
    def _set_table_full_borders(self, table):
        """Set full borders on all cells of a table"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), BLACK)
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
    
    def _add_responsables_section(self, data: dict):
        """Add Section 12: RESPONSABLES - NOMBRE/CARGO left, FIRMA right"""
        # Table: 1 header row + 1 content row with 2 columns
        table = self.doc.add_table(rows=2, cols=2)
        self._set_outer_border_only(table)
        
        # Header row - merge and style
        header_a = table.rows[0].cells[0]
        header_b = table.rows[0].cells[1]
        header_a.merge(header_b)
        header_a.text = "12. RESPONSABLES"
        self._style_header_cell(header_a)
        self._set_cell_bottom_border(header_a)
        
        # Content row - left side (name/cargo), right side (firma)
        left = table.rows[1].cells[0]
        lp = left.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        r1 = lp.add_run(f"NOMBRE: {data.get('responsable', '').upper()}")
        r1.bold = True
        r1.font.size = Pt(8)
        lp.add_run("\n")
        r2 = lp.add_run(f"CARGO: {data.get('cargo', 'Secretario de Planeación Municipal')}")
        r2.bold = True
        r2.font.size = Pt(8)
        
        right = table.rows[1].cells[1]
        rp = right.paragraphs[0]
        rp.paragraph_format.space_after = Pt(0)
        rf = rp.add_run("FIRMA:")
        rf.bold = True
        rf.font.size = Pt(8)
    
    def _add_formatted_text(self, cell, text: str):
        """Add text to cell with line break and bold support"""
        if not text:
            return
        
        # Handle <br> tags
        text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
        
        paragraphs = text.split('\n')
        
        first = True
        for p_text in paragraphs:
            p_text = p_text.strip()
            if not p_text:
                continue
            
            if first:
                para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                first = False
            else:
                para = cell.add_paragraph()
            
            para.paragraph_format.space_after = Pt(1)  # Very tight
            para.paragraph_format.space_before = Pt(0)
            
            if p_text.startswith('• ') or p_text.startswith('- ') or p_text.startswith('* '):
                para.paragraph_format.left_indent = Cm(0.3)
                p_text = "• " + p_text[2:]
            
            # Handle **bold**
            pattern = r'(\*\*(.+?)\*\*|([^*]+))'
            for match in re.finditer(pattern, p_text):
                m = match.group(0)
                if m.startswith('**') and m.endswith('**'):
                    run = para.add_run(m[2:-2])
                    run.bold = True
                    run.font.size = Pt(8)  # SMALLER
                else:
                    run = para.add_run(m)
                    run.font.size = Pt(8)  # SMALLER
    
    def _style_header_cell(self, cell):
        """Style a SECTION header cell with light gray background"""
        self._shade_cell(cell, GRAY_HEADER)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(8)
    
    def _style_table_column_header(self, cell):
        """Style a TABLE COLUMN header cell with GRAY background (same as section headers)"""
        self._shade_cell(cell, GRAY_HEADER)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(8)
    
    def _shade_cell(self, cell, color: str):
        """Add background color to cell"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        
    def _set_cell_bottom_border(self, cell):
        """Set bottom border for a single cell"""
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
            
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), BLACK)
        tcBorders.append(bottom)
    
    def _set_table_borders(self, table):
        """Set black borders on outer table (section container) - Full Grid"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        tblBorders = OxmlElement('w:tblBorders')
        for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), BLACK)
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
            
    def _set_outer_border_only(self, table):
        """Set black borders ONLY on outer table box - No internal grid lines"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        tblBorders = OxmlElement('w:tblBorders')
        
        # Set outer borders
        for name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), BLACK)
            tblBorders.append(border)
            
        # Remove inner borders
        for name in ['insideH', 'insideV']:
            border = OxmlElement(f'w:{name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
    
    def _set_inner_table_style(self, table, with_header: bool = False):
        """
        Set style for inner tables:
        - NO vertical column dividers
        - Only horizontal row separators
        - If with_header=True, add bottom border under header row
        """
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        # Set table borders - only horizontal lines
        tblBorders = OxmlElement('w:tblBorders')
        
        # No top, left, right, bottom outer borders
        # Only inside horizontal lines
        insideH = OxmlElement('w:insideH')
        insideH.set(qn('w:val'), 'single')
        insideH.set(qn('w:sz'), '4')
        insideH.set(qn('w:color'), BLACK)
        tblBorders.append(insideH)
        
        # No vertical dividers
        insideV = OxmlElement('w:insideV')
        insideV.set(qn('w:val'), 'nil')
        tblBorders.append(insideV)
        
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
