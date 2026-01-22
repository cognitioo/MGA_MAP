"""
Word Document Builder
Converts generated content to Word (.docx) documents following MGA templates
Professional styling with yellow section headers and bordered boxes
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
from datetime import datetime
from generators.markdown_converter import MarkdownToWordConverter, add_signature_table

# Color constants
HEADER_COLOR = "FFFFCC"  # Light yellow


def sanitize_filename(filename: str) -> str:
    """Remove invalid characters from filename"""
    invalid_chars = r'[<>:"/\\|?*\t\n\r]'
    sanitized = re.sub(invalid_chars, '', filename)
    sanitized = re.sub(r'[\s_]+', '_', sanitized)
    return sanitized.strip('_')


def shade_cell(cell, color: str):
    """Add background color to table cell"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


class DocumentBuilder:
    """Builds Word documents from generated content"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def _apply_base_styles(self, doc: Document):
        """Apply base styles to document"""
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
    
    def _add_header_table(self, doc: Document, data: dict, doc_type: str):
        """Add header table with contract metadata - styled with yellow first column"""
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.autofit = True
        
        # Header data based on document type
        if doc_type == "estudios_previos":
            headers = [
                ("DEPENDENCIA QUE PROYECTA", data.get("dependencia", "SECRETARÍA DE PLANEACIÓN")),
                ("FECHA", datetime.now().strftime("%B DE %Y").upper()),
                ("PROCESO", data.get("proceso", "CONTRATACIÓN DIRECTA"))
            ]
        else:  # analisis_sector
            headers = [
                ("NUMERO DE CONTRATO", data.get("numero_contrato", "")),
                ("MODALIDAD", data.get("modalidad", "Convenio Interadministrativo")),
                ("ENTIDAD", data.get("entidad", ""))
            ]
        
        for i, (label, value) in enumerate(headers):
            cell_label = table.rows[i].cells[0]
            cell_value = table.rows[i].cells[1]
            
            cell_label.text = label
            cell_value.text = value
            
            # Make label bold
            for para in cell_label.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)
            
            # Style value
            for para in cell_value.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        
        doc.add_paragraph()
    
    def _add_title(self, doc: Document, title: str, subtitle: str = None):
        """Add document title with optional subtitle"""
        # Main title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_after = Pt(6)
        
        # Subtitle if provided
        if subtitle:
            sub_para = doc.add_paragraph()
            sub_run = sub_para.add_run(subtitle)
            sub_run.bold = True
            sub_run.font.size = Pt(11)
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_para.space_after = Pt(12)
        
        doc.add_paragraph()
    
    def build_estudios_previos(self, content: str, data: dict) -> str:
        """Build Estudios Previos Word document matching MGA template"""
        doc = Document()
        self._apply_base_styles(doc)
        
        # Add main title (matching template format)
        municipio = data.get('municipio', '').upper()
        departamento = data.get('departamento', '').upper()
        
        title = f"ESTUDIOS PREVIOS PARA LA SUSCRIPCIÓN DEL CONVENIO INTERADMINISTRATIVO ENTRE EL MUNICIPIO DE {municipio} – {departamento} Y EMACALA S.A.S E.S.P. PARA LA ACTUALIZACIÓN Y REVISIÓN DEL PLAN DE SANEAMIENTO Y MANEJO DE VERTIMIENTOS – PSMV"
        
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_after = Pt(12)
        
        # Add header table
        self._add_header_table(doc, data, "estudios_previos")
        
        # Add content using markdown converter
        converter = MarkdownToWordConverter(doc)
        converter.add_formatted_content(content)
        
        # Add signature table (matching template)
        add_signature_table(doc, data)
        
        # Save document
        bpin_clean = sanitize_filename(data.get('bpin', 'DRAFT'))
        filename = f"Estudios_Previos_{bpin_clean}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def build_analisis_sector(self, content: str, data: dict) -> str:
        """Build Análisis del Sector Word document"""
        doc = Document()
        self._apply_base_styles(doc)
        
        # Add title
        title = "ANÁLISIS DEL SECTOR ECONÓMICO Y DE LOS OFERENTES POR PARTE DE LAS ENTIDADES ESTATALES"
        self._add_title(doc, title)
        
        # Add header table
        self._add_header_table(doc, data, "analisis_sector")
        
        # Add content using markdown converter
        converter = MarkdownToWordConverter(doc)
        converter.add_formatted_content(content)
        
        # Add signature table
        add_signature_table(doc, data)
        
        # Save document
        bpin_clean = sanitize_filename(data.get('bpin', 'DRAFT'))
        filename = f"Analisis_Sector_{bpin_clean}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def build_from_result(self, result: dict) -> str:
        """Build document from generator result"""
        doc_type = result.get("metadata", {}).get("tipo", "estudios_previos")
        content = result.get("documento_completo", "")
        data = result.get("metadata", {})
        
        if doc_type == "estudios_previos":
            return self.build_estudios_previos(content, data)
        else:
            return self.build_analisis_sector(content, data)
