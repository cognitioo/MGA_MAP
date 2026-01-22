"""
DTS (Documento Técnico de Soporte) Document Builder
Generates Word documents for Technical Support Documents
"""

import os
import re
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class DTSBuilder:
    """Builder for DTS (Documento Técnico de Soporte) Word documents"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.doc = None
    
    def build(self, data: dict, ai_content: dict, letterhead_file=None) -> str:
        """Build the complete DTS document"""
        self._has_letterhead = letterhead_file is not None
        
        if letterhead_file:
            self.doc = self._load_template(letterhead_file)
        else:
            self.doc = Document()
        
        self._apply_styles()
        
        # Main Title
        self._add_main_title(ai_content.get("titulo_proyecto", "DOCUMENTO TÉCNICO DE SOPORTE"))
        
        # 1. CONTRIBUCIÓN A LA POLÍTICA PÚBLICA
        self._add_section_header("1. CONTRIBUCIÓN A LA POLÍTICA PÚBLICA")
        
        self._add_subsection("1.1 CONTRIBUCIÓN AL PLAN NACIONAL DE DESARROLLO")
        self._add_content(ai_content.get("contribucion_plan_nacional", ""))
        
        self._add_subsection("1.2 PLAN DE DESARROLLO DEPARTAMENTAL O SECTORIAL")
        self._add_content(ai_content.get("contribucion_plan_departamental", ""))
        
        self._add_subsection("1.3 PLAN DE DESARROLLO DISTRITAL O MUNICIPAL")
        self._add_content(ai_content.get("contribucion_plan_municipal", ""))
        
        # 2. IDENTIFICACIÓN Y DESCRIPCIÓN DEL PROBLEMA
        self._add_section_header("2. IDENTIFICACIÓN Y DESCRIPCIÓN DEL PROBLEMA")
        
        self._add_subsection("2.1 DESCRIPCIÓN DE LA SITUACIÓN ACTUAL")
        self._add_content(ai_content.get("descripcion_situacion_actual", ""))
        
        # 3. PROBLEMA CENTRAL
        self._add_subsection("3.2 PROBLEMA CENTRAL")
        self._add_content(ai_content.get("problema_central", ""))
        
        self._add_subsection("3.3 MAGNITUD ACTUAL DEL PROBLEMA – INDICADORES DE REFERENCIA")
        self._add_content(ai_content.get("magnitud_problema", ""))
        
        self._add_subsection("3.4 DESCRIPCIÓN DE CAUSAS Y EFECTOS DIRECTOS E INDIRECTOS")
        self._add_content(f"**Causas directas:**<br>{ai_content.get('causas_directas', '')}<br><br>**Efectos directos:**<br>{ai_content.get('efectos_directos', '')}")
        
        # 4. ANÁLISIS DE PARTICIPANTES
        self._add_section_header("4. ANÁLISIS DE PARTICIPANTES")
        self._add_participants_table(ai_content.get("analisis_participantes", []))
        
        # 5. POBLACIONES AFECTADAS
        self._add_section_header("5. POBLACIONES AFECTADAS")
        self._add_content(f"**Tipo de población:** {ai_content.get('poblacion_afectada', 'Personas')}")
        self._add_content(f"**Localización:** {ai_content.get('localizacion', '')}")
        
        # 3.5 POBLACIÓN OBJETIVO DE LA INTERVENCIÓN
        self._add_subsection("3.5 POBLACIÓN OBJETIVO DE LA INTERVENCIÓN")
        self._add_content(ai_content.get("poblacion_objetivo", ""))
        
        # 4. OBJETIVO GENERAL Y ESPECÍFICO
        self._add_section_header("4. OBJETIVO GENERAL Y ESPECÍFICO")
        self._add_subsection("OBJETIVO GENERAL")
        self._add_content(ai_content.get("objetivo_general", ""))
        
        self._add_subsection("PROPÓSITO: PARA EL OBJETIVO GENERAL")
        self._add_indicators_table(ai_content.get("indicadores", []))
        
        # 5. ALTERNATIVAS DE LA SOLUCIÓN
        self._add_section_header("5. ALTERNATIVAS DE LA SOLUCIÓN")
        self._add_subsection("5.1 ALTERNATIVA DE LA SOLUCIÓN")
        self._add_content(ai_content.get("poblacion_objetivo", ""))
        
        self._add_subsection("5.2 DESARROLLO DE LA ALTERNATIVA DE SOLUCIÓN")
        self._add_content(ai_content.get("desarrollo_alternativa", ""))
        
        # 6. ESTUDIO DE NECESIDADES
        self._add_section_header("6. ESTUDIO DE NECESIDADES")
        
        self._add_subsection("6.1 BIEN O SERVICIO - ACUEDUCTO")
        self._add_content("Apoyo financiero para los usuarios de servicios públicos domiciliarios de acueducto en los estratos 1 y 2.")
        self._add_oferta_demanda_table(ai_content.get("tabla_subsidios_acueducto", []))
        
        self._add_subsection("ALCANTARILLADO")
        self._add_content("Apoyo financiero para los usuarios de servicios públicos domiciliarios de alcantarillado en los estratos 1 y 2.")
        self._add_oferta_demanda_table(ai_content.get("tabla_subsidios_alcantarillado", []))
        
        self._add_subsection("ASEO")
        self._add_content("Apoyo financiero para los usuarios de servicios públicos domiciliarios de aseo en los estratos 1 y 2.")
        self._add_oferta_demanda_table(ai_content.get("tabla_subsidios_aseo", []))
        
        # 7. CADENA DE VALOR DE LA ALTERNATIVA
        self._add_section_header("7. CADENA DE VALOR DE LA ALTERNATIVA")
        self._add_cadena_valor_table(ai_content.get("cadena_valor_productos", []))
        
        # 8. FUENTES DE FINANCIACIÓN
        self._add_section_header("8. FUENTES DE FINANCIACIÓN")
        self._add_financiacion_table(ai_content)
        
        self._add_content("<br>NOTA: Se anexa presupuesto.")
        
        # Signature
        self._add_signature(data)
        
        # Save document
        return self._save_document(data)
    
    def _load_template(self, letterhead_file):
        """Load template from uploaded file"""
        try:
            if hasattr(letterhead_file, 'read'):
                return Document(BytesIO(letterhead_file.read()))
            else:
                return Document(letterhead_file)
        except Exception:
            return Document()
    
    def _apply_styles(self):
        """Apply document styles"""
        for section in self.doc.sections:
            if not self._has_letterhead:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
    
    def _add_main_title(self, title):
        """Add main document title"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run("DOCUMENTO TÉCNICO SOPORTE DEL PROYECTO DE INVERSIÓN")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        
        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f'"{title}"')
        run2.bold = True
        run2.font.size = Pt(11)
        run2.font.name = 'Arial'
        
        self.doc.add_paragraph()
    
    def _add_section_header(self, title):
        """Add main section header"""
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    
    def _add_subsection(self, title):
        """Add subsection header"""
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    
    def _add_content(self, content):
        """Add content paragraph with formatting"""
        if not content:
            return
        
        # Replace <br> with actual line breaks
        content = content.replace("<br><br>", "\n\n").replace("<br>", "\n")
        
        for para_text in content.split("\n\n"):
            if not para_text.strip():
                continue
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            
            # Handle bold markers
            parts = re.split(r'\*\*(.*?)\*\*', para_text)
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Bold
                    run = p.add_run(part)
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                else:
                    for line in part.split("\n"):
                        run = p.add_run(line)
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                        if line != parts[-1]:
                            p.add_run("\n")
    
    def _add_participants_table(self, participants):
        """Add participants analysis table"""
        if not participants:
            return
        
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row
        headers = ["ACTOR", "ENTIDAD", "POSICIÓN", "TIPO"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            self._set_cell_shading(cell, "FFEB9C")
        
        # Data rows
        for participant in participants:
            row = table.add_row()
            row.cells[0].text = participant.get("actor", "")
            row.cells[1].text = participant.get("entidad", "")
            row.cells[2].text = participant.get("posicion", "")
            row.cells[3].text = participant.get("tipo", "")
            
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Arial'
        
        self.doc.add_paragraph()
    
    def _add_oferta_demanda_table(self, data):
        """Add offer/demand table"""
        if not data:
            return
        
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header
        headers = ["Año", "Oferta", "Demanda", "Déficit"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            self._set_cell_shading(cell, "FFEB9C")
        
        # Data
        for item in data:
            row = table.add_row()
            row.cells[0].text = str(item.get("ano", ""))
            row.cells[1].text = str(item.get("oferta", ""))
            row.cells[2].text = str(item.get("demanda", ""))
            row.cells[3].text = str(item.get("deficit", ""))
            
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in para.runs:
                        run.font.size = Pt(9)
        
        self.doc.add_paragraph()
    
    def _add_indicators_table(self, indicators):
        """Add indicators table"""
        if not indicators:
            return
        
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        headers = ["Indicador objetivo de desarrollo", "Meta"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            self._set_cell_shading(cell, "FFEB9C")
        
        for ind in indicators:
            row = table.add_row()
            row.cells[0].text = str(ind.get("objetivo", ""))
            row.cells[1].text = str(ind.get("meta", ""))
            
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
        
        self.doc.add_paragraph()
    
    def _add_cadena_valor_table(self, productos):
        """Add value chain table"""
        if not productos:
            return
        
        for prod in productos:
            if "producto" in prod:
                p = self.doc.add_paragraph()
                run = p.add_run(f"{prod.get('codigo', '')} {prod.get('producto', '')}")
                run.bold = True
                run.font.size = Pt(10)
                
                p2 = self.doc.add_paragraph()
                p2.add_run(f"Medida a través de: {prod.get('medida', '')}\n")
                p2.add_run(f"Cantidad: {prod.get('cantidad', '')}\n")
                p2.add_run(f"Costo: {prod.get('costo', '')}")
            
            elif "actividad" in prod:
                p = self.doc.add_paragraph()
                run = p.add_run(f"{prod.get('codigo', '')} {prod.get('actividad', '')}")
                run.font.size = Pt(10)
                
                p2 = self.doc.add_paragraph()
                p2.add_run(f"Etapa: {prod.get('etapa', '')}\n")
                p2.add_run(f"Costo: {prod.get('costo', '')}")
        
        self.doc.add_paragraph()
    
    def _add_financiacion_table(self, ai_content):
        """Add financing sources table"""
        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        table.rows[0].cells[0].text = ai_content.get("fuente_financiacion", "Recursos SGP - APSB")
        table.rows[0].cells[1].text = f"${ai_content.get('total_recursos', '0')}"
        table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        table.rows[1].cells[0].text = "Total del recurso:"
        table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[1].cells[1].text = f"${ai_content.get('total_recursos', '0')}"
        table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        table.rows[1].cells[1].paragraphs[0].runs[0].bold = True
        
        self.doc.add_paragraph()
    
    def _add_signature(self, data):
        """Add signature section"""
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        p = self.doc.add_paragraph()
        run = p.add_run(data.get("responsable", "").upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        
        p2 = self.doc.add_paragraph()
        p2.add_run(data.get("cargo", "Secretario de Planeación Municipal"))
        
        p3 = self.doc.add_paragraph()
        p3.add_run(f"Municipio de {data.get('municipio', '')}")
    
    def _set_cell_shading(self, cell, color):
        """Set cell background color"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        tcPr.append(shading)
    
    def _save_document(self, data):
        """Save the document to file"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Sanitize municipio - remove invalid characters for filenames
        municipio = data.get("municipio", "documento")
        # Remove tabs, newlines, and other invalid chars
        municipio = re.sub(r'[\t\n\r\\/:"*?<>|]', '', municipio)
        municipio = municipio.replace(" ", "_").strip()
        if not municipio:
            municipio = "documento"
        filename = f"DTS_{municipio}_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        self.doc.save(filepath)
        return filepath
