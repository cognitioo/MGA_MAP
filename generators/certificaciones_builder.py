"""
Certificaciones Document Builder
Generates Word documents for Presentation Letter and Certificates
"""

import os
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class CertificacionesBuilder:
    """Builder for Presentation and Certificates Word documents"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.doc = None
    
    def build(self, data: dict, ai_content: dict, letterhead_file=None) -> str:
        """Build the complete Certificaciones document with all certificates"""
        self._has_letterhead = letterhead_file is not None
        
        if letterhead_file:
            self.doc = self._load_template(letterhead_file)
        else:
            self.doc = Document()
        
        self._apply_styles()
        
        # 1. Carta de Presentación
        self._add_carta_presentacion(data, ai_content.get("carta_presentacion", {}))
        
        # Page break
        self.doc.add_page_break()
        
        # 2. Certificación Plan de Desarrollo
        self._add_certificacion(data, ai_content.get("cert_plan_desarrollo", {}))
        self.doc.add_page_break()
        
        # 3. Certificación Precios Unitarios
        self._add_certificacion(data, ai_content.get("cert_precios_unitarios", {}))
        self.doc.add_page_break()
        
        # 4. Certificación No Financiación
        self._add_certificacion(data, ai_content.get("cert_no_financiacion", {}))
        self.doc.add_page_break()
        
        # 5. Certificación Sostenibilidad
        self._add_certificacion(data, ai_content.get("cert_sostenibilidad", {}))
        self.doc.add_page_break()
        
        # 6. Certificación Viabilidad
        self._add_certificacion(data, ai_content.get("cert_viabilidad", {}))
        self.doc.add_page_break()
        
        # 7. Certificación Localización
        self._add_certificacion(data, ai_content.get("cert_localizacion", {}))
        self.doc.add_page_break()
        
        # 8. Certificación Normas Técnicas
        self._add_certificacion(data, ai_content.get("cert_normas_tecnicas", {}))
        
        return self._save_document(data)
    
    def _load_template(self, letterhead_file):
        """Load template from uploaded file"""
        try:
            if hasattr(letterhead_file, 'read'):
                return Document(BytesIO(letterhead_file.read()))
            else:
                return Document(letterhead_file)
        except Exception:
            return Document()
    
    def _apply_styles(self):
        """Apply document styles"""
        for section in self.doc.sections:
            if not self._has_letterhead:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
    
    def _add_carta_presentacion(self, data: dict, content: dict):
        """Add presentation letter"""
        # Date
        fecha = data.get("fecha", datetime.now().strftime("%d de %B de %Y"))
        p_fecha = self.doc.add_paragraph()
        p_fecha.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_fecha.add_run(f"{data.get('municipio', '')}, {fecha}")
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        
        self.doc.add_paragraph()
        
        # Destinatario
        destinatario = content.get("destinatario", "").replace("\\n", "\n")
        for line in destinatario.split("\n"):
            p = self.doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            if "Doctor" in line or "Alcalde" in line:
                pass
            else:
                run.bold = True
        
        self.doc.add_paragraph()
        
        # Referencia
        p_ref = self.doc.add_paragraph()
        run_ref = p_ref.add_run("Ref: ")
        run_ref.font.size = Pt(11)
        run_ref2 = p_ref.add_run(content.get("referencia", "Presentación Proyecto"))
        run_ref2.bold = True
        run_ref2.font.size = Pt(11)
        
        self.doc.add_paragraph()
        
        # Cuerpo
        cuerpo = content.get("cuerpo", "").replace("\\n", "\n")
        for para in cuerpo.split("\n\n"):
            if not para.strip():
                continue
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            
            for line in para.split("\n"):
                run = p.add_run(line + "\n")
                run.font.size = Pt(11)
                run.font.name = 'Arial'
        
        self.doc.add_paragraph()
        
        # Despedida
        p_desp = self.doc.add_paragraph()
        run = p_desp.add_run(content.get("despedida", "Agradeciendo su atención,"))
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        
        # Firma
        self._add_signature(data)
    
    def _add_certificacion(self, data: dict, content: dict):
        """Add a certification document"""
        if not content:
            return
        
        # Firma header (arriba a la derecha)
        p_header = self.doc.add_paragraph()
        p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_name = p_header.add_run(data.get("responsable", "").upper())
        run_name.bold = True
        run_name.font.size = Pt(11)
        run_name.font.name = 'Arial'
        
        p_cargo_header = self.doc.add_paragraph()
        p_cargo_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_cargo = p_cargo_header.add_run(data.get("cargo", "SECRETARIO DE PLANEACIÓN").upper())
        run_cargo.font.size = Pt(10)
        run_cargo.font.name = 'Arial'
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Título
        titulo = content.get("titulo", "").replace("\\n", "\n")
        for line in titulo.split("\n"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # CERTIFICA
        p_cert = self.doc.add_paragraph()
        p_cert.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_cert.add_run(content.get("encabezado", "CERTIFICA"))
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        
        self.doc.add_paragraph()
        
        # Contenido
        contenido = content.get("contenido", "").replace("\\n", "\n")
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(contenido)
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        
        self.doc.add_paragraph()
        
        # Fecha de expedición
        fecha_exp = content.get("fecha_expedicion", "")
        if fecha_exp:
            # Replace placeholders
            now = datetime.now()
            fecha_exp = fecha_exp.replace("{dia}", str(now.day))
            fecha_exp = fecha_exp.replace("{mes}", self._get_month_name(now.month))
            fecha_exp = fecha_exp.replace("{ano}", str(now.year))
            
            p = self.doc.add_paragraph()
            run = p.add_run(fecha_exp)
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Firma al final
        self._add_signature(data)
    
    def _add_signature(self, data: dict):
        """Add signature section"""
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        p_name = self.doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_name.add_run(data.get("responsable", "").upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        
        p_cargo = self.doc.add_paragraph()
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_cargo.add_run(data.get("cargo", "Secretario de Planeación").upper())
        run.font.size = Pt(10)
        run.font.name = 'Arial'
    
    def _get_month_name(self, month: int) -> str:
        """Get Spanish month name"""
        months = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
        ]
        return months[month - 1]
    
    def _save_document(self, data: dict) -> str:
        """Save the document to file"""
        import re
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Sanitize municipio - remove invalid characters for filenames
        municipio = data.get("municipio", "documento")
        municipio = re.sub(r'[\t\n\r\\/:"*?<>|]', '', municipio)
        municipio = municipio.replace(" ", "_").strip()
        if not municipio:
            municipio = "documento"
        filename = f"Certificaciones_{municipio}_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        self.doc.save(filepath)
        return filepath
