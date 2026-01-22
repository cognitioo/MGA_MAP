"""
Análisis del Sector Document Builder
Builds Word documents matching the MGA Sector Analysis template
- Multiple economic sections with subsections
- Matplotlib graphs for PIB and SMLMV
- Risk tables and economic data tables
"""

import os
import re
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Optional: matplotlib for graphs
try:
    import matplotlib
    matplotlib.use('Agg')  # Non-interactive backend
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

# Colors matching client template
GRAY_HEADER = "BFBFBF"  # For table headers and riesgos
GOLD_HEADER = "F4C430"  # Yellow/gold for main header table labels
BLACK = "000000"

# Spanish month names
SPANISH_MONTHS = {
    1: "ENERO", 2: "FEBRERO", 3: "MARZO", 4: "ABRIL",
    5: "MAYO", 6: "JUNIO", 7: "JULIO", 8: "AGOSTO",
    9: "SEPTIEMBRE", 10: "OCTUBRE", 11: "NOVIEMBRE", 12: "DICIEMBRE"
}


def sanitize_filename(filename: str) -> str:
    """Remove invalid characters from filename"""
    invalid_chars = r'[<>:"/\\|?*\t\n\r]'
    sanitized = re.sub(invalid_chars, '', filename)
    sanitized = re.sub(r'[\s_]+', '_', sanitized)
    return sanitized.strip('_')


class AnalisisSectorBuilder:
    """Builds Análisis del Sector Word documents"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.doc = None
    
    def build(self, data: dict, ai_content: dict, letterhead_file=None, section_toggles=None) -> str:
        """Build the complete Análisis del Sector document
        
        Args:
            data: Form data from user input
            ai_content: AI-generated content for each section
            letterhead_file: Optional .docx file with letterhead
            section_toggles: Dict of section_name -> bool to enable/disable sections
        """
        # Default all sections to enabled if no toggles provided
        if section_toggles is None:
            section_toggles = {}
        self.section_toggles = section_toggles
        
        # Helper function to check if section is enabled
        def is_enabled(section_name):
            return self.section_toggles.get(section_name, True)
        
        # Track if letterhead is used (to preserve margins in _apply_styles)
        self._has_letterhead = letterhead_file is not None
        
        # If letterhead provided, use it as base template (preserves images/logos)
        if letterhead_file:
            self.doc = self._load_template(letterhead_file)
        else:
            self.doc = Document()
        
        self._apply_styles()
        
        # 1. Main Title
        self._add_main_title()
        
        # 2. Header Table (ENTIDAD, DEPENDENCIA, FUNCIONARIO, MODALIDAD)
        self._add_header_table(data)
        
        # 3. OBJETO section (bold inline, no gray box) - matches client template
        if is_enabled("objeto"):
            self._add_subsection_with_content(
                "OBJETO",
                ai_content.get("objeto", "")
            )
        
        # ALCANCE section (bold inline, no gray box)
        if is_enabled("alcance"):
            self._add_subsection_with_content(
                "ALCANCE",
                ai_content.get("alcance", "")
            )
        
        # DESCRIPCIÓN DE LA NECESIDAD
        if is_enabled("descripcion_necesidad"):
            self._add_subsection_with_content(
                "DESCRIPCIÓN DE LA NECESIDAD",
                ai_content.get("descripcion_necesidad", "")
            )
        
        # INTRODUCCIÓN
        if is_enabled("introduccion"):
            self._add_subsection_with_content(
                "INTRODUCCIÓN",
                ai_content.get("introduccion", "")
            )
        
        # DEFINICIONES
        if is_enabled("definiciones"):
            self._add_subsection_with_content(
                "DEFINICIONES",
                ai_content.get("definiciones", "")
            )
        
        # 1. DESARROLLO DEL ESTUDIO DEL SECTOR (main section - bold, slightly larger)
        if is_enabled("desarrollo_estudio"):
            self._add_main_section_header("1. DESARROLLO DEL ESTUDIO DEL SECTOR")
            
            # 1.1 Banco de Programas y Proyectos
            self._add_subsection_with_content(
                "1.1. Banco de Programas y Proyectos",
                ai_content.get("banco_programas", "")
            )
            
            # 1.2 Consideraciones para la realización del Estudio del Sector
            self._add_subsection_with_content(
                "1.2. Consideraciones para la realización del Estudio del Sector",
                ai_content.get("consideraciones_estudio", "")
            )
            
            # 1.3 Preparación del Estudio del Sector
            self._add_subsection_with_content(
                "1.3. Preparación del Estudio del Sector",
                ai_content.get("preparacion_estudio", "")
            )
            
            # 1.4 Estructura del Estudio del Sector
            self._add_subsection_with_content(
                "1.4. Estructura del Estudio del Sector",
                ai_content.get("estructura_estudio", "")
            )
            
            # 1.4.1 Aspectos generales del mercado
            self._add_subsection_with_content(
                "1.4.1. Aspectos generales del mercado",
                ai_content.get("aspectos_mercado", "")
            )
            
            # 1.4.2 Comportamiento del gasto histórico
            self._add_subsection_with_content(
                "1.4.2. Comportamiento del gasto histórico",
                ai_content.get("gasto_historico", "")
            )
            
            # 1.4.3 Estudio de la oferta
            self._add_subsection_with_content(
                "1.4.3. Estudio de la oferta",
                ai_content.get("estudio_oferta", "")
            )
            
            # 1.4.4 Estudio de mercado
            self._add_subsection_with_content(
                "1.4.4. Estudio de mercado",
                ai_content.get("estudio_mercado", "")
            )
            
            # 1.4.5 Objeto del contrato
            self._add_subsection_with_content(
                "1.4.5. Objeto del contrato",
                ai_content.get("objeto_contrato", "")
            )
            
            # 1.4.6 Sector económico de la necesidad
            self._add_subsection_with_content(
                "1.4.6. Sector económico de la necesidad",
                ai_content.get("sector_economico", "")
            )
        
        # 1.5 ANÁLISIS DEL SECTOR
        if is_enabled("analisis_sector"):
            self._add_subsection_with_content(
                "1.5. ANÁLISIS DEL SECTOR",
                ai_content.get("analisis_sector_intro", "")
            )
            
            # 1.5.1 Descripción del sector económico
            self._add_subsection_with_content(
                "1.5.1. Descripción del sector económico",
                ai_content.get("descripcion_sector_economico", "")
            )
            
            # 1.5.2 Sector terciario o de servicios
            self._add_subsection_with_content(
                "1.5.2. Sector terciario o de servicios",
                ai_content.get("sector_terciario", "")
            )
            
            # 1.5.3 Comportamiento de la economía nacional
            self._add_subsection_with_content(
                "1.5.3. Comportamiento de la economía nacional en el primer trimestre de 2025",
                ai_content.get("comportamiento_economia", "")
            )
            
            # Add PIB graph if matplotlib available AND enabled
            if MATPLOTLIB_AVAILABLE and is_enabled("grafico_pib"):
                self._add_pib_graph(data)
            
            # 1.5.4 Variables económicas - SMLMV
            self._add_subsection_with_content(
                "1.5.4. Variables económicas",
                ai_content.get("variables_economicas", "")
            )
            
            # Add SMLMV table if enabled
            if is_enabled("tabla_smlmv"):
                self._add_smlmv_table()
            
            # 1.5.5 Relevancia para el PSMV
            self._add_subsection_with_content(
                "1.5.5. Relevancia para el PSMV",
                ai_content.get("relevancia_psmv", "")
            )
            
            # 1.5.6 Perspectivas Legales Del Sector
            self._add_subsection_with_content(
                "1.5.6. Perspectivas Legales Del Sector",
                ai_content.get("perspectivas_legales", "")
            )
        
        # 1.5.7 Riesgos - with table
        if is_enabled("riesgos"):
            self._add_subsection_with_content(
                "1.5.7. Riesgos",
                ai_content.get("riesgos_texto", "")
            )
            self._add_riesgos_table(ai_content.get("riesgos", []))
        
        # 2. ESTUDIOS DEL SECTOR EN LOS PROCESOS DE CONTRATACIÓN
        if is_enabled("estudios_contratacion"):
            self._add_main_section_header("2. ESTUDIOS DEL SECTOR EN LOS PROCESOS DE CONTRATACIÓN")
            self._add_subsection_with_content(
                "",  # No separate title, content follows header
                ai_content.get("estudios_sector_contratacion", "")
            )
            
            # 2.1 Contratación directa
            self._add_subsection_with_content(
                "2.1. Contratación directa",
                ai_content.get("contratacion_directa", "")
            )
            
            # 2.2 Mínima cuantía
            self._add_subsection_with_content(
                "2.2. Mínima cuantía",
                ai_content.get("minima_cuantia", "")
            )
            
            # 3. ANÁLISIS DEL SECTOR PARA CRITERIOS DIFERENCIALES DE MIPYME Y EMPRESAS DE MUJERES
            self._add_main_section_header("3. ANÁLISIS DEL SECTOR PARA CRITERIOS DIFERENCIALES DE MIPYME Y EMPRESAS DE MUJERES")
            self._add_subsection_with_content(
                "",
                ai_content.get("analisis_mga", "")
            )
        
        # 4. RECOMENDACIONES PARA EL MANEJO DE DATOS Y ANÁLISIS ESTADÍSTICO
        if is_enabled("recomendaciones"):
            self._add_main_section_header("4. RECOMENDACIONES PARA EL MANEJO DE DATOS Y ANÁLISIS ESTADÍSTICO")
            self._add_subsection_with_content(
                "",
                ai_content.get("recomendaciones", "")
            )
            
            # 4.1 ¿Cuándo usar estadística descriptiva?
            self._add_subsection_with_content(
                "4.1. ¿Cuándo usar estadística descriptiva?",
                ai_content.get("estadistica_descriptiva", "")
            )
            
            # 4.2 Preparación de datos
            self._add_subsection_with_content(
                "4.2. Preparación de datos",
                ai_content.get("preparacion_datos", "")
            )
            
            # 4.3 Análisis gráfico
            self._add_subsection_with_content(
                "4.3. Análisis gráfico",
                ai_content.get("analisis_grafico", "")
            )
        
        # 5. FUENTES DE INFORMACIÓN
        if is_enabled("fuentes"):
            self._add_main_section_header("5. FUENTES DE INFORMACIÓN")
            self._add_subsection_with_content(
                "",
                ai_content.get("fuentes_informacion", "")
            )
            
            # 6. HERRAMIENTAS DE BÚSQUEDA DE INFORMACIÓN
            self._add_main_section_header("6. HERRAMIENTAS DE BÚSQUEDA DE INFORMACIÓN")
            self._add_subsection_with_content(
                "",
                ai_content.get("herramientas_busqueda", "")
            )
        
        # 7. Estimación y justificación del valor del contrato
        if is_enabled("estimacion_valor"):
            self._add_main_section_header("7. Estimación y justificación del valor del contrato")
            self._add_subsection_with_content(
                "",
                ai_content.get("estimacion_valor", "")
            )
        
        # Signature Section
        self._add_signature_section(data)
        
        # Save document
        bpin = sanitize_filename(data.get('bpin', 'DRAFT'))
        filename = f"Analisis_Sector_{bpin}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        self.doc.save(filepath)
        
        return filepath
    
    def _apply_styles(self):
        """Apply document-wide styles (preserves letterhead margins)"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)
        
        # Only set margins if no letterhead was loaded (letterhead margins set in _load_template)
        if not getattr(self, '_has_letterhead', False):
            for section in self.doc.sections:
                section.top_margin = Cm(1.5)
                section.bottom_margin = Cm(1.5)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
    
    def _load_template(self, letterhead_file):
        """
        Load the letterhead template as the base document.
        This preserves all images, logos, and graphics in headers/footers.
        """
        import io
        
        try:
            if hasattr(letterhead_file, 'read'):
                template_bytes = letterhead_file.read()
                letterhead_file.seek(0)
                doc = Document(io.BytesIO(template_bytes))
            else:
                doc = Document(letterhead_file)
            
            # Clear all body content (keep headers/footers)
            for element in doc.element.body[:]:
                if not element.tag.endswith('sectPr'):
                    doc.element.body.remove(element)
            
            # Set margins to leave room for header/footer graphics on ALL pages
            for section in doc.sections:
                section.different_first_page_header_footer = False
                # Set proper margins for content area (leaves room for letterhead)
                section.top_margin = Cm(4)        # Distance from top of page to content
                section.bottom_margin = Cm(3.5)   # Distance from bottom of page to content
                # header_distance = space from page edge to header
                # footer_distance = space from page edge to footer
                section.header_distance = Cm(1)   # Keep header at top
                section.footer_distance = Cm(1)   # Keep footer at bottom
            
            return doc
            
        except Exception as e:
            print(f"Warning: Could not load letterhead template: {e}")
            return Document()
    
    def _add_main_title(self):
        """Add main document title"""
        title = "ANÁLISIS DEL SECTOR ECONÓMICO Y DE LOS OFERENTES POR PARTE DE LAS ENTIDADES ESTATALES"
        
        para = self.doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(24)  # Push down from letterhead header
        para.paragraph_format.space_after = Pt(6)
    
    def _add_header_table(self, data: dict):
        """Add header table with gold/yellow left column - matches client template"""
        table = self.doc.add_table(rows=4, cols=2)
        table.autofit = True
        self._set_table_borders(table)
        
        now = datetime.now()
        spanish_date = f"{SPANISH_MONTHS[now.month]} DE {now.year}"
        
        # Client template uses these labels
        rows_data = [
            ("ENTIDAD QUE CONTRATA", data.get("entidad", f"MUNICIPIO DE {data.get('municipio', '').upper()}")),
            ("DEPENDENCIA QUE POSEE LA NECESIDAD", data.get("dependencia", "SECRETARÍA DE PLANEACIÓN")),
            ("FUNCIONARIO RESPONSABLE", data.get("responsable", "")),
            ("MODALIDAD DEL PROCESO", data.get("proceso", "CONVENIO INTERADMINISTRATIVO"))
        ]
        
        for i, (label, value) in enumerate(rows_data):
            cell_label = table.rows[i].cells[0]
            cell_value = table.rows[i].cells[1]
            
            cell_label.text = label
            cell_value.text = value
            
            # Keep cells white (no shading) - matches client template
            
            for para in cell_label.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(8)
            
            for para in cell_value.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(8)
        
        self.doc.add_paragraph()  # Space after header
    
    def _add_main_section_header(self, title: str):
        """Add a main section header (bold, larger font, no gray box) - matches client template"""
        para = self.doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(10)  # Slightly larger than subsections
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    
    def _add_section_header(self, title: str):
        """Add a section header (gray background)"""
        table = self.doc.add_table(rows=1, cols=1)
        self._set_table_borders(table)
        
        cell = table.rows[0].cells[0]
        cell.text = title
        self._shade_cell(cell, GRAY_HEADER)
        
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    
    def _add_section_with_content(self, title: str, content: str):
        """Add a section with gray header and text content"""
        table = self.doc.add_table(rows=2, cols=1)
        self._set_table_borders(table)
        
        # Header
        header_cell = table.rows[0].cells[0]
        header_cell.text = title
        self._shade_cell(header_cell, GRAY_HEADER)
        for para in header_cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        
        # Content
        content_cell = table.rows[1].cells[0]
        content_cell.paragraphs[0].clear()
        self._add_formatted_text(content_cell, content)
    
    def _add_subsection_with_content(self, title: str, content: str):
        """Add a subsection with bold title (no box) and content"""
        # Title paragraph
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(9)
        title_para.paragraph_format.space_before = Pt(6)
        title_para.paragraph_format.space_after = Pt(3)
        
        # Content paragraph
        if content:
            content_para = self.doc.add_paragraph()
            self._add_formatted_text_to_para(content_para, content)
    
    def _add_pib_graph(self, data: dict):
        """Generate and add PIB evolution graph"""
        if not MATPLOTLIB_AVAILABLE:
            return
        
        # Sample PIB data (2018-2025)
        years = [2018, 2019, 2020, 2021, 2022, 2023, 2024, 2025]
        pib_values = [1.1, 3.2, -7.0, 10.8, 7.3, 0.6, 2.7, 2.5]  # % growth
        
        fig, ax = plt.subplots(figsize=(8, 4))
        ax.bar(years, pib_values, color=['#4472C4' if v >= 0 else '#C0504D' for v in pib_values])
        ax.axhline(y=0, color='black', linewidth=0.5)
        ax.set_xlabel('Año', fontsize=9)
        ax.set_ylabel('Crecimiento PIB (%)', fontsize=9)
        ax.set_title('Gráfico 1. Producto Interno Bruto (PIB)\nTasa de crecimiento anual en volumen', fontsize=10)
        ax.set_xticks(years)
        
        # Add value labels
        for i, v in enumerate(pib_values):
            ax.text(years[i], v + 0.3, f'{v}%', ha='center', fontsize=8)
        
        plt.tight_layout()
        
        # Save to buffer and add to document
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        
        self.doc.add_picture(buf, width=Inches(5.5))
        
        # Caption
        caption = self.doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run("Fuente: DANE")
        run.italic = True
        run.font.size = Pt(8)
    
    def _add_smlmv_graph(self):
        """Generate and add SMLMV evolution graph"""
        if not MATPLOTLIB_AVAILABLE:
            return
        
        # SMLMV historical data (COP)
        years = list(range(2000, 2026))
        smlmv = [260100, 286000, 309000, 332000, 358000, 381500, 408000, 433700,
                 461500, 496900, 515000, 535600, 566700, 589500, 616000, 644350,
                 689454, 737717, 781242, 828116, 877803, 908526, 1000000, 1160000,
                 1300000, 1423500]
        
        fig, ax = plt.subplots(figsize=(8, 4))
        ax.plot(years, [s/1000000 for s in smlmv], marker='o', markersize=4, 
                color='#4472C4', linewidth=2)
        ax.fill_between(years, [s/1000000 for s in smlmv], alpha=0.3, color='#4472C4')
        ax.set_xlabel('Año', fontsize=9)
        ax.set_ylabel('Millones COP', fontsize=9)
        ax.set_title('Evolución del Salario Mínimo en Colombia (2000-2025)', fontsize=10)
        ax.set_xlim(2000, 2025)
        ax.grid(True, alpha=0.3)
        
        plt.tight_layout()
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        
        self.doc.add_picture(buf, width=Inches(5.5))
        
        caption = self.doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run("Fuente: Ministerio de Trabajo")
        run.italic = True
        run.font.size = Pt(8)
    
    def _add_smlmv_table(self):
        """Add SMLMV historical data table - 4 columns matching client template"""
        # Client template has: Año, Salario mínimo mensual, Auxilio de Transporte, Normatividad Decreto
        # Historical data from 2000 to 2025
        data_rows = [
            ("2025", "1,423,500", "200,000", "1572 de dic 24 de 2024"),
            ("2024", "1,300,000", "162,000", "2292 de dic 29 2023"),
            ("2023", "1,160,000", "140,606", "2623 de dic 23 2022"),
            ("2022", "1,000,000", "117,172", "1724 de dic 15 de 2021"),
            ("2021", "908,526", "106,454", "1788 de dic 29 de 2020"),
            ("2020", "877,803", "102,854", "2360 de dic 26 de 2019"),
            ("2019", "828,116", "97,032", "2451 de dic 27 de 2018"),
            ("2018", "781,242", "88,211", "2269 de dic 30 de 2017"),
            ("2017", "737,717", "83,140", "2209 de dic 30 de 2016"),
            ("2016", "689,455", "77,700", "2552 de dic 30 de 2015"),
            ("2015", "644,350", "74,000", "2731 de dic 30 de 2014"),
            ("2014", "616,000", "72,000", "3068 de dic 30 de 2013"),
            ("2013", "589,500", "70,500", "2738 de dic 28 de 2012"),
            ("2012", "566,700", "67,800", "4919 de dic 26 de 2011"),
            ("2011", "535,600", "63,600", "033 de enero 11 de 2011"),
            ("2010", "515,000", "61,500", "5053 de dic 30 de 2009"),
            ("2009", "496,900", "59,300", "4868 de dic 30 de 2008"),
            ("2008", "461,500", "55,000", "4965 de dic 27 de 2007"),
            ("2007", "433,700", "50,800", "4580 de dic 27 de 2006"),
            ("2006", "408,000", "47,700", "4686 de dic 21 de 2005"),
            ("2005", "381,500", "44,500", "4360 de dic 22 de 2004"),
            ("2004", "358,000", "41,600", "3770 de dic 26 de 2003"),
            ("2003", "332,000", "37,500", "3232 de dic 27 de 2002"),
            ("2002", "309,000", "34,000", "2910 de dic 31 de 2001"),
            ("2001", "286,000", "30,000", "2579 de dic 13 de 2000"),
            ("2000", "260,100", "26,413", "2647 de dic 28 de 1999"),
        ]
        
        num_rows = len(data_rows) + 2  # +1 for header, +1 for "Fuente" row
        table = self.doc.add_table(rows=num_rows, cols=4)
        self._set_table_borders(table)
        
        # Headers
        headers = ["Año", "Salario mínimo mensual", "Auxilio de Transporte", "Normatividad Decreto"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            self._shade_cell(cell, GRAY_HEADER)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(8)
        
        # Data rows
        for row_idx, (year, salary, auxilio, decreto) in enumerate(data_rows):
            row = table.rows[row_idx + 1]
            row.cells[0].text = year
            row.cells[1].text = salary
            row.cells[2].text = auxilio
            row.cells[3].text = decreto
            
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        run.font.size = Pt(7)  # Smaller font for many rows
        
        # Source row (merge all cells)
        source_row = table.rows[-1]
        source_row.cells[0].merge(source_row.cells[1])
        source_row.cells[0].merge(source_row.cells[2])
        source_row.cells[0].merge(source_row.cells[3])
        source_cell = source_row.cells[0]
        source_cell.text = "Fuente: DANE."
        for para in source_cell.paragraphs:
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                run.italic = True
                run.font.size = Pt(8)
    
    def _add_riesgos_table(self, riesgos: list):
        """Add risk matrix table"""
        if not riesgos:
            riesgos = [
                {"riesgo": "Demora en recolección de información", "descripcion": "Dificultades de acceso", "probabilidad": "Media", "mitigacion": "Coordinar cronogramas"},
                {"riesgo": "Baja participación comunitaria", "descripcion": "Falta de interés", "probabilidad": "Media", "mitigacion": "Campaña de difusión"},
                {"riesgo": "Inconsistencias en la información", "descripcion": "Datos incompletos", "probabilidad": "Baja", "mitigacion": "Validación cruzada"}
            ]
        
        num_rows = len(riesgos) + 1
        table = self.doc.add_table(rows=num_rows, cols=4)
        self._set_table_borders(table)
        
        # Header row
        headers = ["Riesgo identificado", "Descripción", "Probabilidad", "Medida de mitigación"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            self._shade_cell(cell, GRAY_HEADER)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(8)
        
        # Data rows
        for row_idx, r in enumerate(riesgos):
            if isinstance(r, dict):
                row = table.rows[row_idx + 1]
                row.cells[0].text = r.get("riesgo", "")
                row.cells[1].text = r.get("descripcion", "")
                row.cells[2].text = r.get("probabilidad", "")
                row.cells[3].text = r.get("mitigacion", "")
                
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.space_after = Pt(0)
                        for run in para.runs:
                            run.font.size = Pt(8)
    
    def _add_signature_section(self, data: dict):
        """Add signature section"""
        self.doc.add_paragraph()  # Space before
        
        para = self.doc.add_paragraph()
        para.add_run("Atentamente,")
        para.paragraph_format.space_after = Pt(30)
        
        name_para = self.doc.add_paragraph()
        name_run = name_para.add_run(data.get("responsable", "").upper())
        name_run.bold = True
        name_run.font.size = Pt(10)
        
        cargo_para = self.doc.add_paragraph()
        cargo_run = cargo_para.add_run(f"CARGO: {data.get('cargo', 'Secretario de Planeación Municipal')}")
        cargo_run.font.size = Pt(10)
    
    def _add_formatted_text(self, cell, text: str):
        """Add formatted text to a cell"""
        if not text:
            return
        
        text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
        paragraphs = text.split('\n')
        
        first = True
        for p_text in paragraphs:
            p_text = p_text.strip()
            if not p_text:
                continue
            
            if first:
                para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                first = False
            else:
                para = cell.add_paragraph()
            
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(0)
            
            if p_text.startswith('• ') or p_text.startswith('- ') or p_text.startswith('* '):
                para.paragraph_format.left_indent = Cm(0.5)
                p_text = "• " + p_text[2:]
            
            # Handle **bold**
            pattern = r'(\*\*(.+?)\*\*|([^*]+))'
            for match in re.finditer(pattern, p_text):
                m = match.group(0)
                if m.startswith('**') and m.endswith('**'):
                    run = para.add_run(m[2:-2])
                    run.bold = True
                    run.font.size = Pt(9)
                else:
                    run = para.add_run(m)
                    run.font.size = Pt(9)
    
    def _add_formatted_text_to_para(self, para, text: str):
        """Add formatted text directly to a paragraph"""
        if not text:
            return
        
        text = re.sub(r'<br\s*/?>', ' ', text, flags=re.IGNORECASE)
        
        para.paragraph_format.space_after = Pt(3)
        
        pattern = r'(\*\*(.+?)\*\*|([^*]+))'
        for match in re.finditer(pattern, text):
            m = match.group(0)
            if m.startswith('**') and m.endswith('**'):
                run = para.add_run(m[2:-2])
                run.bold = True
                run.font.size = Pt(9)
            else:
                run = para.add_run(m)
                run.font.size = Pt(9)
    
    def _set_table_borders(self, table):
        """Set borders on table"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), BLACK)
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
    
    def _shade_cell(self, cell, color: str):
        """Add background color to cell"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
