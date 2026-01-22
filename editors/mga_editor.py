"""
MGA Document Editor
Edits existing MGA documents (PDF/DOCX) based on user prompts using AI.
Preserves document structure while making targeted text changes.
"""

import os
import re
import json
from io import BytesIO
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime

# PDF editing with PyMuPDF
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

# DOCX editing
try:
    from docx import Document
except ImportError:
    Document = None


class MGAEditor:
    """
    Editor for MGA documents that applies AI-driven edits to existing files.
    Supports both PDF (direct text replacement) and DOCX formats.
    """
    
    def __init__(self, llm=None, output_dir: str = "output"):
        """
        Initialize the MGA Editor.
        
        Args:
            llm: LangChain LLM instance for AI-powered edit analysis
            output_dir: Directory to save edited documents
        """
        self.llm = llm
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    # ═══════════════════════════════════════════════════════════════
    # DOCUMENT READING
    # ═══════════════════════════════════════════════════════════════
    
    def read_document(self, file, file_type: str) -> Dict[str, Any]:
        """
        Read and extract content from uploaded document.
        
        Args:
            file: Uploaded file object
            file_type: File extension (.pdf or .docx)
            
        Returns:
            Dictionary with:
                - full_text: Complete document text
                - pages: List of page contents (for PDFs)
                - page_count: Total number of pages
                - structure: Document structure info
        """
        if hasattr(file, 'read'):
            content = file.read()
            file.seek(0)
        else:
            with open(file, 'rb') as f:
                content = f.read()
        
        file_type = file_type.lower().replace('.', '')
        
        if file_type == 'pdf':
            return self._read_pdf(content)
        elif file_type == 'docx':
            return self._read_docx(content)
        else:
            return {"error": f"Unsupported file type: {file_type}"}
    
    def _read_pdf(self, content: bytes) -> Dict[str, Any]:
        """Read PDF and extract text by page"""
        if not fitz:
            return {"error": "PyMuPDF not installed"}
        
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            pages = []
            full_text = ""
            
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                pages.append({
                    "page_number": page_num + 1,
                    "text": page_text,
                    "char_count": len(page_text)
                })
                full_text += f"\n\n=== PÁGINA {page_num + 1} ===\n{page_text}"
            
            result = {
                "full_text": full_text,
                "pages": pages,
                "page_count": len(doc),
                "file_type": "pdf",
                "raw_doc": content  # Store for later editing
            }
            doc.close()
            return result
            
        except Exception as e:
            return {"error": f"PDF read error: {str(e)}"}
    
    def _read_docx(self, content: bytes) -> Dict[str, Any]:
        """Read DOCX and extract text with structure"""
        if not Document:
            return {"error": "python-docx not installed"}
        
        try:
            doc = Document(BytesIO(content))
            
            paragraphs = []
            tables_text = []
            full_text = ""
            
            # Extract paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs.append({
                        "index": i,
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal"
                    })
                    full_text += para.text + "\n"
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_text = f"\n=== TABLA {table_idx + 1} ===\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    table_text += row_text + "\n"
                tables_text.append(table_text)
                full_text += table_text
            
            # Estimate page count (rough: ~3000 chars per page)
            estimated_pages = max(1, len(full_text) // 3000)
            
            return {
                "full_text": full_text,
                "paragraphs": paragraphs,
                "tables_count": len(doc.tables),
                "page_count": estimated_pages,
                "file_type": "docx",
                "raw_doc": content
            }
            
        except Exception as e:
            return {"error": f"DOCX read error: {str(e)}"}
    
    # ═══════════════════════════════════════════════════════════════
    # AI EDIT ANALYSIS
    # ═══════════════════════════════════════════════════════════════
    
    def analyze_edit_request(
        self, 
        doc_content: Dict[str, Any], 
        user_prompt: str,
        target_pages: Optional[List[int]] = None
    ) -> List[Dict[str, Any]]:
        """
        Use AI to analyze the document and determine what edits to make.
        
        Args:
            doc_content: Document content from read_document()
            user_prompt: User's edit instructions
            target_pages: Optional list of page numbers to focus on
            
        Returns:
            List of edit operations:
            [
                {
                    "page": 1,
                    "original_text": "Valor: $100,000",
                    "new_text": "Valor: $500,000",
                    "edit_type": "replace",
                    "reason": "User requested value update"
                }
            ]
        """
        if not self.llm:
            return [{"error": "No LLM configured for AI analysis"}]
        
        from langchain_core.prompts import ChatPromptTemplate
        from langchain_core.output_parsers import StrOutputParser
        
        # Prepare document text (focus on target pages if specified)
        if target_pages and doc_content.get("pages"):
            doc_text = ""
            for page in doc_content["pages"]:
                if page["page_number"] in target_pages:
                    doc_text += f"\n=== PÁGINA {page['page_number']} ===\n{page['text']}"
        else:
            doc_text = doc_content.get("full_text", "")[:50000]  # Limit for API
        
        # Build the prompt
        prompt = ChatPromptTemplate.from_messages([
            ("system", """Eres un experto editor de documentos MGA colombianos.

Tu tarea es analizar el documento y las instrucciones del usuario para generar una lista de EDICIONES ESPECÍFICAS.

REGLAS:
1. Cada edición debe tener el texto EXACTO original y el texto nuevo
2. Los textos deben ser lo suficientemente específicos para encontrarlos en el documento
3. Solo sugiere cambios que el usuario solicitó
4. Mantén el formato y estructura del documento

Responde SOLO con JSON válido en este formato:
{{
    "edits": [
        {{
            "page": 1,
            "original_text": "texto exacto a reemplazar",
            "new_text": "texto nuevo",
            "edit_type": "replace",
            "reason": "explicación breve"
        }}
    ],
    "summary": "resumen de los cambios"
}}

Si no encuentras qué editar, responde:
{{
    "edits": [],
    "summary": "No se encontraron elementos para editar según las instrucciones."
}}"""),
            ("human", """INSTRUCCIONES DEL USUARIO:
{user_prompt}

DOCUMENTO A EDITAR:
{doc_text}

Genera las ediciones necesarias en formato JSON.""")
        ])
        
        try:
            chain = prompt | self.llm | StrOutputParser()
            response = chain.invoke({
                "user_prompt": user_prompt,
                "doc_text": doc_text
            })
            
            # Parse JSON response
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                result = json.loads(json_match.group(0))
                return result.get("edits", []), result.get("summary", "")
            
            return [], "No se pudo analizar la respuesta del AI"
            
        except Exception as e:
            return [], f"Error en análisis AI: {str(e)}"
    
    # ═══════════════════════════════════════════════════════════════
    # APPLY EDITS
    # ═══════════════════════════════════════════════════════════════
    
    def apply_edits_pdf(
        self, 
        content: bytes, 
        edits: List[Dict[str, Any]]
    ) -> Tuple[bytes, List[Dict[str, Any]]]:
        """
        Apply text replacements to PDF using PyMuPDF.
        
        Note: PyMuPDF can do text replacement but it's limited.
        For complex edits, consider converting to DOCX first.
        
        Args:
            content: Original PDF bytes
            edits: List of edit operations
            
        Returns:
            Tuple of (edited PDF bytes, list of applied edits)
        """
        if not fitz:
            return content, [{"error": "PyMuPDF not installed"}]
        
        applied_edits = []
        
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            
            for edit in edits:
                original_text = edit.get("original_text", "")
                new_text = edit.get("new_text", "")
                target_page = edit.get("page")
                
                if not original_text or not new_text:
                    continue
                
                # Search and replace in specific page or all pages
                pages_to_search = [doc[target_page - 1]] if target_page else doc
                
                for page in pages_to_search:
                    # Find text instances
                    text_instances = page.search_for(original_text)
                    
                    if text_instances:
                        for inst in text_instances:
                            # Add redaction annotation to remove old text
                            page.add_redact_annot(inst)
                        
                        # Apply redactions
                        page.apply_redactions()
                        
                        # Insert new text at first instance location
                        if text_instances:
                            rect = text_instances[0]
                            # Insert new text
                            page.insert_text(
                                (rect.x0, rect.y1),
                                new_text,
                                fontsize=10
                            )
                        
                        applied_edits.append({
                            "original": original_text[:50] + "...",
                            "new": new_text[:50] + "...",
                            "status": "applied"
                        })
                    else:
                        applied_edits.append({
                            "original": original_text[:50] + "...",
                            "status": "not_found"
                        })
            
            # Save to bytes
            output = BytesIO()
            doc.save(output)
            doc.close()
            
            return output.getvalue(), applied_edits
            
        except Exception as e:
            return content, [{"error": f"PDF edit error: {str(e)}"}]
    
    def apply_edits_docx(
        self, 
        content: bytes, 
        edits: List[Dict[str, Any]]
    ) -> Tuple[bytes, List[Dict[str, Any]]]:
        """
        Apply text replacements to DOCX using python-docx.
        
        Args:
            content: Original DOCX bytes
            edits: List of edit operations
            
        Returns:
            Tuple of (edited DOCX bytes, list of applied edits)
        """
        if not Document:
            return content, [{"error": "python-docx not installed"}]
        
        applied_edits = []
        
        try:
            doc = Document(BytesIO(content))
            
            for edit in edits:
                original_text = edit.get("original_text", "")
                new_text = edit.get("new_text", "")
                
                if not original_text:
                    continue
                
                found = False
                
                # Search and replace in paragraphs
                for para in doc.paragraphs:
                    if original_text in para.text:
                        # Replace while preserving some formatting
                        for run in para.runs:
                            if original_text in run.text:
                                run.text = run.text.replace(original_text, new_text)
                                found = True
                        
                        # If not found in runs, replace in full paragraph
                        if not found and original_text in para.text:
                            para.text = para.text.replace(original_text, new_text)
                            found = True
                
                # Search and replace in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if original_text in cell.text:
                                for para in cell.paragraphs:
                                    if original_text in para.text:
                                        para.text = para.text.replace(original_text, new_text)
                                        found = True
                
                applied_edits.append({
                    "original": original_text[:50] + ("..." if len(original_text) > 50 else ""),
                    "new": new_text[:50] + ("..." if len(new_text) > 50 else ""),
                    "status": "applied" if found else "not_found"
                })
            
            # Save to bytes
            output = BytesIO()
            doc.save(output)
            
            return output.getvalue(), applied_edits
            
        except Exception as e:
            return content, [{"error": f"DOCX edit error: {str(e)}"}]
    
    # ═══════════════════════════════════════════════════════════════
    # MAIN EDIT FUNCTION
    # ═══════════════════════════════════════════════════════════════
    
    def edit_document(
        self,
        file,
        file_type: str,
        user_prompt: str,
        target_pages: Optional[List[int]] = None
    ) -> Dict[str, Any]:
        """
        Main function to edit a document based on user instructions.
        
        Args:
            file: Uploaded file object
            file_type: File extension (.pdf or .docx)
            user_prompt: User's edit instructions
            target_pages: Optional list of pages to focus on
            
        Returns:
            Dictionary with:
                - success: Boolean
                - edited_file: Bytes of edited document
                - file_name: Suggested filename
                - edits_applied: List of changes made
                - summary: AI summary of changes
                - error: Error message if failed
        """
        # Step 1: Read the document
        doc_content = self.read_document(file, file_type)
        
        if "error" in doc_content:
            return {
                "success": False,
                "error": doc_content["error"]
            }
        
        # Step 2: Analyze with AI what needs to be edited
        edits, summary = self.analyze_edit_request(doc_content, user_prompt, target_pages)
        
        if not edits:
            return {
                "success": False,
                "error": summary or "No se identificaron ediciones para aplicar"
            }
        
        # Step 3: Apply the edits
        raw_content = doc_content.get("raw_doc")
        file_type_clean = file_type.lower().replace('.', '')
        
        if file_type_clean == 'pdf':
            edited_bytes, applied = self.apply_edits_pdf(raw_content, edits)
        elif file_type_clean == 'docx':
            edited_bytes, applied = self.apply_edits_docx(raw_content, edits)
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {file_type}"
            }
        
        # Step 4: Generate output filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"MGA_editado_{timestamp}.{file_type_clean}"
        
        # Step 5: Save to output directory
        output_path = os.path.join(self.output_dir, output_filename)
        with open(output_path, 'wb') as f:
            f.write(edited_bytes)
        
        return {
            "success": True,
            "edited_file": edited_bytes,
            "file_path": output_path,
            "file_name": output_filename,
            "edits_applied": applied,
            "summary": summary,
            "page_count": doc_content.get("page_count", 0)
        }


# Convenience function for use in app.py
def edit_mga_document(
    file,
    user_prompt: str,
    llm=None,
    target_pages: Optional[List[int]] = None
) -> Dict[str, Any]:
    """
    Convenience function to edit an MGA document.
    
    Args:
        file: Uploaded file (Streamlit UploadedFile)
        user_prompt: User's edit instructions
        llm: LangChain LLM instance
        target_pages: Optional list of pages to focus on
        
    Returns:
        Result dictionary from MGAEditor.edit_document()
    """
    # Get file type from filename
    filename = file.name if hasattr(file, 'name') else str(file)
    file_type = filename.split('.')[-1].lower()
    
    editor = MGAEditor(llm=llm)
    return editor.edit_document(file, file_type, user_prompt, target_pages)
